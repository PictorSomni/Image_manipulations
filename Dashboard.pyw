# -*- coding: utf-8 -*-
"""
Tableau de bord principal de l'application Image Manipulations.

Ce module fournit une interface graphique Flet permettant de :
  - Sélectionner un dossier de travail et parcourir son contenu
    dans un panneau de prévisualisation avec chargement paresseux des miniatures.
  - Lancer les scripts de traitement d'images du répertoire ``Data/``
    en tant que sous-processus Python isolés, avec injection automatique des
    variables d'environnement requises (``FOLDER_PATH``, ``DATA_PATH``, etc.).
  - Afficher la sortie standard / erreur de chaque script dans un terminal
    intégré mis à jour en temps réel.
  - Effectuer des opérations sur fichiers : sélection par checkbox, copier/
    coller, suppression avec confirmation, création de dossier.

Raccourcis clavier :
  Ctrl/Cmd+A  — sélectionner tout / désélectionner tout.
  Ctrl/Cmd+C  — copier les fichiers sélectionnés dans le presse-papiers interne.
  Ctrl/Cmd+D  — sélectionner tous les fichiers de la même date que la référence.
  Ctrl/Cmd+↓ — basculer entre Terminal/Favoris et IA+Notes.
  Ctrl/Cmd+I  — inverser la sélection.
  Ctrl/Cmd+N  — créer un nouveau dossier.
  Ctrl/Cmd+R  — rafraîchir la prévisualisation.
  Ctrl/Cmd+V  — coller dans le dossier actuel.
    Ctrl/Cmd+↑ — agrandir/réduire le terminal.
    Ctrl/Cmd+← — IA en mode colonne (avec preview à droite).
    Ctrl/Cmd+→ — Bloc-notes en mode colonne (avec preview à droite).
    Ctrl/Cmd+Shift+← — IA en plein écran réel (moins la barre du haut).
    Ctrl/Cmd+Shift+→ — Bloc-notes en plein écran réel (moins la barre du haut).

Dépendances :
  flet >= 0.80, modules standard (os, subprocess, sys, platform, shutil,
  threading, re, zipfile, time).
"""

__version__ = "2.9.4"
overlay_fullscreen = {"mode": None}

# ==============================================================================
# TABLE DES MATIÈRES — Dashboard.pyw
# ==============================================================================
# 1. IMPORTS & CONFIGURATION ....................................... ~L 55
# 2. CONSTANTES .................................................... ~L 110
# 3. INTERFACE PRINCIPALE main() .................................. ~L 128
#    3.1  COULEURS ................................................. ~L 155
#    3.2  PROPRIÉTÉS & ÉTAT ....................................... ~L 171
#    3.3  ÉLÉMENTS UI ............................................. ~L 358
#    3.4  MÉTHODES ................................................ ~L 727
#    3.5  CONNEXIONS UI ........................................... ~L 7399
#    3.6  STRIP MODE (bandeau tactile) ............................. ~L 7806
#    3.7  INTERFACE FLET .......................................... ~L 7848
# ==============================================================================



#############################################################
#                          IMPORTS                          #
#############################################################
import flet as ft
import flet_code_editor as fce
import os
import subprocess
import sys
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "Data"))
import CONSTANTS
import platform
import shutil
import threading
import re
import zipfile
import json
import asyncio
import datetime
import concurrent.futures
import time
import hashlib
import urllib.request
import base64

try:
    from PIL import Image as _PILImage
except ImportError:
    _PILImage = None


from ai_tools import (
    _fetch_url_content, _web_search, _ollama_chat_once, _ollama_chat_stream,
    _ollama_chat_stream_with_tools, _gemini_chat_stream_with_tools,
    _parse_text_tool_calls, _strip_text_tool_calls,
    _format_ai_conversation, _folder_tool_definitions, _gemini_tool_definitions, _folder_list_contents,
    _folder_read_file, _folder_create_file, _folder_delete_files, _folder_move_file,
    _folder_copy_file, _folder_create_folder, _resolve_path,
    _folder_read_exif, _folder_zip_files, _folder_unzip_file,
    _encode_image_for_analysis, _analyze_images_batched, _take_screenshot,
    _gemini_generate_image, _gemini_refine_image_prompt, _gemini_generate_music,
    _WEB_TOOLS, _TERMINAL_TOOLS, _MEMORY_TOOLS, _SCREENSHOT_TOOLS, _NOTEPAD_TOOLS,
    _UI_TOOLS, _run_terminal_command,
    _EDIT_TOOLS, _READ_LINES_TOOLS, _SEARCH_TOOLS, _GIT_TOOLS, _TASK_TOOLS, _PDF_TOOLS, _SUBAGENT_TOOLS, _SCHEDULE_TOOLS,
    _HTTP_TOOLS, _SPREADSHEET_TOOLS, _PYAUTOGUI_TOOLS,
    _edit_file, _read_file_lines, _search_in_files, _find_files, _git_command, _manage_tasks, _read_pdf,
    _ask_subagent, _schedule_task, _http_request, _read_spreadsheet,
    _mouse_click, _keyboard_type, _keyboard_hotkey,
    _is_network_error,
    _update_memory_file, _build_system_content,
    _gemini_tts_stream, _gemini_live_tts_stream, _gemini_tts, _voice_play_audio,
    _claude_chat_stream_with_tools,
)
import thumb_cache
#############################################################
#                         CONSTANTS                         #
#############################################################
_IMAGE_VIEWER_EXTS = CONSTANTS.IMAGE_EXTS
_NOTEPAD_EXTS      = CONSTANTS.NOTEPAD_EXTS
_ANSI_ESCAPE       = re.compile(r'\x1b\[[0-9;]*[a-zA-Z]|\x1b\[\?[0-9;]*[a-zA-Z]')

_OS_JUNK = {
    ".ds_store", "thumbs.db", "thumbs.db:encryptable",
    "ehthumbs.db", "ehthumbs_vista.db", "desktop.ini",
    ".directory", ".spotlight-v100", ".trashes",
    ".thumbcache.db",
}



def _is_os_junk(entry):
    """Retourne True si l'entrée est un fichier système à ignorer."""
    name_lower = entry.name.lower()
    return (
        name_lower in _OS_JUNK
        or name_lower.startswith("._")
        or entry.name == "$RECYCLE.BIN"
        or (entry.name.startswith(".Trash-") and entry.is_dir())
    )



#############################################################
#                        -={MAIN}=-                         #
#############################################################
def main(page: ft.Page):
    """
    Point d'entrée Flet du Dashboard.

    Configure la fenêtre principale (titre, thème, dimensions, barre de titre
    personnalisée), initialise toutes les variables d'état, enregistre les
    canaux PubSub et construit l'interface avec trois zones :
      - Grille « Applications disponibles » (gauche).
      - Panneau de prévisualisation du contenu du dossier sélectionné (droite).
      - Terminal intégré affichant la sortie des scripts (bas).

    Parameters
    ----------
    page : ft.Page
        Objet page Flet injecté automatiquement par ``ft.run(main)``.

    ---
    PubSub (publish/subscribe) est un système de messagerie interne
    qui permet la communication thread-safe entre les threads de fond
    (scan, lecture de processus) et le thread UI de Flet :
    les threads publient des messages sur des "canaux" nommés
    (ex. "terminal", "refresh", "navigate"), et les callbacks abonnés les reçoivent
    et mettent à jour l'interface.
    ---
    """


# ===================== COULEURS ===================== #
    DARK        = CONSTANTS.COLOR_DARK
    BACKGROUND  = CONSTANTS.COLOR_BACKGROUND
    GREY        = CONSTANTS.COLOR_GREY
    LIGHT_GREY  = CONSTANTS.COLOR_LIGHT_GREY
    BLUE        = CONSTANTS.COLOR_BLUE
    VIOLET      = CONSTANTS.COLOR_VIOLET
    GREEN       = CONSTANTS.COLOR_GREEN
    YELLOW      = CONSTANTS.COLOR_YELLOW
    HOVER_YELLOW = CONSTANTS.COLOR_HOVER_YELLOW
    ORANGE      = CONSTANTS.COLOR_ORANGE
    RED         = CONSTANTS.COLOR_RED
    WHITE       = CONSTANTS.COLOR_WHITE



# ===================== PROPRIÉTÉS ===================== #
    page.title = "Dashboard - Image Manipulations"
    page.theme_mode = ft.ThemeMode.DARK
    page.bgcolor = BACKGROUND
    page.window.title_bar_hidden = True
    page.window.title_bar_buttons_hidden = True
    page.window.width = CONSTANTS.WINDOW_WIDTH
    page.window.height = CONSTANTS.WINDOW_HEIGHT
    if platform.system() == "Darwin" and CONSTANTS.MAXIMIZED:
        page.window.maximized = False
    else:
        page.window.maximized = CONSTANTS.MAXIMIZED
    page.window.icon = "assets/icon.png"



    async def on_window_event(event):
        if event.data == "close":
            if note_mode["value"]:
                try:
                    with open(note_target_file["path"], "w", encoding="utf-8") as _f:
                        _f.write(notepad_field.value or "")
                except Exception:
                    pass
            proc = ollama_process["proc"]
            if proc is not None and proc.poll() is None:
                try:
                    proc.terminate()
                except Exception:
                    pass
        elif event.data in ("resize", "maximize", "unmaximize"):
            if bottom_panel_container is None:
                return
            if overlay_fullscreen["mode"] in ("ai", "notepad", "ai_full", "notepad_full", "both_full"):
                if overlay_fullscreen["mode"] in ("ai_full", "notepad_full", "both_full"):
                    bottom_panel_container.left = 0
                    bottom_panel_container.right = 0
                    bottom_panel_container.width = None
                    bottom_panel_container.top = 0
                    bottom_panel_container.height = None
                else:
                    win_w = page.window.width or CONSTANTS.WINDOW_WIDTH
                    bottom_panel_container.width = int((win_w - 8) * 6 / 15 + 4)
                try:
                    bottom_panel_container.update()
                except Exception:
                    pass



    page.window.on_event = on_window_event
    selected_folder = {"path": None}
    current_browse_folder = {"path": None}
    kiosk_tariff = {"value": "PRINTS"}  # Tarif kiosk actif : "STUDIOS" ou "PRINTS"
    app_directory = os.path.dirname(os.path.abspath(__file__))
    selected_files = []  # Liste des fichiers sélectionnés (ordre de clic préservé)

    clipboard = {"files": [], "cut": False}  # Presse-papiers pour copier/coller/couper des fichiers



    # ── Dossiers récents ──────────────────────────────────────────────
    recent_folders_file_path = os.path.join(app_directory, ".recent_folders.json")



    def _load_recent() -> list:
        try:
            with open(recent_folders_file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return [p for p in data if os.path.isdir(p)]
        except Exception:
            return []



    def _save_recent(folders: list) -> None:
        try:
            with open(recent_folders_file_path, "w", encoding="utf-8") as f:
                json.dump(folders[:10], f, ensure_ascii=False, indent=2)
        except Exception:
            pass



    def _add_to_recent(path: str) -> None:
        path = os.path.normpath(path)
        recent = _load_recent()
        if path in recent:
            recent.remove(path)
        recent.insert(0, path)
        _save_recent(recent[:10])



    # ── Dossiers favoris ──────────────────────────────────────────────
    favorites_file_path = os.path.join(app_directory, ".favorites.json")



    # ── Programmes "Ouvrir avec" ──────────────────────────────────────
    open_with_config_file_path = os.path.join(app_directory, "open_with.json")



    def _load_favorites() -> list:
        try:
            with open(favorites_file_path, "r", encoding="utf-8") as f:
                raw = json.load(f)
            result = []
            for item in raw:
                if isinstance(item, str):
                    result.append({"path": item, "label": ""})
                elif isinstance(item, dict) and "path" in item:
                    result.append({"path": item["path"], "label": item.get("label", "")})
            return result
        except Exception:
            return []



    def _save_favorites(folders: list) -> None:
        try:
            with open(favorites_file_path, "w", encoding="utf-8") as f:
                json.dump(folders, f, ensure_ascii=False, indent=2)
        except Exception:
            pass



    def _resolve_favorite_path(p: str) -> str:
        """Sur macOS, résout /Volumes/NOM vers /Volumes/NOM-1 si nécessaire.
        Cas 1 : NOM n'existe pas → cherche NOM-1, -2…
        Cas 2 : NOM existe mais NOM-1 aussi → NOM est un stub, préférer NOM-1."""
        if sys.platform != "darwin":
            return p
        if not p.startswith("/Volumes/"):
            return p
        rest = p[len("/Volumes/"):]
        vol_name = rest.split("/")[0]
        sub_path = rest[len(vol_name):]
        # Cherche le variant -N avec le numéro le plus élevé qui existe
        for suffix in ["-1", "-2", "-3", "-4"]:
            candidate_vol = f"/Volumes/{vol_name}{suffix}"
            candidate = f"{candidate_vol}{sub_path}"
            if os.path.isdir(candidate_vol):
                return candidate
        # Aucun variant -N : retourne le chemin original (existant ou non)
        return p



    # Configuration: nom du fichier -> True si l'app est locale (pas besoin de dossier sélectionné)
    apps = {
        "Transfert vers TEMP.py": (True, BLUE),
        "Conversion JPG.py": (False, BLUE),
        "Renommer séquence.py": (False, BLUE),
        "Comparaison.pyw": (False, VIOLET, os.path.join(os.path.dirname(os.path.abspath(__file__)), "Data", "Comparaison.pyw")),
        "Recadrage automatique.py": (False, GREEN, os.path.join(os.path.dirname(os.path.abspath(__file__)), "Data", "Recadrage automatique.py")),
        "Recadrage manuel.pyw": (False, RED, os.path.join(os.path.dirname(os.path.abspath(__file__)), "Data", "Recadrage manuel.pyw")),
        "Fichiers identiques.py": (False, VIOLET),
        "Débruiter.py": (False, BLUE),
        "Grain pellicule.py": (False, YELLOW),
        "Redimensionner filigrane.py": (False, WHITE),
        "2 en 1.py": (False, HOVER_YELLOW),
        "Redimensionner.py": (False, WHITE),
        "Augmentation IA.py": (False, YELLOW),
        "Copyright.py": (False, VIOLET),
        "IA / Bloc-notes": (True, BLUE),
    }



    # ===================== Valeurs par défaut ===================== #

    resize_size = {"value": str(CONSTANTS.RESIZE_DEFAULT)}  # Taille par défaut pour le redimensionnement
    resize_watermark_size = {"value": str(CONSTANTS.RESIZE_DEFAULT)}  # Taille par défaut pour le redimensionnement avec watermark
    sort_mode = {"value": 2}  # 0 = A→Z, 1 = Z→A, 2 = par date de modification
    show_only_selection = {"value": False}  # True = afficher uniquement les fichiers sélectionnés
    removable_drives_state = {"list": []}  # [(name, path), ...]
    _image_cache_busters = {}  # {normpath: temp_path_unique} pour invalider le cache navigateur
    _image_last_mtime = {}     # {normpath: mtime} pour détecter les modifications externes
    _checkbox_refs = {}        # {file_path: ft.Checkbox} — refs aux checkboxes rendues (mise à jour in-place)
    _thumb_cache = {}          # {normpath: b64_string} — miniatures PIL générées (chargement asynchrone)
    _pending_thumb_refs = {}   # {normpath: (ft.Container, file_path, icon, icon_color)} — widgets en attente
    PAGE_SIZE = 100             # Nb d'éléments max par page dans la prévisualisation
    preview_page = {"value": 0}  # Page courante (0-indexé)
    all_entries_data = {"list": [], "error": ""}  # Données brutes du dernier scan
    pending_file_selection = {"names": None}  # Noms à sélectionner après le prochain scan
    preview_refresh_token = {"value": 0}   # incrémenté à chaque refresh pour annuler les anciens threads
    search_query = {"value": ""}  # Requête de recherche active dans la preview
    command_history = []           # Historique des commandes du terminal
    history_index = {"value": -1}  # -1 = nouvelle saisie en cours
    history_draft = {"value": ""}  # Saisie en cours avant navigation dans l'historique
    terminal_is_expanded = {"value": False}
    terminal_input_focused = {"value": False}
    keyboard_shortcuts_suspended = {"count": 0}
    _solo_left_state   = {"container": None}   # Référence au conteneur solo (mode pleine hauteur)
    ai_mode            = {"value": False}
    ai_conversation    = []              # Historique de conversation [{role, content}]
    ai_streaming       = {"value": False}
    ollama_process     = {"proc": None}  # Process ollama serve lancé par nous
    ai_pending_images  = []              # Images jointes en attente [{path, b64}]
    ai_pending_files   = []              # Documents joints en attente [path]
    _live_print_counts = {}     # {file_path: int} — cache local/live des nombres d'impressions avant renommage disque
    _print_rename_timers = {}   # {file_path: threading.Timer} — timers pour débouncer le renommage disque
    _print_count_text_refs = {} # {file_path: ft.Text}
    _print_minus_btn_refs = {}  # {file_path: ft.Container}



    def _generate_thumbnail(file_path):
        """
        Génère une miniature via le cache persistant thumb_cache.
        Retourne une chaîne base64, ou None en cas d'échec.
        """
        return thumb_cache.get_or_generate(file_path)



# ===================== ÉLÉMENTS UI ===================== #
    def _short_path(p, max_len=35):
        if not p or len(p) <= max_len:
            return p or ""
        return "…" + p[-(max_len - 1):]

    folder_path = ft.TextField(
        label="Dossier sélectionné",
        hint_text="Cliquez sur Parcourir... ou collez un chemin",
        width=300,
        bgcolor=DARK,
        border_color=GREY,
    )
    recent_folders_btn = ft.PopupMenuButton(
        icon=ft.Icons.HISTORY,
        icon_color=LIGHT_GREY,
        tooltip="Dossiers récents",
        items=[],
    )
    apps_list = ft.Column(expand=True, spacing=8)
    quick_tools_col = ft.Column([], spacing=10, horizontal_alignment=ft.CrossAxisAlignment.CENTER)
    preview_list = ft.ListView(expand=True, auto_scroll=False, spacing=4)
    preview_loading = ft.Container(
        content=ft.Row([
            ft.ProgressRing(width=16, height=16, stroke_width=2, color=BLUE),
            ft.Text("Chargement...", size=12, color=LIGHT_GREY),
        ], spacing=8, alignment=ft.MainAxisAlignment.CENTER),
        alignment=ft.Alignment(0, 0),
        expand=True,
        visible=False,
    )



    terminal_output = ft.ListView(expand=True, spacing=2, auto_scroll=True)
    app_progress_bar = ft.ProgressBar(value=None, visible=False, color=GREEN, height=2)
    terminal_cmd_input = ft.TextField(
        hint_text="> Terminal",#  (tapez /note pour ouvrir le bloc-notes)",
        border_color=GREEN,
        text_style=ft.TextStyle(font_family="monospace", size=CONSTANTS.TERMINAL_FONT_SIZE),
        dense=True,
        expand=True,
        color=GREEN,
        on_submit=lambda e: on_terminal_command_submit(e),
        on_focus=lambda e: terminal_input_focused.update({"value": True}),
        on_blur=lambda e: terminal_input_focused.update({"value": False}),
    )
    terminal_cmd_row = ft.Row([terminal_cmd_input])

    # ── Bloc-notes ────────────────────────────────────────────────────
    notes_file_path      = os.path.join(app_directory, ".notes.md")
    constants_file_path  = os.path.join(app_directory, "Data", "CONSTANTS.py")
    ai_history_file_path = os.path.join(app_directory, ".ai_conversation.json")
    note_mode            = {"value": False}
    note_target_file     = {"path": notes_file_path}

    _HAS_CODE_EDITOR = platform.system() != "Linux"

    if _HAS_CODE_EDITOR:
        notepad_field = fce.CodeEditor(
            text_style=ft.TextStyle(font_family="monospace", size=CONSTANTS.TERMINAL_FONT_SIZE),
            language=fce.CodeLanguage.PYTHON,
            code_theme=fce.CodeTheme.ATOM_ONE_DARK,
            gutter_style=fce.GutterStyle(width=85),
            expand=True,
        )
    else:
        notepad_field = ft.TextField(
            multiline=True,
            expand=True,
            min_lines=4,
            text_style=ft.TextStyle(font_family="monospace", size=CONSTANTS.TERMINAL_FONT_SIZE),
            color=WHITE,
            border_color=ft.Colors.TRANSPARENT,
            border_radius=6,
            bgcolor=DARK,
            filled=True,
            hint_text="Écrivez vos notes ici…",
            hint_style=ft.TextStyle(color=LIGHT_GREY, italic=True),
        )
    notepad_is_preview       = {"value": False}
    notepad_markdown_preview = ft.Markdown(
        "",
        selectable=True,
        extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,

        code_theme=ft.MarkdownCodeTheme.ATOM_ONE_DARK,

        md_style_sheet=ft.MarkdownStyleSheet(
            blockquote_text_style=ft.TextStyle(color=WHITE), # Texte blanc
            blockquote_decoration=ft.BoxDecoration(
                bgcolor=GREY, # Fond sombre au lieu du bleu clair
                border=ft.Border.all(1, BLUE),
                border_radius=5,
            ),
        ),
        expand=True,
    )
    notepad_preview_scroll = ft.ListView(
        controls=[notepad_markdown_preview],
        expand=True,
        visible=False,
    )
    _notepad_autosave_timer = {"task": None}

    notepad_header_icon  = ft.Icon(ft.Icons.EDIT_NOTE, color=VIOLET, size=16)
    notepad_header_title = ft.Text("Notes", color=VIOLET, size=12, weight=ft.FontWeight.BOLD)

    expand_button_terminal = ft.IconButton(
        icon=ft.Icons.VERTICAL_SPLIT,
        tooltip="Agrandir  (Ctrl+↑)",
        icon_color=YELLOW,
        icon_size=16,
        on_click=lambda e: toggle_terminal_overlay(),
    )
    expand_button_overlay = ft.IconButton(
        icon=ft.Icons.OPEN_IN_FULL,
        tooltip="IA seule (Ctrl/Cmd+←)",
        icon_color=BLUE,
        icon_size=16,
        on_click=lambda e: toggle_ai_fullscreen(),
    )
    expand_button_notepad = ft.IconButton(
        icon=ft.Icons.OPEN_IN_FULL,
        tooltip="Bloc-notes seul (Ctrl/Cmd+→)",
        icon_color=VIOLET,
        icon_size=16,
        on_click=lambda e: toggle_notepad_fullscreen(),
    )

    notepad_container = ft.Container(
        content=ft.Column([
            notepad_field,
            notepad_preview_scroll,
        ], spacing=4, expand=True),
        expand=True,
        visible=True,
        bgcolor=DARK,
    )



    # ── Conversation IA ───────────────────────────────────────────────
    ai_chat_view = ft.ListView(expand=True, spacing=4, auto_scroll=True)
    ai_input_field = ft.TextField(
        hint_text="Posez votre question… (Entrée pour envoyer)",
        border_color=BLUE,
        text_style=ft.TextStyle(font_family="monospace", size=CONSTANTS.TERMINAL_FONT_SIZE),
        dense=True,
        expand=True,
        color=WHITE,
        bgcolor=DARK,
        shift_enter=True,
        on_submit=lambda e: _on_ai_submit(),
    )
    ai_model_dropdown = ft.Dropdown(
        value=CONSTANTS.AI_MODEL_TEXT,
        options=[ft.dropdown.Option(model) for model in CONSTANTS.AI_DROPDOWN_MODELS],
        text_size=11,
        dense=True,
        color=LIGHT_GREY,
        bgcolor=DARK,
        border_color=GREY,
        content_padding=ft.Padding.symmetric(horizontal=6, vertical=0),
        width=150,
    )
    ai_status_text    = ft.Text("", color=LIGHT_GREY, size=11, italic=True)
    ai_progress_bar  = ft.ProgressBar(value=None, visible=False, color=BLUE, height=2)
    ai_stop_button   = ft.IconButton(
        icon=ft.Icons.STOP_CIRCLE,
        icon_color=LIGHT_GREY,
        icon_size=16,
        tooltip="Libérer le modèle (ollama stop)",
        on_click=lambda e: _ai_stop_model(),
    )
    ai_attach_row    = ft.Row([], spacing=4, visible=False, wrap=True)
    ai_send_button   = ft.IconButton(
        icon=ft.Icons.SEND,
        icon_color=BLUE,
        icon_size=18,
        tooltip="Envoyer",
        on_click=lambda e: _on_ai_submit(),
    )
    ai_attach_button = ft.IconButton(
        icon=ft.Icons.ATTACH_FILE,
        icon_color=LIGHT_GREY,
        icon_size=18,
        tooltip="Joindre une image ou un document",
        on_click=lambda e: page.run_task(_ai_pick_any),
    )
    ai_tts_enabled = {"value": CONSTANTS.AI_VOICE_TTS_ENABLED}
    ai_send_original_images = {"value": CONSTANTS.AI_IMAGE_ATTACH_DEFAULT_ORIGINAL}
    ai_tts_stop_event = {"event": None}
    ai_speaker_button = ft.IconButton(
        icon=ft.Icons.VOLUME_UP if CONSTANTS.AI_VOICE_TTS_ENABLED else ft.Icons.VOLUME_OFF,
        icon_color=CONSTANTS.COLOR_BLUE if CONSTANTS.AI_VOICE_TTS_ENABLED else CONSTANTS.COLOR_LIGHT_GREY,
        icon_size=18,
        tooltip="Désactiver la lecture vocale" if CONSTANTS.AI_VOICE_TTS_ENABLED else "Activer la lecture vocale",
        visible=CONSTANTS.AI_VOICE_TTS_BTN_VISIBLE,
        on_click=lambda e: _toggle_tts(),
    )
    ai_image_mode_label = ft.Text(
        "REEL" if CONSTANTS.AI_IMAGE_ATTACH_DEFAULT_ORIGINAL else "1024",
        color=GREEN if CONSTANTS.AI_IMAGE_ATTACH_DEFAULT_ORIGINAL else BLUE,
        size=10,
        weight=ft.FontWeight.BOLD,
    )
    ai_image_size_button = ft.IconButton(
        icon=ft.Icons.IMAGE,
        icon_color=GREEN if CONSTANTS.AI_IMAGE_ATTACH_DEFAULT_ORIGINAL else BLUE,
        icon_size=16,
        tooltip=(
            "Mode images IA en taille réelle (fichier original) — affecte uniquement les nouveaux fichiers joints"
            if CONSTANTS.AI_IMAGE_ATTACH_DEFAULT_ORIGINAL
            else "Mode images IA optimisé (1024px max) — affecte uniquement les nouveaux fichiers joints"
        ),
        on_click=lambda e: _toggle_ai_image_size_mode(),
    )

    ai_container = ft.Container(
        content=ft.Column([
            ai_chat_view,
            ai_attach_row,
            ai_progress_bar,
            ft.Row(
                [ai_attach_button, ai_input_field, ai_send_button],
                spacing=4,
                vertical_alignment=ft.CrossAxisAlignment.CENTER,
            ),
        ], spacing=4, expand=True),
        expand=True,
        visible=True,
        bgcolor=DARK,
    )


    file_count_text = ft.Text("", size=14, color=WHITE, text_align=ft.TextAlign.RIGHT)
    selection_count_text = ft.Text("", size=14, color=BLUE, text_align=ft.TextAlign.RIGHT)
    select_toggle_button = ft.IconButton(
        icon=ft.Icons.SELECT_ALL,
        icon_color=VIOLET,
        icon_size=22,
        tooltip="Tout sélectionner",
    )



    invert_selection_button = ft.IconButton(
        icon=ft.Icons.PUBLISHED_WITH_CHANGES,
        icon_color=VIOLET,
        icon_size=22,
        tooltip="Inverser la sélection",
    )



    select_same_date_button = ft.IconButton(
        icon=ft.Icons.EVENT,
        icon_color=VIOLET,
        icon_size=22,
        tooltip="Sélectionner même date",
    )



    filter_sel_btn = ft.IconButton(
        icon=ft.Icons.FILTER_LIST,
        icon_color=LIGHT_GREY,
        icon_size=22,
        tooltip="Afficher uniquement la sélection",
    )



    sort_segment = ft.CupertinoSlidingSegmentedButton(
        selected_index=2,
        bgcolor=GREY,
        thumb_color=DARK,
        controls=[
            ft.Text("A→Z",  size=11, color=WHITE),
            ft.Text("Z→A",  size=11, color=WHITE),
            ft.Text("Date", size=11, color=WHITE),
        ],
        tooltip="Tri alphabétique (A→Z / Z→A) ou par date de modification",
    )



    prev_page_btn = ft.IconButton(
        icon=ft.Icons.CHEVRON_LEFT, icon_size=18, icon_color=DARK, bgcolor=YELLOW,
        tooltip="Page précédente", visible=False, hover_color=HOVER_YELLOW
    )



    next_page_btn = ft.IconButton(
        icon=ft.Icons.CHEVRON_RIGHT, icon_size=18, icon_color=DARK, bgcolor=YELLOW,
        tooltip="Page suivante", visible=False, hover_color=HOVER_YELLOW
    )



    page_indicator_text = ft.Text("", size=12, color=LIGHT_GREY)
    selected_files_prefix = "SELECTED_FILES:"



    # ── Barre de recherche dans la preview ────────────────────────────
    search_field = ft.TextField(
        hint_text="Rechercher...",
        border_color=BLUE,
        text_size=13,
        height=45,
        width=180,
        content_padding=ft.Padding(8, 2, 8, 2),
        prefix_icon=ft.Icons.SEARCH,
        bgcolor=DARK,
    )
    search_close_btn = ft.IconButton(
        icon=ft.Icons.CLOSE,
        icon_color=LIGHT_GREY,
        icon_size=16,
        tooltip="Fermer la recherche",
        style=ft.ButtonStyle(padding=ft.Padding.all(4)),
    )
    search_active_row = ft.Row(
        [search_field, search_close_btn],
        spacing=0, tight=True,
        vertical_alignment=ft.CrossAxisAlignment.CENTER,
    )



    # ── Champs de saisie Redimensionner ──────────────────────────────
    resize_input = ft.TextField(
        value=str(CONSTANTS.RESIZE_DEFAULT),
        width=80,
        height=35,
        text_size=13,
        text_align=ft.TextAlign.CENTER,
        keyboard_type=ft.KeyboardType.NUMBER,
        border_color=BLUE,
        content_padding=ft.Padding(5, 5, 5, 5),
    )
    resize_watermark_input = ft.TextField(
        value=str(CONSTANTS.RESIZE_DEFAULT),
        width=80,
        height=35,
        text_size=13,
        text_align=ft.TextAlign.CENTER,
        keyboard_type=ft.KeyboardType.NUMBER,
        border_color=ORANGE,
        content_padding=ft.Padding(5, 5, 5, 5),
    )



    # ── Section favoris ──────────────────────────────────────────────
    favorites_list_view = ft.ReorderableListView(expand=True, spacing=2, auto_scroll=False, padding=ft.Padding(12, 6, 12, 6), show_default_drag_handles=False)
    favorites_panel = ft.Container(
        content=ft.Column([
            ft.Row([
                ft.Icon(ft.Icons.STAR, size=14, color=BLUE),
                ft.Text("Favoris", size=14, color=BLUE, weight=ft.FontWeight.BOLD),
                ft.Container(expand=True),
                ft.Container(
                    content=ft.Icon(ft.Icons.ADD, size=13, color=DARK),
                    bgcolor=BLUE,
                    border_radius=10,
                    padding=ft.Padding(3, 1, 3, 1),
                    tooltip="Ajouter le dossier courant aux favoris",
                    on_click=lambda e: _add_favorite_current(),
                    ink=True,
                ),
            ], spacing=6, tight=True, vertical_alignment=ft.CrossAxisAlignment.CENTER),
            favorites_list_view,
        ], spacing=4, expand=True),
        bgcolor=GREY,
        border=ft.Border.all(1, BLUE),
        border_radius=6,
        padding=ft.Padding(12, 6, 12, 6),
        expand=True,
    )



    # ── Section périphériques amovibles ──────────────────────────────
    drives_list_view = ft.ReorderableListView(expand=True, spacing=4, auto_scroll=False, padding=ft.Padding(12, 6, 12, 6), show_default_drag_handles=False)
    drives_panel = ft.Container(
        content=ft.Column([
            ft.Row([
                ft.Icon(ft.Icons.USB, size=14, color=VIOLET),
                ft.Text("Périphériques détectés", size=14, color=VIOLET,
                        weight=ft.FontWeight.BOLD),
            ], spacing=6, tight=True),
            drives_list_view,
        ], spacing=4, expand=True),
        bgcolor=GREY,
        border=ft.Border.all(1, VIOLET),
        border_radius=6,
        padding=ft.Padding(12, 6, 12, 6),
        expand=True,
        visible=False,
    )



# ===================== MÉTHODES ===================== #
    # ================================================================ #
    #                    PUBSUB & ÉVÉNEMENTS                           #
    # ================================================================ #
    # (terminal géré directement dans log_to_terminal)



    def on_refresh_preview(topic, message):
        """Callback pour rafraîchir la preview depuis un thread"""
        refresh_preview(reset_page=False)
    
    # S'abonner au canal refresh
    page.pubsub.subscribe_topic("refresh", on_refresh_preview)



    def on_navigate_request(topic, folder_path):
        """Callback pour naviguer vers un dossier depuis un thread"""
        if folder_path and os.path.isdir(folder_path):
            navigate_to_folder(folder_path)
    
    # S'abonner au canal navigate
    page.pubsub.subscribe_topic("navigate", on_navigate_request)



    def on_select_files_request(topic, selected_names_str):
        """Callback pour sélectionner des fichiers depuis la sortie d'un script"""
        apply_selected_files_by_name(selected_names_str)

    # S'abonner au canal select-files
    page.pubsub.subscribe_topic("select_files", on_select_files_request)



    def on_quit_request(topic, message):
        """Callback pour fermer la fenêtre depuis un thread de fond (thread-safe)"""
        page.window.close()

    # S'abonner au canal quit
    page.pubsub.subscribe_topic("quit", on_quit_request)



    def on_deselect_request(topic, message):
        """Callback pour désélectionner tous les fichiers depuis un thread de fond."""
        selected_files.clear()
        if show_only_selection["value"]:
            show_only_selection["value"] = False
            _update_filter_sel_btn()
            _render_preview_page()
        else:
            _update_visible_checkboxes()

    # S'abonner au canal deselect
    page.pubsub.subscribe_topic("deselect", on_deselect_request)

    dashboard_window_cycle_config = {
        "SidePanel.pyw": {"restore_previous_maximized_state": True},
        "kiosk_flet.pyw": {"restore_previous_maximized_state": True},
        "Comparaison.pyw": {"restore_previous_maximized_state": True},
        "Recadrage manuel.pyw": {"restore_previous_maximized_state": True},
        "Augmentation IA.py": {"restore_previous_maximized_state": True},
    }



    def _get_dashboard_window_cycle_options(app_name: str):
        """Retourne la configuration de cycle fenêtre pour une app, ou None."""
        return dashboard_window_cycle_config.get(app_name)



    def _restore_dashboard_window(previous_maximized_state: bool = None):
        """Restaure Dashboard (dé-minimisation, état maximisé optionnel, puis premier plan)."""
        page.window.minimized = False

        if previous_maximized_state is not None:
            if previous_maximized_state and platform.system() == "Darwin":
                page.window.maximized = False
                page.update()
                time.sleep(0.05)
                page.window.maximized = True
            else:
                page.window.maximized = bool(previous_maximized_state)

        page.run_task(page.window.to_front)
        page.update()



    def _launch_with_dashboard_restore(
        command: list,
        env: dict,
        *,
        restore_previous_maximized_state: bool = False,
        on_exit_topic: str | None = None,
        popen_kwargs: dict = None,
    ):
        """Lance une app externe en minimisant Dashboard, puis restaure Dashboard à la fermeture."""
        process_kwargs = dict(popen_kwargs or {})
        restore_maximized_state = bool(page.window.maximized) if restore_previous_maximized_state else None
        process = subprocess.Popen(command, env=env, **process_kwargs)
        page.window.minimized = True
        page.update()

        def _watch_process_closure():
            process.wait()
            if restore_previous_maximized_state:
                _restore_dashboard_window(bool(restore_maximized_state))
            else:
                _restore_dashboard_window()
            if on_exit_topic:
                page.pubsub.send_all_on_topic(on_exit_topic, None)

        threading.Thread(target=_watch_process_closure, daemon=True).start()



    def _launch_side_panel(extra_env: dict = None):
        """Lance Side Panel, minimise Dashboard, puis le restaure à la fermeture de Side Panel."""
        cycle_options = _get_dashboard_window_cycle_options("SidePanel.pyw") or {}
        env = {
            **os.environ,
            "SELECTEUR_INITIAL_FOLDER": (
                current_browse_folder["path"] or selected_folder["path"] or ""
            ),
        }
        if extra_env:
            env.update(extra_env)

        _launch_with_dashboard_restore(
            [sys.executable, os.path.join(app_directory, "Data", "SidePanel.pyw")],
            env,
            restore_previous_maximized_state=bool(cycle_options.get("restore_previous_maximized_state", False)),
            popen_kwargs={
                "stdout": subprocess.DEVNULL,
                "stderr": subprocess.DEVNULL,
            },
        )



    def _launch_comparaison(second_folder: str = ""):
        """Lance Comparaison.pyw, minimise Dashboard, puis le restaure à la fermeture."""
        browse = current_browse_folder["path"] or ""
        base   = selected_folder["path"] or ""
        selected_image_files = [
            file_path for file_path in selected_files
            if os.path.isfile(file_path)
            and os.path.splitext(file_path)[1].lower() in CONSTANTS.IMAGE_EXTS
        ]

        # Dossier 1 : dossier de la paire sélectionnée si possible,
        # sinon dossier courant de navigation (ou le dossier sélectionné).
        folder1 = browse or base
        if len(selected_image_files) == 2:
            folder1 = os.path.normpath(os.path.dirname(selected_image_files[0]))

        if not folder1:
            log_to_terminal("[ERREUR] Veuillez sélectionner un dossier avant de lancer la Comparaison", RED)
            return

        # Dossier 2 : si browse et base sont distincts, utiliser base comme dossier 2
        if browse and base and os.path.normpath(browse) != os.path.normpath(base):
            folder2 = base
        else:
            folder2 = second_folder

        def _do_launch(f2: str):
            cycle_options = _get_dashboard_window_cycle_options("Comparaison.pyw") or {}
            env = {**os.environ, "FOLDER_PATH": folder1}
            if f2:
                env["SECOND_FOLDER"] = f2
            # Si des fichiers sont sélectionnés dans folder1, les transmettre.
            # Cas spécial: 2 images exactement -> comparaison directe de cette paire.
            if len(selected_image_files) == 2:
                env["SELECTED_PAIR_FILES"] = "|".join(os.path.basename(f) for f in selected_image_files)
                env["SELECTED_PAIR_PATHS"] = "|".join(selected_image_files)
            else:
                files_in_folder1 = [
                    f for f in selected_files
                    if os.path.isfile(f) and os.path.normpath(os.path.dirname(f)) == os.path.normpath(folder1)
                ]
                image_files_in_folder1 = [
                    f for f in files_in_folder1
                    if os.path.splitext(f)[1].lower() in CONSTANTS.IMAGE_EXTS
                ]

                if image_files_in_folder1:
                    env["SELECTED_FILES"] = "|".join(os.path.basename(f) for f in image_files_in_folder1)
            comparaison_path = os.path.join(app_directory, "Data", "Comparaison.pyw")
            _launch_with_dashboard_restore(
                [sys.executable, comparaison_path],
                env,
                restore_previous_maximized_state=bool(cycle_options.get("restore_previous_maximized_state", False)),
                on_exit_topic="refresh",
            )

        if len(selected_image_files) == 2:
            _do_launch("")
            return

        # Si un ou deux dossiers sont sélectionnés dans la preview_list, les utiliser
        selected_dirs = [f for f in selected_files if os.path.isdir(f)]
        if len(selected_dirs) >= 2 and not folder2:
            folder1 = os.path.normpath(selected_dirs[0])
            folder2 = os.path.normpath(selected_dirs[1])
        elif len(selected_dirs) == 1 and not folder2:
            folder2 = os.path.normpath(selected_dirs[0])

        # Si le second dossier est déjà connu, lancer directement
        if folder2:
            _do_launch(folder2)
            return
        # Sinon, ouvrir un sélecteur de dossier dans Dashboard pour choisir le second dossier
        async def _pick_and_launch(e=None):
            picked = await ft.FilePicker().get_directory_path(
                dialog_title="Sélectionner le second dossier à comparer",
                initial_directory=folder1)
            if picked:
                _do_launch(os.path.normpath(picked))
            else:
                log_to_terminal("[INFO] Comparaison annulée (pas de second dossier sélectionné)", LIGHT_GREY)

        page.run_task(_pick_and_launch)



    def _launch_kiosk_flet():
        """Lance kiosk_flet.pyw avec le cycle fenêtre standard (comme SidePanel)."""
        cycle_options = _get_dashboard_window_cycle_options("kiosk_flet.pyw") or {}
        folder = current_browse_folder["path"] or selected_folder["path"] or ""
        env = {
            **os.environ,
            "FOLDER_PATH": folder,
            "TARIFF_TYPE": kiosk_tariff["value"],
        }
        kiosk_path = os.path.join(app_directory, "Data", "kiosk_flet.pyw")
        _launch_with_dashboard_restore(
            [sys.executable, kiosk_path],
            env,
            restore_previous_maximized_state=bool(cycle_options.get("restore_previous_maximized_state", False)),
            on_exit_topic="refresh",
            popen_kwargs={
                "stdout": subprocess.DEVNULL,
                "stderr": subprocess.DEVNULL,
            },
        )



    _kiosk_tariff_label = ft.Text("PRINTS", size=12, color=DARK)
    kiosk_tariff_btn = ft.Button(
        content=_kiosk_tariff_label,
        bgcolor=GREEN,
        style=ft.ButtonStyle(
            padding=ft.Padding.symmetric(horizontal=10, vertical=2),
            shape=ft.StadiumBorder(),
        ),
        height=40,
        width=76,
        tooltip="Tarif kiosk actif — cliquer pour changer (STUDIOS / PRINTS)",
    )



    def _toggle_kiosk_tariff(e) -> None:
        kiosk_tariff["value"] = "PRINTS" if kiosk_tariff["value"] == "STUDIOS" else "STUDIOS"
        if kiosk_tariff["value"] == "STUDIOS":
            _kiosk_tariff_label.value = "STUDIOS"
            kiosk_tariff_btn.bgcolor = YELLOW
        else:
            _kiosk_tariff_label.value = "PRINTS"
            kiosk_tariff_btn.bgcolor = GREEN
        kiosk_tariff_btn.update()

    kiosk_tariff_btn.on_click = _toggle_kiosk_tariff


    def on_preview_ready(topic, payload):
        """Reçoit les données brutes du thread bg et déclenche le rendu de la page courante."""
        token, entries_data, new_file_count_text, error_text = payload
        if preview_refresh_token["value"] != token:
            return
        all_entries_data["list"] = entries_data
        all_entries_data["error"] = error_text
        file_count_text.value = new_file_count_text
        # Appliquer la sélection en attente si un script l'a demandé
        # (ex: Fichiers manquants). On le fait ICI avec les données fraîches
        # pour éviter tout race condition entre threads.
        if pending_file_selection["names"] is not None:
            names_to_apply = pending_file_selection["names"]
            pending_file_selection["names"] = None
            apply_selected_files_by_name(names_to_apply)
        else:
            _render_preview_page()

    # S'abonner au canal preview_ready
    page.pubsub.subscribe_topic("preview_ready", on_preview_ready)



    def _start_thumb_loader():
        """Lance un thread qui génère les miniatures PIL pour la page courante."""
        pending_snapshot = list(_pending_thumb_refs.items())
        load_token = preview_refresh_token["value"]

        def _load():
            _ts = CONSTANTS.DASHBOARD_THUMB_SIZE
            for norm_path, (img_ref, file_path, icon, icon_color) in pending_snapshot:
                if preview_refresh_token["value"] != load_token:
                    return  # Navigation survenue, annuler
                thumb = _generate_thumbnail(file_path)
                if thumb and preview_refresh_token["value"] == load_token:
                    _thumb_cache[norm_path] = thumb
                    img_ref.bgcolor = None
                    img_ref.content = ft.Image(
                        src=thumb,
                        width=_ts, height=_ts,
                        fit=ft.BoxFit.COVER,
                        border_radius=ft.BorderRadius.all(4),
                    )

                    async def _apply():
                        try:
                            page.update()
                        except Exception:
                            pass

                    page.run_task(_apply)

        threading.Thread(target=_load, daemon=True).start()



    def request_quit():
        """Ferme la fenêtre principale de façon thread-safe via pubsub"""
        page.pubsub.send_all_on_topic("quit", None)



    def request_refresh():
        """Demande un rafraîchissement de la preview (thread-safe)"""
        page.pubsub.send_all_on_topic("refresh", None)



    def run_refresh_preview_command():
        """Exécute un rafraîchissement manuel de la preview et journalise la commande."""
        log_to_terminal("[CMD] refresh_preview(force_reload=True)", BLUE)
        refresh_preview(force_reload=True)



    def _suspend_keyboard_shortcuts() -> None:
        """Suspend temporairement les raccourcis globaux (ex: pendant une saisie en dialog)."""
        keyboard_shortcuts_suspended["count"] += 1



    def _resume_keyboard_shortcuts() -> None:
        """Réactive les raccourcis globaux suspendus."""
        keyboard_shortcuts_suspended["count"] = max(0, keyboard_shortcuts_suspended["count"] - 1)



    def on_keyboard_event(e: ft.KeyboardEvent):
        """Gestionnaire des événements clavier pour les raccourcis"""
        ctrl_pressed = e.ctrl or e.meta
        shift_pressed = bool(getattr(e, "shift", False))
        key_upper = (e.key or "").upper()

        # Si une boîte de dialogue de saisie est ouverte, laisser le champ texte
        # gérer ses propres raccourcis (copier/coller/supprimer, etc.).
        if keyboard_shortcuts_suspended["count"] > 0:
            return

        # Ctrl+↑ / Ctrl+↓ sont globaux : fonctionnent quelle que soit la zone active
        if ctrl_pressed and e.key in ("Arrow Up", "ArrowUp"):
            toggle_terminal_overlay()
            return

        if ctrl_pressed and e.key in ("Arrow Down", "ArrowDown"):
            if ai_mode["value"]:
                switch_to_terminal_mode()
            else:
                switch_to_ai_mode()
            return

        if ctrl_pressed and (ai_mode["value"] or note_mode["value"]):
            if shift_pressed and e.key in ("Arrow Left", "ArrowLeft"):
                toggle_ai_true_fullscreen()
                return
            if shift_pressed and e.key in ("Arrow Right", "ArrowRight"):
                toggle_notepad_true_fullscreen()
                return
            if e.key in ("Arrow Left", "ArrowLeft"):
                toggle_ai_fullscreen()
                return
            if e.key in ("Arrow Right", "ArrowRight"):
                toggle_notepad_fullscreen()
                return

        # Raccourcis globaux de gestion fichiers, même si le terminal a le focus.
        if ctrl_pressed and key_upper == "R":
            run_refresh_preview_command()
            return
        if ctrl_pressed and key_upper == "D":
            select_same_date(None)
            return

        if terminal_input_focused["value"]:
            # Laisser le champ terminal gérer ses raccourcis d'édition
            # (Ctrl/Cmd+C, Ctrl/Cmd+V, Backspace, Delete, etc.)
            if ctrl_pressed or e.key in ("Delete", "Backspace"):
                return
            if e.key in ("Arrow Up", "ArrowUp"):
                if command_history:
                    if history_index["value"] == -1:
                        history_draft["value"] = terminal_cmd_input.value or ""
                    history_index["value"] = min(history_index["value"] + 1, len(command_history) - 1)
                    terminal_cmd_input.value = command_history[history_index["value"]]
                    terminal_cmd_input.update()
                return
            elif e.key in ("Arrow Down", "ArrowDown"):
                if history_index["value"] >= 0:
                    history_index["value"] -= 1
                    terminal_cmd_input.value = (
                        history_draft["value"] if history_index["value"] == -1
                        else command_history[history_index["value"]]
                    )
                    terminal_cmd_input.update()
                return
            elif e.key == "Enter":
                on_terminal_command_submit(e)
                return

        if e.key == "Escape" and (note_mode["value"] or ai_mode["value"]):
            switch_to_terminal_mode()
            return

        # Ne pas intercepter les raccourcis clavier si le bloc-notes ou l'IA est actif
        # (laisser le TextField gérer ses propres Ctrl+A, Ctrl+C, etc.)
        if note_mode["value"] or ai_mode["value"]:
            return

        if ctrl_pressed:
            if key_upper == "A":
                toggle_select_all(None)
            elif key_upper == "C":
                copy_selected_files(None)
            elif key_upper == "I":
                invert_selection(None)
            elif key_upper == "N":
                create_new_folder(None)
            elif key_upper == "V":
                paste_files(None)
            elif key_upper == "X":
                cut_selected_files(None)
        elif e.key == "Delete" or (e.key == "Backspace" and bool(getattr(e, "meta", False))):
            delete_selected_files(None)

    # Activer la gestion des événements clavier
    page.on_keyboard_event = on_keyboard_event



    # ================================================================ #
    #                          TERMINAL                                #
    # ================================================================ #
    _terminal_update_timer = {"timer": None}
    _terminal_update_lock  = threading.Lock()

    def log_to_terminal(message, color=WHITE):
        """Ajoute un message au terminal intégré"""
        clean_message = _ANSI_ESCAPE.sub('', message).strip()
        if not clean_message:
            return
        try:
            terminal_output.controls.append(
                ft.Text(clean_message, size=CONSTANTS.TERMINAL_FONT_SIZE, color=color, font_family="monospace")
            )
            if len(terminal_output.controls) > 1000:
                terminal_output.controls.pop(0)
            with _terminal_update_lock:
                if _terminal_update_timer["timer"] is not None:
                    _terminal_update_timer["timer"].cancel()
                def _do_update():
                    page.update()
                    async def _scroll():
                        try:
                            await terminal_output.scroll_to(offset=-1)
                        except Exception:
                            pass
                    page.run_task(_scroll)
                t = threading.Timer(0.05, _do_update)
                _terminal_update_timer["timer"] = t
                t.start()
        except Exception:
            pass



    def clear_terminal(e):
        """Efface le contenu du terminal"""
        terminal_output.controls.clear()
        page.update()

        async def _refocus_after_clear():
            try:
                await terminal_cmd_input.focus()
            except Exception:
                pass

        page.run_task(_refocus_after_clear)



    def copy_terminal_to_clipboard():
        """Copie tout le contenu du terminal dans le presse-papiers"""
        if not terminal_output.controls:
            return
        terminal_text = "\n".join([ctrl.value for ctrl in terminal_output.controls if hasattr(ctrl, 'value')])
        try:
            if platform.system() == "Darwin":
                process = subprocess.Popen(['pbcopy'], stdin=subprocess.PIPE)
                process.communicate(terminal_text.encode('utf-8'))
            elif platform.system() == "Windows":
                process = subprocess.Popen(['clip'], stdin=subprocess.PIPE)
                process.communicate(terminal_text.encode('utf-16'))
            else:
                try:
                    process = subprocess.Popen(['xclip', '-selection', 'clipboard'], stdin=subprocess.PIPE)
                    process.communicate(terminal_text.encode('utf-8'))
                except FileNotFoundError:
                    process = subprocess.Popen(['xsel', '--clipboard', '--input'], stdin=subprocess.PIPE)
                    process.communicate(terminal_text.encode('utf-8'))
            log_to_terminal("[OK] Terminal copié dans le presse-papiers", GREEN)
        except Exception as e:
            log_to_terminal(f"[ERREUR] Copie presse-papiers: {e}", RED)



    # ── Bloc-notes : fonctions ──────────────────────────────────────────
    async def _notepad_save_as():
        """Exporte le contenu du bloc-notes dans un fichier choisi par l'utilisateur."""
        result = await ft.FilePicker().save_file(
            dialog_title="Enregistrer le bloc-notes sous…",
            file_name="notes.md",
            allowed_extensions=["txt", "md"],
        )
        if not result:
            return
        target_path = result if isinstance(result, str) else getattr(result, "path", None)
        if not target_path:
            return
        try:
            with open(target_path, "w", encoding="utf-8") as exported_file:
                exported_file.write(notepad_field.value or "")
            log_to_terminal(f"[OK] Bloc-notes exporté → {target_path}", GREEN)
        except Exception as export_error:
            log_to_terminal(f"[ERREUR] Export bloc-notes : {export_error}", RED)



    def _notepad_clear():
        """Efface tout le contenu du bloc-notes (sans sauvegarder)."""
        notepad_field.value = ""
        if notepad_is_preview["value"]:
            notepad_is_preview["value"] = False
            notepad_field.visible = True
            notepad_preview_scroll.visible = False
        try:
            page.update()
        except Exception:
            pass



    def _prepare_notepad_markdown(text: str) -> str:
        """Prépare le texte brut pour l'affichage Markdown en préservant les sauts
        de ligne : deux espaces en fin de ligne non vide (force <br>), et &nbsp;
        pour les lignes vides afin de conserver l'espacement vertical."""
        processed_lines = []
        for line in text.split("\n"):
            if line.strip() == "":
                processed_lines.append("&nbsp;")
            else:
                processed_lines.append(line + "  ")
        return "\n".join(processed_lines)



    def _notepad_toggle_preview():
        """Bascule entre édition et prévisualisation Markdown du bloc-notes."""
        notepad_is_preview["value"] = not notepad_is_preview["value"]
        is_preview = notepad_is_preview["value"]
        if is_preview:
            save_notes()
            notepad_markdown_preview.value = _prepare_notepad_markdown(notepad_field.value or "")
            notepad_preview_scroll.visible = True
            notepad_field.visible = False
        else:
            notepad_field.visible = True
            notepad_preview_scroll.visible = False
        try:
            page.update()
        except Exception:
            pass



    async def _notepad_autosave_after_delay():
        await asyncio.sleep(CONSTANTS.NOTEPAD_AUTOSAVE_DELAY)
        if note_mode["value"]:
            save_notes(restart_if_constants=False)

    def _notepad_on_change(e):
        t = _notepad_autosave_timer["task"]
        if t is not None and not t.done():
            t.cancel()
        _notepad_autosave_timer["task"] = page.run_task(
            _notepad_autosave_after_delay
        )

    notepad_field.on_change = _notepad_on_change

    def save_notes(restart_if_constants=True):
        """Sauvegarde le contenu du bloc-notes dans le fichier cible."""
        is_constants = (note_target_file["path"] == constants_file_path)
        try:
            with open(note_target_file["path"], "w", encoding="utf-8") as _f:
                _f.write(notepad_field.value or "")
            label = os.path.basename(note_target_file["path"])
            log_to_terminal(f"[OK] {label} sauvegardé", GREEN)
        except Exception as _err:
            log_to_terminal(f"[ERREUR] Sauvegarde : {_err}", RED)
            return
        if is_constants and restart_if_constants:
            log_to_terminal("[INFO] Redémarrage pour appliquer les nouvelles constantes…", ORANGE)
            dashboard_path = os.path.abspath(__file__)
            async def _restart_async():
                import time as _time
                _time.sleep(0.4)
                subprocess.Popen([sys.executable, dashboard_path])
                _time.sleep(0.2)
                try:
                    await page.window.close()
                except Exception:
                    pass
                os._exit(0)
            page.run_task(_restart_async)



    def load_notes():
        """Charge le contenu du bloc-notes depuis le fichier cible et adapte la coloration syntaxique."""
        # Toujours revenir en mode édition lors du chargement d'un nouveau fichier
        if notepad_is_preview["value"]:
            notepad_is_preview["value"] = False
            notepad_field.visible = True
            notepad_preview_scroll.visible = False
            
        # Détection dynamique de l'extension pour appliquer la bonne coloration
        path = note_target_file.get("path", "")
        ext = os.path.splitext(path)[1].lower() if path else ""
        
        if _HAS_CODE_EDITOR:
            if ext in [".py", ".pyw"]:
                notepad_field.language = fce.CodeLanguage.PYTHON
            elif ext == ".json":
                notepad_field.language = fce.CodeLanguage.JSON
            elif ext in [".md", ".markdown"]:
                notepad_field.language = fce.CodeLanguage.MARKDOWN
            else:
                notepad_field.language = fce.CodeLanguage.PLAINTEXT

        try:
            if os.path.exists(note_target_file["path"]):
                with open(note_target_file["path"], "r", encoding="utf-8") as _f:
                    content = _f.read()
            else:
                content = ""
        except Exception:
            content = ""
        notepad_field.value = content
        if notepad_is_preview["value"]:
            notepad_markdown_preview.value = content
            notepad_preview_scroll.visible = True
            notepad_field.visible = False


    def _open_notepad_ui(title, icon, color, hint):
        """Affiche la zone bloc-notes (+ IA) avec le titre et la couleur donnés."""
        notepad_header_icon.name  = icon
        notepad_header_icon.color = color
        notepad_header_title.value = title
        notepad_header_title.color = color
        notepad_field.hint_text   = hint
        load_notes()
        note_mode["value"] = True
        ai_mode["value"]   = True
        terminal_output.visible  = False
        terminal_cmd_row.visible = False
        update_overlay_visibility()
        try:
            page.update()
        except Exception:
            pass

        async def _focus_note():
            try:
                await notepad_field.focus()
            except Exception:
                pass
        page.run_task(_focus_note)



    def switch_to_note():
        """Bascule la zone bas en mode bloc-notes (fichier .notes.md)."""
        if note_mode["value"]:
            save_notes()
        note_target_file["path"] = notes_file_path
        _open_notepad_ui("Notes", ft.Icons.EDIT_NOTE, VIOLET, "Écrivez vos notes ici…")



    def switch_to_options():
        """Bascule la zone bas en mode édition CONSTANTS.py."""
        if note_mode["value"]:
            save_notes()
        note_target_file["path"] = constants_file_path
        _open_notepad_ui("CONSTANTS.py", ft.Icons.TUNE, ORANGE, "Modifiez les constantes ici…")
        return



    def open_file_in_notepad(file_path):
        """Ouvre un fichier texte dans le bloc-notes intégré et affiche le panneau."""
        if note_mode["value"]:
            save_notes()
        note_target_file["path"] = file_path
        title = os.path.basename(file_path)
        _open_notepad_ui(title, ft.Icons.DESCRIPTION, VIOLET, "")



    def _create_and_open_info_txt(e=None):
        """Crée INFO.txt dans le dossier courant (si inexistant) et l'ouvre dans le bloc-notes."""
        folder = current_browse_folder["path"] or selected_folder.get("path")
        if not folder or not os.path.isdir(folder):
            log_to_terminal("[INFO] Aucun dossier sélectionné", LIGHT_GREY)
            return
        file_path = os.path.join(folder, "INFO.txt")
        if not os.path.exists(file_path):
            try:
                with open(file_path, "w", encoding="utf-8") as _f:
                    pass
            except Exception as err:
                log_to_terminal(f"[ERREUR] Création INFO.txt : {err}", RED)
                return
            refresh_preview(reset_page=False)
        open_file_in_notepad(file_path)



    # ── Intelligence artificielle ──────────────────────────────────────
    def switch_to_ai_mode():
        """Bascule la zone bas en mode conversation IA + notes."""
        note_target_file["path"] = notes_file_path
        load_notes()
        note_mode["value"] = True
        ai_mode["value"] = True
        terminal_output.visible  = False
        terminal_cmd_row.visible = False
        update_overlay_visibility()
        terminal_output.update()
        terminal_cmd_row.update()
        try:
            page.update()
        except Exception:
            pass

        # Pré-démarrer Ollama en silence pendant que l'utilisateur tape
        def _silent_prestart():
            try:
                urllib.request.urlopen(f"{CONSTANTS.AI_OLLAMA_URL}/api/tags", timeout=3).close()
            except Exception:
                try:
                    ollama_process["proc"] = subprocess.Popen(
                        ["ollama", "serve"],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                    )
                except Exception:
                    pass
        threading.Thread(target=_silent_prestart, daemon=True).start()

        async def _focus_ai():
            try:
                await ai_input_field.focus()
            except Exception:
                pass
        page.run_task(_focus_ai)



    def _clear_ai_conversation():
        """Efface l'historique de la conversation IA et supprime le fichier .ai_conversation.json."""
        ai_conversation.clear()
        ai_chat_view.controls.clear()
        ai_status_text.value = ""
        try:
            if os.path.isfile(ai_history_file_path):
                os.remove(ai_history_file_path)
        except Exception:
            pass
        try:
            page.update()
        except Exception:
            pass



    def _ai_save_history():
        """Sauvegarde ai_conversation dans .ai_conversation.json."""
        try:
            # Ne sauvegarder que role + content (pas les images base64)
            serializable = [
                {"role": message["role"], "content": message["content"]}
                for message in ai_conversation
                if message.get("role") in ("user", "assistant")
            ]
            with open(ai_history_file_path, "w", encoding="utf-8") as history_file:
                json.dump(serializable, history_file, ensure_ascii=False, indent=2)
        except Exception:
            pass



    def _ai_load_history():
        """Charge .ai_conversation.json dans ai_conversation et reconstruit ai_chat_view."""
        if not os.path.isfile(ai_history_file_path):
            return
        try:
            with open(ai_history_file_path, "r", encoding="utf-8") as history_file:
                saved_messages = json.load(history_file)
            for message in saved_messages:
                role = message.get("role")
                content = message.get("content", "")
                if role not in ("user", "assistant"):
                    continue
                ai_conversation.append({"role": role, "content": content})
                is_user = role == "user"
                if is_user:
                    bubble_text = ft.Text(
                        content,
                        size=CONSTANTS.TERMINAL_FONT_SIZE,
                        color=BLUE,
                        font_family="monospace",
                        selectable=True,
                        no_wrap=False,
                    )
                else:
                    bubble_text = ft.Markdown(
                        content,
                        selectable=True,
                        extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,
                        code_theme=ft.MarkdownCodeTheme.ATOM_ONE_DARK,
                        expand=True,
                    )
                bubble = ft.Container(
                    content=bubble_text,
                    bgcolor=DARK if is_user else GREY,
                    border_radius=6,
                    padding=ft.Padding(8, 4, 8, 4),
                    expand=True,
                )
                if is_user:
                    ai_chat_view.controls.append(
                        ft.Row(
                            [bubble],
                            alignment=ft.MainAxisAlignment.END,
                        )
                    )
                else:
                    raw_text = content

                    def _speak_loaded_bubble_text(text_control, fallback_text):
                        current_text = getattr(text_control, "value", "")
                        if not isinstance(current_text, str) or not current_text.strip():
                            current_text = fallback_text
                        threading.Thread(
                            target=_speak_bubble,
                            args=(current_text,),
                            kwargs={"force_chunked": True},
                            daemon=True,
                        ).start()

                    speak_btn = ft.IconButton(
                        icon=ft.Icons.VOLUME_UP,
                        icon_color=LIGHT_GREY,
                        icon_size=14,
                        tooltip="Lire cette réponse (lecture fidèle)",
                        on_click=lambda e, text_control=bubble_text, fallback_text=raw_text: _speak_loaded_bubble_text(
                            text_control,
                            fallback_text,
                        ),
                    )
                    ai_chat_view.controls.append(
                        ft.Row(
                            [bubble, speak_btn],
                            alignment=ft.MainAxisAlignment.START,
                            vertical_alignment=ft.CrossAxisAlignment.START,
                        )
                    )
        except Exception:
            pass



    def _export_ai_conversation(to_notepad=False, event=None):
        """Copie la conversation IA dans le presse-papiers, et la transfère dans le bloc-notes si demandé."""
        if not ai_conversation:
            log_to_terminal("[IA] Aucune conversation à exporter", LIGHT_GREY)
            return
        text = _format_ai_conversation(ai_conversation, CONSTANTS.AI_USER_NAME, CONSTANTS.AI_SEPARATOR_WIDTH)
        async def _copy():
            try:
                await ft.Clipboard().set(text)
                log_to_terminal("[IA] Conversation copiée dans le presse-papiers", BLUE)
            except Exception as copy_error:
                log_to_terminal(f"[ERREUR] Copie IA : {copy_error}", RED)
        page.run_task(_copy)
        if to_notepad:
            switch_to_note()
            existing = notepad_field.value or ""
            sep = "\n\n" + "#" * CONSTANTS.AI_SEPARATOR_WIDTH + "\n\n" if existing.strip() else ""
            notepad_field.value = existing + sep + text
            try:
                notepad_field.update()
            except Exception:
                pass



    def _ai_stop_model():
        """Libère le modèle chargé en RAM via `ollama stop`."""
        def _run_stop():
            try:
                current_model = ai_model_dropdown.value or CONSTANTS.AI_MODEL_TEXT
                if not (current_model or "").startswith(("gemini", "claude")):
                    subprocess.run(["ollama", "stop", CONSTANTS.AI_MODEL_VISION], timeout=10)
                    subprocess.run(["ollama", "stop", CONSTANTS.AI_MODEL_TEXT],   timeout=10)
            except Exception:
                pass
            ai_stop_button.visible = False
            ai_stop_button.icon_color = LIGHT_GREY
            ai_status_text.value = ""
            try:
                page.update()
            except Exception:
                pass
        threading.Thread(target=_run_stop, daemon=True).start()



    # ── Gestion des images jointes ────────────────────────────────────
    def _ai_refresh_attach_row():
        """Reconstruit la barre de pièces jointes visuellement."""
        ai_attach_row.controls.clear()
        for image_entry in ai_pending_images:
            name = os.path.basename(image_entry["path"])
            image_mode = image_entry.get("mode", "optimized")
            mode_text = "R" if image_mode == "original" else "1024"
            mode_color = BLUE if image_mode == "original" else LIGHT_GREY
            # Copie locale pour la lambda
            entry_ref = image_entry
            ai_attach_row.controls.append(
                ft.Container(
                    content=ft.Row([
                        ft.Icon(ft.Icons.IMAGE, size=13, color=ORANGE),
                        ft.Container(
                            content=ft.Text(mode_text, size=9, color=mode_color, weight=ft.FontWeight.BOLD),
                            bgcolor=DARK,
                            border=ft.Border.all(1, GREY),
                            border_radius=3,
                            padding=ft.Padding(3, 0, 3, 0),
                        ),
                        ft.Text(name, size=11, color=ORANGE),
                        ft.IconButton(
                            icon=ft.Icons.CLOSE,
                            icon_color=RED,
                            icon_size=12,
                            tooltip="Retirer",
                            style=ft.ButtonStyle(padding=ft.Padding.all(2)),
                            on_click=lambda e, ref=entry_ref: _ai_remove_image(ref),
                        ),
                    ], spacing=2, tight=True),
                    bgcolor=GREY,
                    border_radius=4,
                    padding=ft.Padding(4, 2, 4, 2),
                )
            )
        for file_path in ai_pending_files:
            name = os.path.basename(file_path)
            icon_name = ft.Icons.DESCRIPTION
            entry_ref = file_path
            ai_attach_row.controls.append(
                ft.Container(
                    content=ft.Row([
                        ft.Icon(icon_name, size=13, color=YELLOW),
                        ft.Text(name, size=11, color=YELLOW),
                        ft.IconButton(
                            icon=ft.Icons.CLOSE,
                            icon_color=RED,
                            icon_size=12,
                            tooltip="Retirer",
                            style=ft.ButtonStyle(padding=ft.Padding.all(2)),
                            on_click=lambda e, ref=entry_ref: _ai_remove_file(ref),
                        ),
                    ], spacing=2, tight=True),
                    bgcolor=GREY,
                    border_radius=4,
                    padding=ft.Padding(4, 2, 4, 2),
                )
            )
        ai_attach_row.visible = bool(ai_pending_images) or bool(ai_pending_files)
        try:
            page.update()
        except Exception:
            pass



    def _ai_attach_image(image_path, use_original=None):
        """Encode une image en base64 (optimisé ou taille réelle) et l'ajoute aux pièces jointes."""
        # Vérifier si déjà jointe
        if any(entry["path"] == image_path for entry in ai_pending_images):
            return
        if use_original is None:
            use_original = ai_send_original_images["value"]
        try:
            if use_original:
                with open(image_path, "rb") as image_file:
                    b64_data = base64.b64encode(image_file.read()).decode("utf-8")
            else:
                from PIL import Image as PilImage
                import io as _io
                with PilImage.open(image_path) as pil_img:
                    pil_img = pil_img.convert("RGB")
                    max_side = 1024
                    width, height = pil_img.size
                    if width > max_side or height > max_side:
                        ratio = min(max_side / width, max_side / height)
                        new_size = (int(width * ratio), int(height * ratio))
                        pil_img = pil_img.resize(new_size, PilImage.LANCZOS)
                    buffer = _io.BytesIO()
                    pil_img.save(buffer, format="JPEG", quality=85)
                    b64_data = base64.b64encode(buffer.getvalue()).decode("utf-8")
        except Exception:
            # Fallback : lecture brute si Pillow échoue
            try:
                with open(image_path, "rb") as image_file:
                    b64_data = base64.b64encode(image_file.read()).decode("utf-8")
            except Exception as exc:
                _ai_add_bubble("assistant", f"[ERREUR] Impossible de lire l'image : {exc}")
                return
        image_mode = "original" if use_original else "optimized"
        ai_pending_images.append({"path": image_path, "b64": b64_data, "mode": image_mode})
        _ai_refresh_attach_row()
        # Avertir si la configuration vision n'est pas reconnue
        vision_model = CONSTANTS.AI_MODEL_VISION
        is_vision = any(
            vision_model == entry[1] or vision_model.startswith(entry[1] + ":")
            for entry in CONSTANTS.AI_AVAILABLE_MODELS
            if entry[2]
        )
        if not is_vision:
            _ai_add_bubble(
                "assistant",
                "⚠️ La configuration vision actuelle n'est pas reconnue comme compatible.\n"
                "Vérifiez AI_MODEL_VISION dans CONSTANTS.py.",
            )



    def _ai_remove_image(image_entry):
        """Retire une image des pièces jointes en attente."""
        if image_entry in ai_pending_images:
            ai_pending_images.remove(image_entry)
        _ai_refresh_attach_row()



    # ── Extensions reconnues comme documents ───────────────────────────
    _AI_DOCUMENT_EXTS = CONSTANTS.AI_DOCUMENT_EXTS

    def _ai_attach_document_file(file_path):
        """Ajoute un document aux pièces jointes en attente."""
        if file_path in ai_pending_files:
            return
        ai_pending_files.append(file_path)
        _ai_refresh_attach_row()



    def _ai_remove_file(file_entry):
        """Retire un document des pièces jointes en attente."""
        if file_entry in ai_pending_files:
            ai_pending_files.remove(file_entry)
        _ai_refresh_attach_row()



    def _ai_extract_file_content(file_path):
        """
        Extrait le contenu textuel d'un document.
        Retourne (nom_affiché, texte_extrait).
        Lève une exception si l'extraction échoue.
        """
        ext = os.path.splitext(file_path)[1].lower()
        name = os.path.basename(file_path)

        # ── Documents ────────────────────────────────────────────────
        if ext == ".pdf":
            try:
                import fitz as _fitz
                pdf_doc = _fitz.open(file_path)
                text = "\n".join(pdf_page.get_text("text") for pdf_page in pdf_doc)
                pdf_doc.close()
                return name, text
            except ImportError:
                raise ImportError("PyMuPDF non disponible pour lire les PDF.")

        if ext in (".docx", ".doc"):
            try:
                import docx as _docx
                word_doc = _docx.Document(file_path)
                text = "\n".join(para.text for para in word_doc.paragraphs)
                return name, text
            except ImportError:
                raise ImportError(
                    "python-docx non installé. Installez-le avec : pip install python-docx"
                )

        # Fichier texte brut (tous les autres formats)
        with open(file_path, "r", encoding="utf-8", errors="replace") as text_file:
            return name, text_file.read()



    async def _ai_pick_any():
        """Ouvre un sélecteur de fichier pour joindre une image ou un document."""
        _image_exts_pick = {"jpg", "jpeg", "png", "gif", "bmp", "webp"}
        result = await ft.FilePicker().pick_files(
            dialog_title="Joindre une image ou un document",
            allowed_extensions=[
                "jpg", "jpeg", "png", "gif", "bmp", "webp",
                "txt", "md", "py", "js", "ts", "json", "csv", "xml",
                "html", "htm", "yaml", "yml", "toml", "ini", "cfg", "log",
                "rst", "pdf", "docx", "doc", "rtf",
            ],
            allow_multiple=True,
        )
        if result:
            for picked_file in result:
                if picked_file.path:
                    ext = os.path.splitext(picked_file.path)[1].lstrip(".").lower()
                    if ext in _image_exts_pick:
                        _ai_attach_image(picked_file.path)
                    else:
                        _ai_attach_document_file(picked_file.path)



    def ai_send_selected_images(e=None):
        """Joint les fichiers image et document sélectionnés dans la preview à la conversation IA."""
        image_exts = CONSTANTS.IMAGE_EXTS
        image_paths = [
            file_path for file_path in selected_files
            if os.path.splitext(file_path)[1].lower() in image_exts
        ]
        file_paths = [
            file_path for file_path in selected_files
            if os.path.splitext(file_path)[1].lower() in _AI_DOCUMENT_EXTS
        ]
        if not image_paths and not file_paths:
            log_to_terminal("[IA] Aucun fichier compatible sélectionné dans la preview", LIGHT_GREY)
            return
        switch_to_ai_mode()
        for image_path in image_paths:
            _ai_attach_image(image_path, use_original=CONSTANTS.AI_IMAGE_ATTACH_SELECTED_ORIGINAL)
        for file_path in file_paths:
            _ai_attach_document_file(file_path)
        total = len(image_paths) + len(file_paths)
        log_to_terminal(f"[IA] {total} fichier(s) joint(s)", BLUE)



    def _ensure_ollama_ready(model_name=None):
        """
        Vérifie qu'Ollama est lancé et que le modèle est disponible.
        Lance le serveur et/ou télécharge le modèle si nécessaire.
        Retourne True si tout est prêt, False en cas d'erreur bloquante.
        Doit être appelé depuis un thread secondaire (bloquant).
        """        
        if model_name is None:
            model_name = CONSTANTS.AI_MODEL_TEXT
        # Les modèles Gemini et Claude n'ont pas besoin d'Ollama
        if (model_name or "").startswith(("gemini", "claude")):
            return True
        def _is_ollama_up():
            try:
                with urllib.request.urlopen(
                    f"{CONSTANTS.AI_OLLAMA_URL}/api/tags", timeout=3
                ) as resp:
                    return resp.status == 200
            except Exception:
                return False

        # ── 1. Démarrer Ollama si nécessaire ──────────────────────────
        if not _is_ollama_up():
            _ai_add_bubble("assistant", "⚙️ Démarrage d'Ollama en arrière-plan…")
            try:
                ollama_process["proc"] = subprocess.Popen(
                    ["ollama", "serve"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
            except FileNotFoundError:
                _ai_add_bubble(
                    "assistant",
                    "[ERREUR] Ollama n'est pas installé sur cette machine.\n"
                    "Téléchargez-le sur https://ollama.com",
                )
                return False
            # Attendre jusqu'à 20 s que le serveur réponde
            for _ in range(40):
                time.sleep(0.5)
                if _is_ollama_up():
                    break
            else:
                _ai_add_bubble(
                    "assistant",
                    "[ERREUR] Ollama n'a pas démarré dans les délais impartis.",
                )
                return False

        # ── 2. Vérifier si le modèle est présent ──────────────────────
        try:
            with urllib.request.urlopen(
                f"{CONSTANTS.AI_OLLAMA_URL}/api/tags", timeout=5
            ) as resp:
                available_names = [
                    model.get("name", "")
                    for model in json.loads(resp.read().decode("utf-8")).get("models", [])
                ]
            model_present = any(
                name == model_name or name.startswith(model_name + ":")
                for name in available_names
            )
        except Exception:
            model_present = False

        if not model_present:
            pull_status_ctrl = _ai_add_bubble(
                "assistant",
                "⬇️ Téléchargement du composant IA…\n"
                "(première utilisation — peut prendre quelques minutes)",
            )
            try:
                pull_payload = json.dumps(
                    {"name": model_name, "stream": True}
                ).encode("utf-8")
                pull_request = urllib.request.Request(
                    f"{CONSTANTS.AI_OLLAMA_URL}/api/pull",
                    data=pull_payload,
                    headers={"Content-Type": "application/json"},
                    method="POST",
                )
                with urllib.request.urlopen(pull_request, timeout=3600) as pull_resp:
                    for raw_line in pull_resp:
                        chunk = json.loads(raw_line.decode("utf-8"))
                        status = chunk.get("status", "")
                        completed = chunk.get("completed", 0)
                        total = chunk.get("total", 0)
                        if total:
                            pct = int(completed / total * 100)
                            pull_status_ctrl.value = (
                                f"⬇️ Téléchargement — {status} {pct}%"
                            )
                        elif status:
                            pull_status_ctrl.value = (
                                f"⬇️ Téléchargement — {status}"
                            )
                        try:
                            page.update()
                        except Exception:
                            pass
                pull_status_ctrl.value = "✅ Composant IA téléchargé et prêt !"
                try:
                    page.update()
                except Exception:
                    pass
            except Exception as exc:
                _ai_add_bubble("assistant", f"[ERREUR] Téléchargement du modèle : {exc}")
                return False

        return True



    def _clean_file_content(raw_content):
        """
        Retire les artefacts de raisonnement inline de Gemma (chain-of-thought)
        qui s'immiscent parfois dans les arguments de create_file.
        Stratégie :
          - Retire les tokens spéciaux Gemma (<channel|>, <|tool_call>…)
          - Dès qu'une ligne de "thinking" est détectée (Wait,/Hmm,/I see…),
            tronque tout le reste (le contenu qui suit est invalide).
          - Retire les lignes de méta-commentaire isolées.
        """
        import re as _re_cfc
        # Tokens spéciaux Gemma et balises tool_call résiduelles
        content = _re_cfc.sub(r'<channel\|>', '', raw_content)
        content = _re_cfc.sub(r'<\|[^|>]+\|>', '', content)
        content = _re_cfc.sub(r'<\|tool_call>.*?(?:<tool_call\|>|$)', '', content, flags=_re_cfc.DOTALL)

        lines = content.split('\n')
        clean_lines = []

        # Patterns de TRONCATURE : dès qu'une de ces lignes apparaît, tout ce qui suit
        # est du raisonnement Gemma — on coupe ici.
        _TRUNCATE_RE = _re_cfc.compile(
            r'^\s*(?:Wait[,\s—]|Hmm[,\s.]|Actually[,\s—]|I see a discrepancy|'
            r'I notice that|Let me reconsider|Let me re-|I need to re-|'
            r'OK so[,\s]|OK, so[,\s]|I will re-run|I should re-)',
            _re_cfc.IGNORECASE,
        )
        # Patterns de lignes individuelles à ignorer (sans tronquer le reste)
        _SKIP_RE = _re_cfc.compile(
            r'^\s*(?:'
            r'\((?:Note|Wait|Self-correction|Correction|Assuming|Final attempt|'
            r'I will|Since the|The prompt|This was|Using the|Given that|'
            r'Final output|OK,? I|Let me re)'
            r'|(?:Actual list from|Final list based|Listing all files and|'
            r'File List:|Assuming the file|I will generate|I will use the|'
            r'I will stop|I will provide|I will list|I will present|'
            r'Since the prompt|The file list has been|I will assume)'
            r')',
            _re_cfc.IGNORECASE,
        )
        for line in lines:
            # Troncature : début du raisonnement Gemma → arrêt immédiat
            if _TRUNCATE_RE.match(line):
                break
            if _SKIP_RE.match(line):
                continue
            # Retire les parenthèses de raisonnement en fin de ligne
            # ex: "fichier.jpg  (Note: Correction: ...)" → "fichier.jpg"
            line = _re_cfc.sub(
                r'\s*\((?:Note|Wait|Self-correction|Correction):.*',
                '', line, flags=_re_cfc.IGNORECASE,
            )
            clean_lines.append(line)
        # Supprime les blocs de 3+ lignes vides consécutives
        result = _re_cfc.sub(r'\n{3,}', '\n\n', '\n'.join(clean_lines))
        return result.strip()



    def _md_dark(text: str) -> str:
        """Remplace les blockquotes Markdown (fond bleu clair de Flutter)
        par un équivalent lisible sur thème sombre."""
        lines = text.split("\n")
        result = []
        for line in lines:
            if line.startswith("> "):
                result.append("**›** " + line[2:])
            elif line == ">":
                result.append("")
            else:
                result.append(line)
        return "\n".join(result)



    def _speak_bubble(text, force_chunked=False):
        """Lit un texte via Gemini TTS.

        force_chunked=True force la lecture fidèle du texte affiché (sans mode Live).
        """
        def _set_tts_feedback(status_text, show_progress):
            ai_status_text.value = status_text
            ai_progress_bar.visible = show_progress

            async def _apply_ui_update():
                try:
                    page.update()
                except Exception:
                    try:
                        ai_status_text.update()
                        ai_progress_bar.update()
                    except Exception:
                        pass

            try:
                page.run_task(_apply_ui_update)
            except Exception:
                try:
                    page.update()
                except Exception:
                    pass

        # Arrêter le TTS précédent s'il tourne encore
        if ai_tts_stop_event["event"] is not None:
            ai_tts_stop_event["event"].set()
        stop_event = threading.Event()
        ai_tts_stop_event["event"] = stop_event
        selected_tts_mode = "chunked" if force_chunked else CONSTANTS.AI_VOICE_TTS_MODE
        if selected_tts_mode == "live":
            _set_tts_feedback(f"🔊 Live — {CONSTANTS.AI_VOICE_TTS_VOICE}…", True)
        else:
            _set_tts_feedback(f"🔊 Préparation de la voix — {CONSTANTS.AI_VOICE_TTS_VOICE}…", True)
        try:
            if selected_tts_mode == "live":
                _gemini_live_tts_stream(
                    text,
                    model=CONSTANTS.AI_VOICE_LIVE_MODEL,
                    voice_name=CONSTANTS.AI_VOICE_TTS_VOICE,
                    sample_rate=CONSTANTS.AI_VOICE_TTS_SAMPLE_RATE,
                    language_code=CONSTANTS.AI_VOICE_TTS_LANGUAGE,
                    stop_event=stop_event,
                )
            else:
                # One-shot par défaut pour un timbre stable sur toute la réponse.
                pcm = _gemini_tts(
                    text,
                    voice_name=CONSTANTS.AI_VOICE_TTS_VOICE,
                    tts_model=CONSTANTS.AI_VOICE_TTS_MODEL,
                    language_code=CONSTANTS.AI_VOICE_TTS_LANGUAGE,
                )
                if pcm:
                    _voice_play_audio(
                        pcm,
                        sample_rate=CONSTANTS.AI_VOICE_TTS_SAMPLE_RATE,
                        stop_event=stop_event,
                    )
                else:
                    # Fallback de sécurité si la requête unique échoue.
                    _gemini_tts_stream(
                        text,
                        voice_name=CONSTANTS.AI_VOICE_TTS_VOICE,
                        tts_model=CONSTANTS.AI_VOICE_TTS_MODEL,
                        sample_rate=CONSTANTS.AI_VOICE_TTS_SAMPLE_RATE,
                        language_code=CONSTANTS.AI_VOICE_TTS_LANGUAGE,
                        stop_event=stop_event,
                    )
        except Exception as tts_exc:
            _set_tts_feedback(f"[❌ TTS] {tts_exc}", False)
            return
        finally:
            is_current_tts = ai_tts_stop_event["event"] is stop_event
            if is_current_tts:
                ai_tts_stop_event["event"] = None
                _set_tts_feedback("", False)



    def _ai_add_bubble(role, text):
        """Ajoute un message dans le panneau IA et retourne le contrôle (pour le streaming)."""
        is_user  = role == "user"
        is_think = role == "think"
        if is_user:
            bubble_text = ft.Text(
                text,
                size=CONSTANTS.TERMINAL_FONT_SIZE,
                color=BLUE,
                font_family="monospace",
                selectable=True,
                no_wrap=False,
            )
        elif is_think:
            bubble_text = ft.Text(
                f"💭 {text}",
                size=CONSTANTS.TERMINAL_FONT_SIZE - 1,
                color=LIGHT_GREY,
                italic=True,
                selectable=True,
                no_wrap=False,
            )
        else:
            bubble_text = ft.Markdown(
                _md_dark(text),
                selectable=True,
                extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,
                code_theme=ft.MarkdownCodeTheme.ATOM_ONE_DARK,
                md_style_sheet=ft.MarkdownStyleSheet(
                    p_text_style=ft.TextStyle(size=CONSTANTS.TERMINAL_FONT_SIZE),
                ),
                expand=True,
            )
        bubble = ft.Container(
            content=bubble_text,
            bgcolor="#1a1c20" if is_think else (DARK if is_user else GREY),
            border=ft.Border.all(1, LIGHT_GREY) if is_think else None,
            border_radius=6,
            padding=ft.Padding(8, 4, 8, 4),
            expand=True,
        )
        if not is_user and not is_think:
            raw_text = text

            def _speak_current_bubble_text(text_control, fallback_text):
                current_text = getattr(text_control, "value", "")
                if not isinstance(current_text, str) or not current_text.strip():
                    current_text = fallback_text
                _speak_bubble(current_text, force_chunked=True)

            speak_btn = ft.IconButton(
                icon=ft.Icons.VOLUME_UP,
                icon_color=LIGHT_GREY,
                icon_size=14,
                tooltip="Lire cette réponse (lecture fidèle)",
                on_click=lambda e, text_control=bubble_text, fallback_text=raw_text: threading.Thread(
                    target=_speak_current_bubble_text,
                    args=(text_control, fallback_text),
                    daemon=True,
                ).start(),
            )
            row = ft.Row(
                [bubble, speak_btn],
                alignment=ft.MainAxisAlignment.START,
                vertical_alignment=ft.CrossAxisAlignment.START,
            )
        else:
            row = ft.Row(
                [bubble],
                alignment=ft.MainAxisAlignment.END if is_user else ft.MainAxisAlignment.START,
            )
        ai_chat_view.controls.append(row)
        async def _update_and_scroll():
            try:
                page.update()
                await asyncio.sleep(0)
                await ai_chat_view.scroll_to(offset=-1)
            except Exception:
                pass
        page.run_task(_update_and_scroll)
        return bubble_text



    def _ai_add_image_bubble(image_path):
        """Affiche une image générée dans le chat IA."""
        image_src = image_path
        try:
            if os.path.isfile(image_path):
                cached_image = thumb_cache.get_or_generate(image_path)
                if cached_image:
                    image_src = cached_image
                else:
                    with open(image_path, "rb") as image_file:
                        image_src = image_file.read()
        except Exception:
            image_src = image_path

        img_widget = ft.Image(
            src=image_src,
            width=400,
            border_radius=8,
            fit=ft.BoxFit.CONTAIN,
        )
        row = ft.Row(
            [ft.Container(img_widget, border_radius=8)],
            alignment=ft.MainAxisAlignment.START,
        )
        ai_chat_view.controls.append(row)
        async def _upd():
            try:
                page.update()
                await asyncio.sleep(0)
                await ai_chat_view.scroll_to(offset=-1)
            except Exception:
                pass
        page.run_task(_upd)



    def _send_ai_message(message_text):
        """Envoie un message à Ollama et streame la réponse dans le panneau IA."""
        if ai_streaming["value"]:
            return
        if not message_text.strip() and not ai_pending_images and not ai_pending_files:
            return
        ai_streaming["value"] = True
        ai_stop_button.visible = True
        ai_stop_button.icon_color = RED
        ai_status_text.value = "⏳ En cours…"
        try:
            page.update()
        except Exception:
            pass

        # Capturer et vider les images jointes avant le thread
        images_b64   = [entry["b64"]  for entry in ai_pending_images]
        images_paths = [entry["path"] for entry in ai_pending_images]
        ai_pending_images.clear()
        _ai_refresh_attach_row()

        # Capturer et vider les documents joints avant le thread
        files_to_inject = list(ai_pending_files)
        ai_pending_files.clear()
        _ai_refresh_attach_row()

        # Choisir le modèle sélectionné par l'utilisateur
        active_model = ai_model_dropdown.value or CONSTANTS.AI_MODEL_TEXT

        enriched_text = message_text
        # Détecter les URLs dans le message et injecter leur contenu
        url_pattern = re.compile(r'https?://[^\s<>"\)\]]+', re.IGNORECASE)
        found_urls = url_pattern.findall(message_text)
        if found_urls:
            url_blocks = []
            for url in found_urls:
                page_content = _fetch_url_content(url, max_chars=CONSTANTS.AI_URL_MAX_CHARS)
                url_blocks.append(f"--- Contenu de {url} ---\n{page_content}\n--- Fin ---")
            enriched_text = enriched_text + "\n\n" + "\n\n".join(url_blocks)

        # Injecter le nom des fichiers image joints pour que l'IA les identifie
        if images_paths:
            filenames_info = ", ".join(os.path.basename(path) for path in images_paths)
            enriched_text = enriched_text + f"\n[Image(s) jointe(s) : {filenames_info}]"

        # Construire l'entrée utilisateur (avec images si présent)
        user_message = {"role": "user", "content": enriched_text}
        if images_b64:
            user_message["images"] = images_b64
        ai_conversation.append(user_message)

        # Afficher la bulle utilisateur avec indicateur image/fichier
        display_text = message_text
        if images_b64:
            display_text = f"🖼️ ({len(images_b64)} image(s))  {message_text}" if message_text else f"🖼️ {len(images_b64)} image(s) jointe(s)"
        if files_to_inject:
            files_label = "  ".join(
                "📄 " + os.path.basename(file_path)
                for file_path in files_to_inject
            )
            display_text = (display_text + "  " if display_text else "") + files_label
        _ai_add_bubble("user", display_text)

        def _run():
            full_response = ""
            response_text_ctrl = None
            try:
                # S'assurer qu'Ollama est prêt (serveur + modèle)
                if not _ensure_ollama_ready(active_model):
                    return

                # Indiquer que le modèle est en cours de chargement
                loading_ctrl = _ai_add_bubble("assistant", "⏳ Réflexion en cours…")

                def _remove_loading():
                    nonlocal loading_ctrl
                    if loading_ctrl is not None:
                        try:
                            ai_chat_view.controls = [
                                row for row in ai_chat_view.controls
                                if not (
                                    hasattr(row, "controls") and row.controls
                                    and hasattr(row.controls[0], "content")
                                    and row.controls[0].content is loading_ctrl
                                )
                            ]
                            loading_ctrl = None
                        except Exception:
                            pass

                # Extraire et injecter le contenu des documents joints
                if files_to_inject:
                    injected_blocks = []
                    for file_path in files_to_inject:
                        file_name = os.path.basename(file_path)
                        try:
                            label, content = _ai_extract_file_content(file_path)
                            injected_blocks.append(
                                f"--- Document : {label} ---\n{content[:50000]}\n--- Fin ---"
                            )
                        except Exception as extraction_exc:
                            _ai_add_bubble("assistant", f"[ERREUR] {file_name} : {extraction_exc}")
                    if injected_blocks:
                        ai_conversation[-1]["content"] += "\n\n" + "\n\n".join(injected_blocks)
                    ai_status_text.value = "⏳ En cours…"
                    try:
                        page.update()
                    except Exception:
                        pass

                # ── Outils dossier (disponibles si un dossier est ouvert) ─────
                _folder_path_for_tools = current_browse_folder["path"] or selected_folder["path"]
                _FOLDER_TOOLS = _folder_tool_definitions(_folder_path_for_tools)
                _NEW_TOOLS = (_EDIT_TOOLS + _READ_LINES_TOOLS + _SEARCH_TOOLS + _GIT_TOOLS + _TASK_TOOLS
                              + _PDF_TOOLS + _SUBAGENT_TOOLS + _SCHEDULE_TOOLS
                              + _HTTP_TOOLS + _SPREADSHEET_TOOLS + _PYAUTOGUI_TOOLS)
                if (active_model or "").startswith("gemini"):
                    _ALL_TOOLS = _WEB_TOOLS + _TERMINAL_TOOLS + _MEMORY_TOOLS + _SCREENSHOT_TOOLS + _NOTEPAD_TOOLS + _UI_TOOLS + _NEW_TOOLS + _gemini_tool_definitions(_folder_path_for_tools)
                else:
                    _ALL_TOOLS = _WEB_TOOLS + _TERMINAL_TOOLS + _MEMORY_TOOLS + _SCREENSHOT_TOOLS + _NOTEPAD_TOOLS + _UI_TOOLS + _NEW_TOOLS + _FOLDER_TOOLS

                today = datetime.date.today().strftime("%d %B %Y")
                _system_content = _build_system_content(
                    _folder_path_for_tools, today
                )
                if _folder_path_for_tools:
                    _system_content += f"\n\nDOSSIER ACTUELLEMENT OUVERT : {_folder_path_for_tools}"
                if selected_files:
                    _sel_basenames = [os.path.basename(f) for f in selected_files]
                    _sel_list = "\n".join(f"- {n}" for n in _sel_basenames[:50])
                    _system_content += f"\n\nFICHIERS SÉLECTIONNÉS DANS L'INTERFACE ({len(selected_files)}) :\n{_sel_list}"
                    if len(selected_files) > 50:
                        _system_content += f"\n(… et {len(selected_files) - 50} autres non listés)"
                # Limiter l'historique : 20 tours pour les modèles cloud capables, 10 pour les modèles locaux
                _history_limit = CONSTANTS.AI_HISTORY_LIMIT_CLOUD if (active_model or "").startswith(("gemini", "claude")) else CONSTANTS.AI_HISTORY_LIMIT_LOCAL
                _history = ai_conversation[-_history_limit:] if len(ai_conversation) > _history_limit else ai_conversation
                # Pour les modèles Ollama : retirer le champ "thinking" de l'historique.
                # Gemma (et la plupart des modèles locaux) n'est pas un modèle thinking natif
                # d'Ollama — ses tokens <think> passent dans msg.content. Si on renvoie le
                # contenu "thinking" extrait dans les messages suivants, Ollama le réinjecte
                # dans la fenêtre de contexte et la réponse réelle rétrécit tour après tour.
                _is_cloud = (active_model or "").startswith(("gemini", "claude"))
                _skip_keys = {"events"} if _is_cloud else {"events", "thinking"}
                messages = [
                    {"role": "system", "content": _system_content},
                    *[{k: v for k, v in m.items() if k not in _skip_keys} for m in _history],
                ]

                # ── Journal permanent en Markdown ────────────────
                _DEBUG_MD = f"{app_directory}/ai_conversations_debug.md"

                def _log_exchange_to_md(user_text: str, assistant_text: str, thinking_text: str = "", events_list: list = None) -> None:
                    """Ajoute l'échange au journal permanent en Markdown et le limite à ~500 lignes."""
                    try:
                        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        
                        # Construire le bloc Markdown de cet échange
                        block = []
                        block.append(f"## 💬 Échange du {timestamp}")
                        block.append(f"**👤 Utilisateur :**\n{user_text.strip()}\n")
                        
                        if thinking_text.strip():
                            # On préfixe chaque ligne de la réflexion par "> " pour utiliser notre superbe blockquote stylisé !
                            formatted_thinking = "\n".join(f"> {line}" for line in thinking_text.strip().splitlines())
                            block.append(f"💭 **Réflexion :**\n{formatted_thinking}\n")
                            
                        block.append(f"**🤖 Assistant :**\n{assistant_text.strip()}\n")
                        
                        if events_list:
                            block.append("🛠️ **Événements d'outils**") # Titre en gras pour faire plus propre
                            for evt in events_list:
                                block.append(f"- {evt}")
                            block.append("\n")
                            
                        block.append("---\n")
                        new_entry = "\n".join(block)
                        
                        # Lire l'historique existant
                        existing_content = ""
                        if os.path.exists(_DEBUG_MD):
                            with open(_DEBUG_MD, "r", encoding="utf-8") as f:
                                existing_content = f.read()
                        
                        # Combiner l'ancien et le nouveau journal
                        full_md = existing_content + "\n" + new_entry
                        lines = full_md.splitlines()
                        
                        # Limiter à environ 500 lignes
                        if len(lines) > 500:
                            # On prend les 500 dernières lignes
                            truncated_lines = lines[-500:]
                            
                            # On cherche le premier début d'échange "## 💬" pour couper proprement
                            start_idx = 0
                            for idx, line in enumerate(truncated_lines):
                                if line.startswith("## 💬"):
                                    start_idx = idx
                                    break
                            
                            # Si on a trouvé un début d'échange, on repart de là, sinon on prend les 500 brutes
                            final_lines = truncated_lines[start_idx:] if start_idx > 0 else truncated_lines
                            full_md = "*(Historique plus ancien tronqué pour rester sous 500 lignes)*\n\n" + "\n".join(final_lines)
                        
                        # Écrire dans le fichier permanent
                        with open(_DEBUG_MD, "w", encoding="utf-8") as f:
                            f.write(full_md.strip() + "\n")
                    except Exception:
                        pass

                # Capturer la demande originale de l'utilisateur (avant tout tour d'outil)
                # pour la réinjecter dans les rounds suivants et éviter que Gemma l'oublie.
                _original_user_request = next(
                    (m["content"] for m in reversed(messages) if m["role"] == "user"),
                    "",
                )
                if len(_original_user_request) > 400:
                    _original_user_request = _original_user_request[:400] + "…"

                # Résultat list_folder_contents conservé entre les rounds pour
                # l'auto-création si Gemma répond en texte sans appeler create_file.
                _last_folder_listing = None
                _text_response_retry_done = False
                _create_file_done = False  # True dès que create_file a été exécuté
                _read_file_done = False    # True dès que read_file_content a été exécuté
                _image_tool_done = False   # True dès qu'une génération/édition image a réussi
                _turn_events = []          # Événements d'outils du tour courant (pour export)

                # ── Boucle agentique (jusqu'à 200 tours, auto-continuation intégrée) ──
                for _tool_round in range(200):
                    # Tous les 40 tours, injecter un rappel silencieux pour que
                    # le modèle ne s'arrête pas mentalement en cours de tâche longue.
                    if _tool_round > 0 and _tool_round % 40 == 0 and ai_streaming["value"]:
                        messages.append({"role": "user", "content": (
                            "Continue la tâche en cours. "
                            "Reprends là où tu t'es arrêté."
                        )})
                    # Streaming avec thinking natif Ollama et capture des tool_calls
                    _streamed = ""
                    _thinking = ""
                    _stream_tool_calls = []
                    _text_parsed_tools = False  # True si tool_calls viennent du parseur texte
                    thinking_ctrl = None
                    _stream_token_count = 0
                    _STREAM_UPDATE_EVERY = 5

                    async def _scroll_and_update():
                        try:
                            page.update()
                            await asyncio.sleep(0)
                            await ai_chat_view.scroll_to(offset=-1)
                        except Exception:
                            pass

                    # ── Chaîne de fallback : Gemini 3.5 → Gemini 3.1 Pro → Gemma ─────
                    _fb_chain = [active_model or ""]
                    if _fb_chain[0].startswith("gemini"):
                        _c_fb = getattr(CONSTANTS, "AI_GEMINI_FALLBACK_CLOUD", "")
                        if _c_fb and _c_fb != _fb_chain[0]:
                            _fb_chain.append(_c_fb)
                    _l_fb = getattr(CONSTANTS, "AI_GEMINI_FALLBACK", "")
                    if _l_fb and _l_fb not in _fb_chain:
                        _fb_chain.append(_l_fb)

                    _fb_last_exc = None
                    _fb_skip_cloud = False
                    _fb_model_used = active_model or ""
                    for _fb_i, _fb_model in enumerate(_fb_chain):
                        if _fb_skip_cloud and _fb_model.startswith(("gemini", "claude")):
                            continue
                        if _fb_model.startswith(("gemini", "claude")):
                            _fb_tools = (_WEB_TOOLS + _TERMINAL_TOOLS + _MEMORY_TOOLS
                                         + _SCREENSHOT_TOOLS + _NOTEPAD_TOOLS + _UI_TOOLS + _NEW_TOOLS
                                         + _gemini_tool_definitions(_folder_path_for_tools))
                        else:
                            _fb_tools = (_WEB_TOOLS + _TERMINAL_TOOLS + _MEMORY_TOOLS
                                         + _SCREENSHOT_TOOLS + _NOTEPAD_TOOLS + _UI_TOOLS + _NEW_TOOLS
                                         + _FOLDER_TOOLS)
                        try:
                            if _fb_model.startswith("gemini"):
                                _stream_iter = _gemini_chat_stream_with_tools(
                                    _fb_model, messages,
                                    tools=_fb_tools, temperature=CONSTANTS.AI_TEMPERATURE)
                            elif _fb_model.startswith("claude"):
                                _stream_iter = _claude_chat_stream_with_tools(
                                    _fb_model, messages,
                                    tools=_fb_tools, temperature=CONSTANTS.AI_TEMPERATURE)
                            else:
                                _stream_iter = _ollama_chat_stream_with_tools(
                                    CONSTANTS.AI_OLLAMA_URL, _fb_model, messages,
                                    tools=_fb_tools, temperature=CONSTANTS.AI_TEMPERATURE,
                                    think=True)
                            for _evt, _dat in _stream_iter:
                                if _evt == "tool_calls":
                                    _stream_tool_calls.extend(_dat)
                                elif _evt == "thinking":
                                    _thinking += _dat
                                    if thinking_ctrl is None:
                                        _remove_loading()
                                        thinking_ctrl = _ai_add_bubble("think", _dat)
                                    else:
                                        thinking_ctrl.value = f"💭 {_thinking}"
                                        page.run_task(_scroll_and_update)
                                else:  # "token"
                                    _streamed += _dat
                                    _stream_token_count += 1
                                    _visible = re.sub(
                                        r'<think>.*?</think>', '', _streamed, flags=re.DOTALL)
                                    if '<think>' in _visible:
                                        _visible = _visible[:_visible.index('<think>')]
                                    _visible = _visible.strip()
                                    if response_text_ctrl is None:
                                        if _visible:
                                            _remove_loading()
                                            response_text_ctrl = _ai_add_bubble(
                                                "assistant", _visible)
                                    elif _stream_token_count % _STREAM_UPDATE_EVERY == 0:
                                        response_text_ctrl.value = _md_dark(_visible)
                                        page.run_task(_scroll_and_update)
                            _fb_model_used = _fb_model
                            _ALL_TOOLS = _fb_tools
                            break  # succès
                        except Exception as exc:
                            if _streamed or _stream_tool_calls or response_text_ctrl is not None:
                                raise
                            _fb_last_exc = exc
                            _fb_skip_cloud = _fb_skip_cloud or _is_network_error(exc)
                            _fb_next = next(
                                (m for m in _fb_chain[_fb_i + 1:]
                                 if not (_fb_skip_cloud and m.startswith(("gemini", "claude")))),
                                None,
                            )
                            if _fb_next is not None and loading_ctrl is not None:
                                loading_ctrl.value = (
                                    f"⚠️ {_fb_model} indisponible"
                                    f" — basculement vers {_fb_next}…"
                                )
                                try:
                                    page.update()
                                except Exception:
                                    pass
                    else:
                        if _fb_last_exc:
                            raise _fb_last_exc

                    _remove_loading()  # Garantit la suppression même si aucun contenu visible n'est arrivé
                    tool_calls = _stream_tool_calls
                    if not _streamed and not _stream_tool_calls:
                        if not _fb_model_used.startswith(("gemini", "claude")):
                            _fallback = _ollama_chat_once(
                                CONSTANTS.AI_OLLAMA_URL, _fb_model_used, messages,
                                tools=_ALL_TOOLS,
                                temperature=CONSTANTS.AI_TEMPERATURE,
                            )
                            tool_calls = _fallback.get("tool_calls") or []
                            _streamed = _fallback.get("content", "")
                            _thinking = _fallback.get("thinking", "")
                    if not tool_calls:
                        text_calls = _parse_text_tool_calls(_streamed)
                        if text_calls:
                            tool_calls = text_calls
                            _text_parsed_tools = True
                            _streamed_for_history = _streamed  # Garder <tool_code> pour l'historique Gemma
                            _streamed = _strip_text_tool_calls(_streamed)

                    if not tool_calls:
                        full_response = _strip_text_tool_calls(_streamed)
                        # Nettoyage des blocs <think> dans le contenu (think=False, Ollama ≥ 0.7)
                        if not _thinking and "<think>" in full_response:
                            _think_match = re.search(r'<think>(.*?)</think>', full_response, re.DOTALL)
                            if _think_match:
                                # Bloc complet <think>…</think>
                                _thinking = _think_match.group(1).strip()
                                full_response = re.sub(r'<think>.*?</think>', '', full_response, flags=re.DOTALL).strip()
                            else:
                                # Bloc <think> non fermé (stream interrompu en plein thinking)
                                full_response = full_response.split("<think>", 1)[0].strip()
                            if response_text_ctrl is not None:
                                response_text_ctrl.value = _md_dark(full_response)
                                try:
                                    page.update()
                                except Exception:
                                    pass
                        if _thinking and thinking_ctrl is None:
                            _ai_add_bubble("think", _thinking)
                        if _fb_model_used != (active_model or ""):
                            full_response = full_response + f"\n\n*↩ {_fb_model_used}*"
                        if response_text_ctrl is not None and full_response:
                            response_text_ctrl.value = _md_dark(full_response)
                            try:
                                page.update()
                            except Exception:
                                pass
                        elif full_response:
                            _remove_loading()
                            response_text_ctrl = _ai_add_bubble("assistant", full_response)
                        break

                    # Tour d'outils — finaliser le texte préliminaire streamé si présent
                    if response_text_ctrl is not None:
                        if _streamed:
                            response_text_ctrl.value = _md_dark(_streamed)
                            try:
                                page.update()
                            except Exception:
                                pass
                        response_text_ctrl = None

                    # ── Exécuter les appels d'outils ──────────────────────────────
                    if _text_parsed_tools:
                        # Gemma : conserver le <tool_code> dans l'assistant message pour que
                        # Gemma reconnaisse sa propre syntaxe et continue d'appeler des outils.
                        messages.append({"role": "assistant", "content": _streamed_for_history})
                    else:
                        # Modèle avec tool_calls natifs : content vide pour éviter HTTP 500.
                        messages.append({"role": "assistant", "content": "", "tool_calls": tool_calls})
                    # Afficher tous les indicateurs, collecter les tâches
                    _tool_tasks = []
                    _folder_tool_results = []  # traités séquentiellement avant le pool
                    _screenshot_b64s = []
                    for tc in tool_calls:
                        fn      = tc.get("function", {})
                        fn_name = fn.get("name", "")
                        fn_args = fn.get("arguments") or {}
                        if fn_name == "web_search":
                            query   = fn_args.get("query", "")
                            short_q = (query[:45] + "…") if len(query) > 45 else query
                            ai_status_text.value = f"🔍 {short_q}"
                            if loading_ctrl is not None:
                                loading_ctrl.value = f"🔍 {short_q}"
                            _ai_add_bubble("assistant", f"🔍 Recherche : {query}")
                            _turn_events.append(f"🔍 Recherche : {query}")
                            _tool_tasks.append((fn_name, fn_args))
                        elif fn_name == "fetch_url":
                            url     = fn_args.get("url", "")
                            short_u = (url[:45] + "…") if len(url) > 45 else url
                            ai_status_text.value = f"🌐 {short_u}"
                            if loading_ctrl is not None:
                                loading_ctrl.value = f"🌐 {short_u}"
                            _ai_add_bubble("assistant", f"🌐 Lecture : {url}")
                            _turn_events.append(f"🌐 Lecture : {url}")
                            _tool_tasks.append((fn_name, fn_args))
                        elif fn_name == "list_folder_contents":
                            _list_path = fn_args.get("path", "").strip() or _folder_path_for_tools or ""
                            _folder_display = os.path.basename(_list_path) if _list_path else "?"
                            ai_status_text.value = "📂 Lecture du dossier…"
                            _ai_add_bubble("assistant", f"📂 Lecture du dossier « {_folder_display} »")
                            _turn_events.append(f"📂 Lecture du dossier « {_folder_display} »")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append((fn_name, _folder_list_contents(_list_path)))
                        elif fn_name == "read_file_content":
                            _read_filename = fn_args.get("filename", "")
                            ai_status_text.value = f"📄 Lecture : {_read_filename}…"
                            _ai_add_bubble("assistant", f"📄 Lecture : {_read_filename}")
                            _turn_events.append(f"📄 Lecture : {_read_filename}")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append((fn_name, _folder_read_file(
                                _folder_path_for_tools, _read_filename,
                                document_exts=CONSTANTS.AI_DOCUMENT_EXTS,
                            )))
                            _read_file_done = True
                        elif fn_name == "organize_files":
                            _org_actions = fn_args.get("actions", [])
                            _org_summary = fn_args.get("summary", "")
                            if not _org_actions:
                                _folder_tool_results.append((fn_name, "Aucune action à exécuter."))
                            elif not _folder_path_for_tools:
                                _folder_tool_results.append((fn_name, "Aucun dossier ouvert."))
                            else:
                                _org_confirmed = True
                                if CONSTANTS.AI_ORGANIZE_CONFIRM:
                                    ai_status_text.value = "📂 Organisation — en attente de confirmation…"
                                    try:
                                        page.update()
                                    except Exception:
                                        pass
                                    _confirm_event  = threading.Event()
                                    _confirm_result = {"confirmed": False}

                                    def _on_org_confirm(event=None):
                                        _confirm_result["confirmed"] = True
                                        _organize_dlg.open = False
                                        page.update()
                                        _confirm_event.set()

                                    def _on_org_cancel(event=None):
                                        _organize_dlg.open = False
                                        page.update()
                                        _confirm_event.set()

                                    _action_rows = [
                                        ft.Text(
                                            f"• {_act.get('filename', '?')}  →  "
                                            f"{_act.get('destination_subfolder', '?')}/",
                                            size=12, color=WHITE,
                                        )
                                        for _act in _org_actions[:40]
                                    ]
                                    if len(_org_actions) > 40:
                                        _action_rows.append(
                                            ft.Text(f"… et {len(_org_actions) - 40} autres", size=12, color=LIGHT_GREY)
                                        )
                                    _organize_dlg = ft.AlertDialog(
                                        modal=True,
                                        title=ft.Text("📂 Organiser les fichiers"),
                                        content=ft.Column(
                                            [
                                                ft.Text(
                                                    _org_summary or "Organisation proposée par l'IA :",
                                                    size=13, color=WHITE,
                                                ),
                                                ft.Container(height=6),
                                                ft.Column(
                                                    _action_rows,
                                                    scroll=ft.ScrollMode.AUTO,
                                                    height=min(320, len(_action_rows) * 24),
                                                ),
                                            ],
                                            tight=True,
                                            width=500,
                                        ),
                                        actions=[
                                            ft.TextButton("Annuler", on_click=_on_org_cancel),
                                            ft.Button(
                                                "Exécuter",
                                                bgcolor=BLUE,
                                                color=WHITE,
                                                on_click=_on_org_confirm,
                                            ),
                                        ],
                                        actions_alignment=ft.MainAxisAlignment.END,
                                    )
                                    page.overlay.append(_organize_dlg)
                                    _organize_dlg.open = True
                                    try:
                                        page.update()
                                    except Exception:
                                        pass
                                    _confirm_event.wait(timeout=300)
                                    _org_confirmed = _confirm_result["confirmed"]
                                if not _org_confirmed:
                                    _folder_tool_results.append((fn_name, "Organisation annulée par l'utilisateur."))
                                else:
                                    _executed_moves = []
                                    _move_errors    = []
                                    for _org_action in _org_actions:
                                        _org_filename  = os.path.basename(_org_action.get("filename", ""))
                                        _org_subfolder = _org_action.get("destination_subfolder", "").strip("/\\")
                                        if not _org_filename or not _org_subfolder:
                                            continue
                                        _org_source   = os.path.join(_folder_path_for_tools, _org_filename)
                                        _org_dest_dir = os.path.join(_folder_path_for_tools, _org_subfolder)
                                        _org_dest     = os.path.join(_org_dest_dir, _org_filename)
                                        if not os.path.isfile(_org_source):
                                            _move_errors.append(f"Introuvable : {_org_filename}")
                                            continue
                                        try:
                                            os.makedirs(_org_dest_dir, exist_ok=True)
                                            shutil.move(_org_source, _org_dest)
                                            _executed_moves.append(f"✓ {_org_filename} → {_org_subfolder}/")
                                        except Exception as _move_exc:
                                            _move_errors.append(f"✗ {_org_filename} : {_move_exc}")
                                    page.pubsub.send_all_on_topic("refresh", None)
                                    _org_result_lines = [f"{len(_executed_moves)} fichier(s) déplacé(s)."] + _executed_moves
                                    if _move_errors:
                                        _org_result_lines += ["Erreurs :"] + _move_errors
                                    _folder_tool_results.append((fn_name, "\n".join(_org_result_lines)))
                        elif fn_name == "analyze_images":
                            _analyze_filenames = fn_args.get("filenames", [])
                            _analyze_question  = fn_args.get("question", "")
                            # Résoudre la liste d'images à analyser
                            if not _analyze_filenames:
                                _analyze_candidates = sorted([
                                    entry.name for entry in os.scandir(_folder_path_for_tools)
                                    if entry.is_file()
                                    and os.path.splitext(entry.name)[1].lower() in CONSTANTS.IMAGE_EXTS
                                ])
                            else:
                                _analyze_candidates = [
                                    os.path.basename(fname) for fname in _analyze_filenames
                                    if os.path.isfile(
                                        os.path.join(_folder_path_for_tools, os.path.basename(fname))
                                    )
                                ]
                            if not _analyze_candidates:
                                _folder_tool_results.append((fn_name, "Aucune image trouvée."))
                            else:
                                _analyze_total   = len(_analyze_candidates)
                                _analyze_model   = ai_model_dropdown.value or CONSTANTS.AI_MODEL_VISION
                                _analyze_batch_n = (
                                    CONSTANTS.AI_GEMINI_FOLDER_BATCH_SIZE
                                    if (_analyze_model or "").startswith("gemini")
                                    else CONSTANTS.AI_FOLDER_SELECT_BATCH_SIZE
                                )
                                _analyze_batches = (_analyze_total + _analyze_batch_n - 1) // _analyze_batch_n
                                _analysis_progress_ctrl = _ai_add_bubble(
                                    "assistant",
                                    f"📸 Analyse de {_analyze_total} image(s) — lot 1/{_analyze_batches}…",
                                )
                                _turn_events.append(f"📸 Analyse de {_analyze_total} image(s)")
                                def _on_analyze_progress(batch_num, total_batches):
                                    ai_status_text.value = f"📸 Analyse lot {batch_num}/{total_batches}…"
                                    if _analysis_progress_ctrl:
                                        _analysis_progress_ctrl.value = _md_dark(
                                            f"📸 Analyse — lot {batch_num}/{total_batches}…"
                                        )
                                    page.run_task(_scroll_and_update)
                                _analyze_results = _analyze_images_batched(
                                    CONSTANTS.AI_OLLAMA_URL,
                                    _analyze_model,
                                    _folder_path_for_tools,
                                    _analyze_candidates,
                                    _analyze_question,
                                    batch_size=_analyze_batch_n,
                                    image_exts=CONSTANTS.IMAGE_EXTS,
                                    max_size=CONSTANTS.AI_FOLDER_SELECT_IMAGE_SIZE,
                                    quality=CONSTANTS.AI_FOLDER_SELECT_QUALITY,
                                    on_progress=_on_analyze_progress,
                                    is_running=lambda: ai_streaming["value"],
                                )
                                if _analysis_progress_ctrl:
                                    _analysis_progress_ctrl.value = _md_dark(
                                        f"📸 {len(_analyze_candidates)} image(s) analysée(s)."
                                    )
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append(
                                    (fn_name, "\n\n".join(_analyze_results) or "Aucun résultat.")
                                )
                        elif fn_name in ("generate_image", "edit_image"):
                            if _image_tool_done:
                                _folder_tool_results.append(
                                    (fn_name, "Action ignorée : une image a déjà été générée/modifiée pour cette demande.")
                                )
                                continue
                            # Une seule tentative image par demande utilisateur pour éviter
                            # les boucles de prompts (réessais en chaîne côté modèle).
                            _image_tool_done = True
                            import datetime as _dt_gi
                            _gi_prompt     = fn_args.get("prompt", "")
                            _gi_aspect     = fn_args.get("aspect_ratio", "1:1")
                            _gi_resolution = fn_args.get("resolution", "1K")
                            _gi_src_name   = ""
                            # Fichier de sortie
                            if fn_name == "generate_image":
                                _gi_out_filename = (
                                    fn_args.get("filename", "").strip()
                                    or f"generated_{_dt_gi.datetime.now():%Y%m%d_%H%M%S}.png"
                                )
                                _gi_src_bytes = None
                                _gi_label = _gi_prompt[:60] + ("…" if len(_gi_prompt) > 60 else "")
                                _ai_add_bubble("assistant", f"🎨 Génération : {_gi_label}")
                                _turn_events.append(f"🎨 Génération : {_gi_label}")
                            else:  # edit_image
                                _gi_src_name = fn_args.get("source_filename", "").strip()
                                _gi_out_filename = (
                                    fn_args.get("output_filename", "").strip()
                                    or f"edited_{_dt_gi.datetime.now():%Y%m%d_%H%M%S}.png"
                                )
                                _gi_src_bytes = None
                                if _gi_src_name and _folder_path_for_tools:
                                    _gi_src_path = os.path.join(
                                        _folder_path_for_tools, os.path.basename(_gi_src_name)
                                    )
                                    if os.path.isfile(_gi_src_path):
                                        with open(_gi_src_path, "rb") as _f:
                                            _gi_src_bytes = _f.read()
                                _ai_add_bubble("assistant", f"🎨 Édition : {_gi_src_name} → {_gi_out_filename}")
                                _turn_events.append(f"🎨 Édition : {_gi_src_name} → {_gi_out_filename}")

                            _gi_prompt_refined = _gi_prompt
                            if (active_model or "").startswith("gemini") and _gi_prompt.strip():
                                _gi_prompt_refined = _gemini_refine_image_prompt(
                                    intent_prompt=_gi_prompt,
                                    user_request=_original_user_request,
                                    mode=fn_name,
                                    source_filename=_gi_src_name,
                                    model=active_model,
                                )
                                if _gi_prompt_refined != _gi_prompt:
                                    if CONSTANTS.AI_SHOW_REFINED_IMAGE_PROMPT:
                                        _ai_add_bubble(
                                            "assistant",
                                            "🧪 Prompt image affiné automatiquement :\n\n"
                                            f"{_gi_prompt_refined}",
                                        )
                                    else:
                                        _ai_add_bubble("assistant", "🧪 Prompt image affiné automatiquement.")

                            # Journaliser le prompt réellement envoyé pour qu'il soit
                            # inclus dans l'export/copie de conversation IA.
                            _turn_events.append(f"🧪 Prompt image : {_gi_prompt_refined}")

                            ai_status_text.value = "🎨 Génération d'image en cours…"
                            ai_progress_bar.visible = True
                            try:
                                page.update()
                            except Exception:
                                pass
                            _gi_timeout_seconds = int(getattr(CONSTANTS, "AI_GEMINI_IMAGE_TIMEOUT", 180))
                            _gi_result_holder = {"value": ("[ERREUR] Timeout Gemini image.", None)}
                            _gi_done_event = threading.Event()

                            def _run_gemini_image_call():
                                try:
                                    _gi_result_holder["value"] = _gemini_generate_image(
                                        _gi_prompt_refined,
                                        input_image_bytes=_gi_src_bytes,
                                        aspect_ratio=_gi_aspect,
                                        resolution=_gi_resolution,
                                    )
                                except Exception as _gi_exc:
                                    _gi_result_holder["value"] = (f"[ERREUR] {str(_gi_exc)}", None)
                                finally:
                                    _gi_done_event.set()

                            threading.Thread(target=_run_gemini_image_call, daemon=True).start()
                            _gi_start_time = time.time()
                            _gi_elapsed_s = 0
                            while not _gi_done_event.wait(timeout=1.0):
                                _gi_elapsed_s = int(time.time() - _gi_start_time)
                                if _gi_elapsed_s >= _gi_timeout_seconds:
                                    break
                                ai_status_text.value = f"🎨 Génération d'image en cours… ({_gi_elapsed_s}s)"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                            if not _gi_done_event.is_set():
                                _gi_text, _gi_bytes = (
                                    f"[ERREUR] Timeout Gemini image après {_gi_timeout_seconds}s.",
                                    None,
                                )
                            else:
                                _gi_text, _gi_bytes = _gi_result_holder["value"]
                            ai_progress_bar.visible = False
                            if _gi_bytes:
                                _gi_dest_folder = _folder_path_for_tools or os.path.join(app_directory, "Generated")
                                os.makedirs(_gi_dest_folder, exist_ok=True)
                                _gi_save_path = os.path.join(_gi_dest_folder, _gi_out_filename)
                                with open(_gi_save_path, "wb") as _fout:
                                    _fout.write(_gi_bytes)
                                _ai_add_image_bubble(_gi_save_path)
                                if _folder_path_for_tools:
                                    page.pubsub.send_all_on_topic("refresh", None)
                                _gi_result = f"Image sauvegardée : {_gi_save_path}"
                                if _gi_text:
                                    _gi_result += (
                                        "\n\n"
                                        f"Réponse du service : {_gi_text}\n"
                                        f"Fichier : {_gi_save_path}"
                                    )
                                _turn_events.append(f"✅ Image sauvegardée : {_gi_out_filename}")
                            else:
                                _gi_result = "[ERREUR] Aucune image n'a été générée/sauvegardée."
                                if _gi_text:
                                    _gi_result += (
                                        "\n\nRéponse texte du service (sans image):\n"
                                        f"{_gi_text}"
                                    )
                                _turn_events.append("❌ Échec génération/édition image (aucun fichier créé)")
                            _folder_tool_results.append((fn_name, _gi_result))
                        elif fn_name == "generate_music":
                            import datetime as _dt_gm
                            _gm_prompt   = fn_args.get("prompt", "")
                            _gm_model    = fn_args.get("model", "lyria-3-clip-preview")
                            _gm_filename = (
                                fn_args.get("filename", "").strip()
                                or f"music_{_dt_gm.datetime.now():%Y%m%d_%H%M%S}.mp3"
                            )
                            _gm_label = _gm_prompt[:60] + ("…" if len(_gm_prompt) > 60 else "")
                            _ai_add_bubble("assistant", f"🎵 Génération musique : {_gm_label}")
                            _turn_events.append(f"🎵 Génération musique : {_gm_label}")
                            ai_status_text.value = "🎵 Génération musicale en cours…"
                            ai_progress_bar.visible = True
                            try:
                                page.update()
                            except Exception:
                                pass
                            _gm_result_holder = {"value": (None, None, "Timeout")}
                            _gm_done_event = threading.Event()

                            def _run_music_call():
                                try:
                                    _gm_result_holder["value"] = _gemini_generate_music(
                                        _gm_prompt, model=_gm_model
                                    )
                                except Exception as _gm_exc:
                                    _gm_result_holder["value"] = (None, None, str(_gm_exc))
                                finally:
                                    _gm_done_event.set()

                            threading.Thread(target=_run_music_call, daemon=True).start()
                            _gm_start = time.time()
                            while not _gm_done_event.wait(timeout=1.0):
                                _gm_elapsed = int(time.time() - _gm_start)
                                if _gm_elapsed >= 180:
                                    break
                                ai_status_text.value = (
                                    f"🎵 Génération musicale en cours… ({_gm_elapsed}s)"
                                )
                                try:
                                    page.update()
                                except Exception:
                                    pass
                            ai_progress_bar.visible = False
                            _gm_bytes, _gm_lyrics, _gm_err = _gm_result_holder["value"]
                            if _gm_bytes:
                                _gm_dest = _folder_path_for_tools or os.path.join(
                                    app_directory, "Generated"
                                )
                                os.makedirs(_gm_dest, exist_ok=True)
                                _gm_save_path = os.path.join(_gm_dest, _gm_filename)
                                with open(_gm_save_path, "wb") as _fout:
                                    _fout.write(_gm_bytes)
                                if _folder_path_for_tools:
                                    page.pubsub.send_all_on_topic("refresh", None)
                                _gm_result = f"Musique sauvegardée : {_gm_save_path}"
                                if _gm_lyrics:
                                    _gm_result += f"\n\nParoles / Structure :\n{_gm_lyrics}"
                                _turn_events.append(f"✅ Musique sauvegardée : {_gm_filename}")
                            else:
                                _gm_result = f"[ERREUR] Génération musicale échouée : {_gm_err}"
                                _turn_events.append("❌ Échec génération musicale")
                            _folder_tool_results.append((fn_name, _gm_result))
                        elif fn_name == "create_file":
                            import datetime as _dt_cf
                            _create_filename = fn_args.get("filename", "").strip()
                            if not _create_filename:
                                _create_filename = f"fichier_{_dt_cf.datetime.now():%Y%m%d_%H%M%S}.txt"
                            # Ne JAMAIS écraser un contenu déjà fourni (script, note, etc.)
                            # par un listing de dossier. Le fallback listing ne s'applique
                            # que si le modèle a envoyé un contenu vide.
                            _create_content_raw = _clean_file_content(fn_args.get("content", ""))
                            if _create_content_raw.strip():
                                _create_content = _create_content_raw
                            elif _last_folder_listing and not _read_file_done:
                                _create_content = _last_folder_listing
                            else:
                                _create_content = ""

                            # Garde-fou : éviter d'écrire un listing de dossier dans un script.
                            _create_ext = os.path.splitext(_create_filename)[1].lower()
                            _script_exts = {".py", ".pyw", ".sh", ".bat", ".ps1"}
                            if (
                                _create_ext in _script_exts
                                and _create_content.strip().startswith("Dossier :")
                            ):
                                _folder_tool_results.append(
                                    (
                                        fn_name,
                                        "Création annulée : contenu invalide pour un script (listing de dossier détecté).",
                                    )
                                )
                                continue
                            ai_status_text.value = f"📝 Création : {_create_filename}…"
                            _ai_add_bubble("assistant", f"📝 Création du fichier : {_create_filename}")
                            _turn_events.append(f"📝 Création du fichier : {_create_filename}")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _create_result = _folder_create_file(
                                _folder_path_for_tools, _create_filename, _create_content
                            )
                            page.pubsub.send_all_on_topic("refresh", None)
                            _create_file_done = True
                            _folder_tool_results.append((fn_name, _create_result))
                        elif fn_name == "run_terminal_command":
                            _cmd      = fn_args.get("command", "")
                            _cmd_desc = fn_args.get("description", _cmd)
                            _cwd = _folder_path_for_tools if _folder_path_for_tools else None
                            if CONSTANTS.AI_TERMINAL_CONFIRM:
                                _cmd_confirm_event  = threading.Event()
                                _cmd_confirm_result = {"confirmed": False}

                                def _on_cmd_confirm(event=None):
                                    _cmd_confirm_result["confirmed"] = True
                                    _cmd_dlg.open = False
                                    page.update()
                                    _cmd_confirm_event.set()

                                def _on_cmd_cancel(event=None):
                                    _cmd_dlg.open = False
                                    page.update()
                                    _cmd_confirm_event.set()

                                _cmd_dlg = ft.AlertDialog(
                                    modal=True,
                                    title=ft.Text("💻 Exécuter une commande"),
                                    content=ft.Column(
                                        [
                                            ft.Text(_cmd_desc, size=13, color=WHITE),
                                            ft.Container(height=8),
                                            ft.Container(
                                                ft.Text(_cmd, size=12, font_family="monospace", color=YELLOW),
                                                bgcolor=DARK,
                                                padding=10,
                                                border_radius=6,
                                            ),
                                        ],
                                        tight=True,
                                        width=500,
                                    ),
                                    actions=[
                                        ft.TextButton("Annuler", on_click=_on_cmd_cancel),
                                        ft.Button(
                                            "Exécuter",
                                            bgcolor=BLUE,
                                            color=WHITE,
                                            on_click=_on_cmd_confirm,
                                        ),
                                    ],
                                    actions_alignment=ft.MainAxisAlignment.END,
                                )
                                page.overlay.append(_cmd_dlg)
                                _cmd_dlg.open = True
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _cmd_confirm_event.wait(timeout=300)
                                if not _cmd_confirm_result["confirmed"]:
                                    _folder_tool_results.append((fn_name, "Commande annulée par l'utilisateur."))
                                    continue
                            ai_status_text.value = "💻 Exécution en cours…"
                            _turn_events.append(f"💻 Commande : {_cmd}")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append(
                                (fn_name, _run_terminal_command(_cmd, cwd=_cwd))
                            )
                        elif fn_name == "update_memory_file":
                            _mem_target  = fn_args.get("target", "")
                            _mem_action  = fn_args.get("action", "")
                            _mem_content = fn_args.get("content", "")
                            _mem_old     = fn_args.get("old_text", "")
                            ai_status_text.value = f"🧠 Mise à jour mémoire ({_mem_target})…"
                            _turn_events.append(f"🧠 Mémoire : {_mem_action} → {_mem_target}")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append(
                                (fn_name, _update_memory_file(_mem_target, _mem_action, _mem_content, _mem_old))
                            )
                        elif fn_name == "read_notepad":
                            _np_current = notepad_field.value or ""
                            ai_status_text.value = "📝 Lecture du bloc-notes…"
                            _ai_add_bubble("assistant", "📝 Lecture du bloc-notes")
                            _turn_events.append("📝 Lecture du bloc-notes")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append(
                                (fn_name, _np_current if _np_current.strip() else "(Bloc-notes vide)")
                            )
                        elif fn_name == "write_notepad":
                            _np_content = fn_args.get("content", "")
                            _np_action  = fn_args.get("action", "replace")
                            if _np_action == "replace":
                                notepad_field.value = _np_content
                            elif _np_action == "append":
                                notepad_field.value = (notepad_field.value or "") + "\n" + _np_content
                            elif _np_action == "prepend":
                                notepad_field.value = _np_content + "\n" + (notepad_field.value or "")
                            if notepad_is_preview["value"]:
                                notepad_markdown_preview.value = _prepare_notepad_markdown(notepad_field.value or "")
                            try:
                                notepad_field.update()
                                if notepad_is_preview["value"]:
                                    notepad_markdown_preview.update()
                            except Exception:
                                pass
                            ai_status_text.value = "📝 Bloc-notes mis à jour"
                            _ai_add_bubble("assistant", f"📝 Bloc-notes mis à jour ({_np_action})")
                            _turn_events.append(f"📝 Bloc-notes mis à jour ({_np_action})")
                            _folder_tool_results.append(
                                (fn_name, f"Bloc-notes mis à jour ({_np_action}). Longueur : {len(notepad_field.value or '')} caractères.")
                            )
                        elif fn_name == "take_screenshot":
                            _ss_region = fn_args.get("region") or None
                            ai_status_text.value = "📸 Capture d'écran…"
                            _ai_add_bubble("assistant", "📸 Capture d'écran")
                            _turn_events.append("📸 Capture d'écran")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _ss_capture = _take_screenshot(region=_ss_region)
                            if _ss_capture:
                                _screenshot_b64s.append(_ss_capture["b64"])
                                _folder_tool_results.append((fn_name, _ss_capture["text"]))
                            else:
                                _folder_tool_results.append((fn_name, "Échec de la capture d'écran."))
                        elif fn_name == "mouse_click":
                            _mc_x      = int(fn_args.get("x", 0))
                            _mc_y      = int(fn_args.get("y", 0))
                            _mc_button = fn_args.get("button", "left")
                            _mc_clicks = fn_args.get("clicks", 1)
                            ai_status_text.value = f"🖱️ Clic ({_mc_x}, {_mc_y})…"
                            _ai_add_bubble("assistant", f"🖱️ Clic {_mc_button} à ({_mc_x}, {_mc_y})")
                            _turn_events.append(f"🖱️ Clic à ({_mc_x}, {_mc_y})")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append(
                                (fn_name, _mouse_click(_mc_x, _mc_y, _mc_button, _mc_clicks))
                            )
                        elif fn_name == "keyboard_type":
                            _kt_text = fn_args.get("text", "")
                            _kt_short = (_kt_text[:30] + "…") if len(_kt_text) > 30 else _kt_text
                            ai_status_text.value = f"⌨️ Saisie : {_kt_short}…"
                            _ai_add_bubble("assistant", f"⌨️ Saisie : « {_kt_short} »")
                            _turn_events.append(f"⌨️ Saisie : « {_kt_short} »")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append((fn_name, _keyboard_type(_kt_text)))
                        elif fn_name == "keyboard_hotkey":
                            _kh_keys = fn_args.get("keys", [])
                            _kh_str  = "+".join(_kh_keys)
                            ai_status_text.value = f"⌨️ Raccourci : {_kh_str}…"
                            _ai_add_bubble("assistant", f"⌨️ Raccourci : {_kh_str}")
                            _turn_events.append(f"⌨️ Raccourci : {_kh_str}")
                            try:
                                page.update()
                            except Exception:
                                pass
                            _folder_tool_results.append(
                                (fn_name, _keyboard_hotkey(*_kh_keys))
                            )
                        elif fn_name == "navigate_to_folder":
                            _nav_path = fn_args.get("path", "").strip()
                            if not _nav_path or not os.path.isdir(_nav_path):
                                _folder_tool_results.append((fn_name, f"Dossier introuvable : {_nav_path or '(vide)'}"))
                            else:
                                navigate_to_folder(_nav_path)
                                _folder_path_for_tools = _nav_path
                                ai_status_text.value = f"📂 Navigation → {os.path.basename(_nav_path)}"
                                _ai_add_bubble("assistant", f"📂 Navigation vers : {_nav_path}")
                                _turn_events.append(f"📂 Navigation vers : {_nav_path}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, f"Dossier ouvert : {_nav_path}"))
                        elif fn_name == "select_files_in_ui":
                            _sel_names = fn_args.get("filenames", [])
                            _sel_mode  = fn_args.get("mode", "replace")
                            _folder_for_sel = _folder_path_for_tools or ""
                            if not _folder_for_sel:
                                _folder_tool_results.append((fn_name, "Aucun dossier ouvert pour sélectionner des fichiers."))
                            else:
                                if _sel_mode == "replace":
                                    selected_files.clear()
                                _changed = 0
                                for _sel_name in _sel_names:
                                    _sel_path = os.path.join(_folder_for_sel, os.path.basename(_sel_name))
                                    if os.path.exists(_sel_path):
                                        if _sel_mode == "remove":
                                            if _sel_path in selected_files:
                                                selected_files.remove(_sel_path)
                                                _changed += 1
                                        else:
                                            if _sel_path not in selected_files:
                                                selected_files.append(_sel_path)
                                                _changed += 1
                                page.pubsub.send_all_on_topic("refresh", None)
                                _verb = "retiré(s)" if _sel_mode == "remove" else "sélectionné(s)"
                                ai_status_text.value = f"✅ {_changed} fichier(s) {_verb}"
                                _ai_add_bubble("assistant", f"✅ {_changed} fichier(s) {_verb}. Sélection totale : {len(selected_files)} fichier(s).")
                                _turn_events.append(f"✅ Sélection : {_changed} fichier(s) {_verb}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append(
                                    (fn_name, f"{_changed} fichier(s) {_verb}. Sélection totale : {len(selected_files)} fichier(s).")
                                )
                        elif fn_name == "delete_files":
                            _del_paths   = fn_args.get("paths", [])
                            _del_summary = fn_args.get("summary", "")
                            if not _del_paths:
                                _folder_tool_results.append((fn_name, "Aucun fichier à supprimer."))
                            else:
                                _del_confirmed = True
                                if CONSTANTS.AI_DELETE_CONFIRM:
                                    ai_status_text.value = "🗑️ Suppression — en attente de confirmation…"
                                    try:
                                        page.update()
                                    except Exception:
                                        pass
                                    _del_event  = threading.Event()
                                    _del_result = {"confirmed": False}

                                    def _on_del_confirm(event=None):
                                        _del_result["confirmed"] = True
                                        _del_dlg.open = False
                                        page.update()
                                        _del_event.set()

                                    def _on_del_cancel(event=None):
                                        _del_dlg.open = False
                                        page.update()
                                        _del_event.set()

                                    _del_rows = [
                                        ft.Text(f"• {p}", size=12, color=WHITE)
                                        for p in _del_paths[:40]
                                    ]
                                    if len(_del_paths) > 40:
                                        _del_rows.append(ft.Text(f"… et {len(_del_paths) - 40} autres", size=12, color=LIGHT_GREY))
                                    _del_dlg = ft.AlertDialog(
                                        modal=True,
                                        title=ft.Text("🗑️ Supprimer des fichiers"),
                                        content=ft.Column(
                                            [
                                                ft.Text(_del_summary or "Fichiers à supprimer :", size=13, color=WHITE),
                                                ft.Container(height=6),
                                                ft.Column(
                                                    _del_rows,
                                                    scroll=ft.ScrollMode.AUTO,
                                                    height=min(320, len(_del_rows) * 24),
                                                ),
                                            ],
                                            tight=True,
                                            width=500,
                                        ),
                                        actions=[
                                            ft.TextButton("Annuler", on_click=_on_del_cancel),
                                            ft.ElevatedButton(
                                                "Supprimer",
                                                bgcolor=ft.Colors.RED_700,
                                                color=WHITE,
                                                on_click=_on_del_confirm,
                                            ),
                                        ],
                                        actions_alignment=ft.MainAxisAlignment.END,
                                    )
                                    page.overlay.append(_del_dlg)
                                    _del_dlg.open = True
                                    try:
                                        page.update()
                                    except Exception:
                                        pass
                                    _del_event.wait(timeout=300)
                                    _del_confirmed = _del_result["confirmed"]
                                if not _del_confirmed:
                                    _folder_tool_results.append((fn_name, "Suppression annulée."))
                                else:
                                    _del_res = _folder_delete_files(_folder_path_for_tools, _del_paths)
                                    page.pubsub.send_all_on_topic("refresh", None)
                                    ai_status_text.value = "🗑️ Suppression effectuée"
                                    _ai_add_bubble("assistant", f"🗑️ Suppression effectuée")
                                    _turn_events.append("🗑️ Suppression effectuée")
                                    try:
                                        page.update()
                                    except Exception:
                                        pass
                                    _folder_tool_results.append((fn_name, _del_res))
                        elif fn_name == "move_file":
                            _mv_src = fn_args.get("source", "").strip()
                            _mv_dst = fn_args.get("destination", "").strip()
                            if not _mv_src or not _mv_dst:
                                _folder_tool_results.append((fn_name, "Paramètres source ou destination manquants."))
                            else:
                                _mv_res = _folder_move_file(_folder_path_for_tools, _mv_src, _mv_dst)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _ai_add_bubble("assistant", f"📦 Déplacement : {os.path.basename(_mv_src)} → {_mv_dst}")
                                _turn_events.append(f"📦 Déplacement : {os.path.basename(_mv_src)} → {_mv_dst}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, _mv_res))
                        elif fn_name == "copy_file":
                            _cp_src = fn_args.get("source", "").strip()
                            _cp_dst = fn_args.get("destination", "").strip()
                            if not _cp_src or not _cp_dst:
                                _folder_tool_results.append((fn_name, "Paramètres source ou destination manquants."))
                            else:
                                _cp_res = _folder_copy_file(_folder_path_for_tools, _cp_src, _cp_dst)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _ai_add_bubble("assistant", f"📋 Copie : {os.path.basename(_cp_src)} → {_cp_dst}")
                                _turn_events.append(f"📋 Copie : {os.path.basename(_cp_src)} → {_cp_dst}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, _cp_res))
                        elif fn_name == "create_folder":
                            _mkdir_path = fn_args.get("path", "").strip()
                            if not _mkdir_path:
                                _folder_tool_results.append((fn_name, "Chemin manquant."))
                            else:
                                _mkdir_res = _folder_create_folder(_folder_path_for_tools, _mkdir_path)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _ai_add_bubble("assistant", f"📁 Dossier créé : {_mkdir_path}")
                                _turn_events.append(f"📁 Dossier créé : {_mkdir_path}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, _mkdir_res))
                        elif fn_name == "read_exif":
                            _exif_files = fn_args.get("filenames", [])
                            if not _exif_files:
                                _folder_tool_results.append((fn_name, "Aucun fichier fourni."))
                            else:
                                _exif_res = _folder_read_exif(_folder_path_for_tools, _exif_files)
                                _folder_tool_results.append((fn_name, _exif_res))
                        elif fn_name == "zip_files":
                            _zip_paths = fn_args.get("paths", [])
                            _zip_name = fn_args.get("zip_name", "archive") or "archive"
                            _zip_dest = fn_args.get("destination", "") or None
                            if not _zip_paths:
                                _folder_tool_results.append((fn_name, "Aucun fichier fourni."))
                            else:
                                _zip_res = _folder_zip_files(_folder_path_for_tools, _zip_paths, _zip_name, _zip_dest)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _ai_add_bubble("assistant", f"🗜️ Archive créée : {_zip_name}")
                                _turn_events.append(f"🗜️ Archive créée : {_zip_name}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, _zip_res))
                        elif fn_name == "unzip_file":
                            _unzip_src = fn_args.get("source", "").strip()
                            _unzip_dest = fn_args.get("destination", "") or None
                            if not _unzip_src:
                                _folder_tool_results.append((fn_name, "Source manquante."))
                            else:
                                _unzip_res = _folder_unzip_file(_folder_path_for_tools, _unzip_src, _unzip_dest)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _ai_add_bubble("assistant", f"📦 Extrait : {os.path.basename(_unzip_src)}")
                                _turn_events.append(f"📦 Extrait : {os.path.basename(_unzip_src)}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append((fn_name, _unzip_res))
                        elif fn_name == "edit_file":
                            _ef_path = fn_args.get("filepath", "").strip()
                            _ef_old  = fn_args.get("old_string", "")
                            _ef_new  = fn_args.get("new_string", "")
                            if not _ef_path or _ef_old == "":
                                _folder_tool_results.append((fn_name, "Paramètres filepath / old_string manquants."))
                            else:
                                ai_status_text.value = f"✏️ Édition : {os.path.basename(_ef_path)}…"
                                _ai_add_bubble("assistant", f"✏️ Édition : {_ef_path}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _ef_res = _edit_file(_folder_path_for_tools, _ef_path, _ef_old, _ef_new)
                                page.pubsub.send_all_on_topic("refresh", None)
                                _folder_tool_results.append((fn_name, _ef_res))
                        elif fn_name == "read_file_lines":
                            _rl_path  = fn_args.get("filepath", "").strip()
                            _rl_start = fn_args.get("start_line", 1)
                            _rl_end   = fn_args.get("end_line", None)
                            if not _rl_path:
                                _folder_tool_results.append((fn_name, "Paramètre filepath manquant."))
                            else:
                                _rl_end_str = str(_rl_end) if _rl_end else "fin"
                                ai_status_text.value = f"📄 Lignes {_rl_start}–{_rl_end_str} : {os.path.basename(_rl_path)}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _folder_tool_results.append(
                                    (fn_name, _read_file_lines(_folder_path_for_tools, _rl_path, _rl_start, _rl_end))
                                )
                        elif fn_name == "search_in_files":
                            _si_pattern = fn_args.get("pattern", "")
                            _si_path    = (fn_args.get("path", "") or "").strip() or None
                            _si_glob    = fn_args.get("file_glob", "*") or "*"
                            _si_max     = int(fn_args.get("max_results", 50) or 50)
                            _si_case    = bool(fn_args.get("case_sensitive", False))
                            if not _si_pattern:
                                _folder_tool_results.append((fn_name, "Paramètre 'pattern' manquant."))
                            else:
                                ai_status_text.value = f"🔎 Grep : {_si_pattern}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _si_res = _search_in_files(
                                    _folder_path_for_tools, _si_pattern,
                                    path=_si_path, file_glob=_si_glob,
                                    max_results=_si_max, case_sensitive=_si_case,
                                )
                                _folder_tool_results.append((fn_name, _si_res))
                        elif fn_name == "find_files":
                            _ff_pattern  = fn_args.get("pattern", "")
                            _ff_basepath = (fn_args.get("base_path", "") or "").strip() or None
                            _ff_max      = int(fn_args.get("max_results", 200) or 200)
                            if not _ff_pattern:
                                _folder_tool_results.append((fn_name, "Paramètre 'pattern' manquant."))
                            else:
                                ai_status_text.value = f"🔎 Glob : {_ff_pattern}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _ff_res = _find_files(
                                    _folder_path_for_tools, _ff_pattern,
                                    base_path=_ff_basepath, max_results=_ff_max,
                                )
                                _folder_tool_results.append((fn_name, _ff_res))
                        elif fn_name == "git_command":
                            _git_args = fn_args.get("args", [])
                            _git_cwd  = (fn_args.get("cwd", "") or "").strip() or _folder_path_for_tools or None
                            if not _git_args:
                                _folder_tool_results.append((fn_name, "Paramètre 'args' manquant."))
                            else:
                                _git_label = " ".join(str(a) for a in _git_args[:3])
                                ai_status_text.value = f"🔀 git {_git_label}…"
                                _ai_add_bubble("assistant", f"🔀 git {' '.join(str(a) for a in _git_args)}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _git_res = _git_command(_git_args, cwd=_git_cwd)
                                _folder_tool_results.append((fn_name, _git_res))
                        elif fn_name == "manage_tasks":
                            _task_res = _manage_tasks(
                                fn_args.get("action", "list"),
                                task_id=fn_args.get("task_id") or None,
                                title=fn_args.get("title") or None,
                                status=fn_args.get("status") or None,
                                notes=fn_args.get("notes") or None,
                            )
                            _folder_tool_results.append((fn_name, _task_res))
                        elif fn_name == "read_pdf":
                            _pdf_path  = fn_args.get("filepath", "").strip()
                            _pdf_pages = fn_args.get("pages") or None
                            if not _pdf_path:
                                _folder_tool_results.append((fn_name, "Paramètre 'filepath' manquant."))
                            else:
                                ai_status_text.value = f"📄 PDF : {os.path.basename(_pdf_path)}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _pdf_res = _read_pdf(_folder_path_for_tools, _pdf_path, pages=_pdf_pages)
                                _folder_tool_results.append((fn_name, _pdf_res))
                        elif fn_name == "ask_subagent":
                            _sa_task    = fn_args.get("task", "")
                            _sa_context = fn_args.get("context") or None
                            _sa_model   = fn_args.get("model") or None
                            if not _sa_task:
                                _folder_tool_results.append((fn_name, "Paramètre 'task' manquant."))
                            else:
                                _sa_short = (_sa_task[:50] + "…") if len(_sa_task) > 50 else _sa_task
                                ai_status_text.value = f"🤖 Sous-agent : {_sa_short}…"
                                _ai_add_bubble("assistant", f"🤖 Sous-agent : {_sa_short}")
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _sa_res = _ask_subagent(_sa_task, context=_sa_context, model=_sa_model)
                                _folder_tool_results.append((fn_name, _sa_res))
                        elif fn_name == "schedule_task":
                            ai_status_text.value = f"⏰ Planificateur : {fn_args.get('action', 'list')}…"
                            try:
                                page.update()
                            except Exception:
                                pass
                            _sched_res = _schedule_task(
                                fn_args.get("action", "list"),
                                name=fn_args.get("name") or None,
                                command=fn_args.get("command") or None,
                                when=fn_args.get("when") or None,
                            )
                            _folder_tool_results.append((fn_name, _sched_res))
                        elif fn_name == "http_request":
                            _hr_method = (fn_args.get("method", "GET") or "GET").upper()
                            _hr_url = fn_args.get("url", "").strip()
                            if not _hr_url:
                                _folder_tool_results.append((fn_name, "Paramètre 'url' manquant."))
                            else:
                                ai_status_text.value = f"🌐 {_hr_method} {_hr_url[:60]}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _hr_res = _http_request(
                                    _hr_method, _hr_url,
                                    headers=fn_args.get("headers") or None,
                                    body=fn_args.get("body") or None,
                                    timeout=fn_args.get("timeout") or 30,
                                )
                                _folder_tool_results.append((fn_name, _hr_res))
                        elif fn_name == "read_spreadsheet":
                            _ss_path = fn_args.get("filepath", "").strip()
                            if not _ss_path:
                                _folder_tool_results.append((fn_name, "Paramètre 'filepath' manquant."))
                            else:
                                ai_status_text.value = f"📊 Tableur : {os.path.basename(_ss_path)}…"
                                try:
                                    page.update()
                                except Exception:
                                    pass
                                _ss_res = _read_spreadsheet(
                                    _folder_path_for_tools, _ss_path,
                                    sheet=fn_args.get("sheet") or None,
                                    max_rows=fn_args.get("max_rows") or 100,
                                )
                                _folder_tool_results.append((fn_name, _ss_res))
                    try:
                        page.update()
                    except Exception:
                        pass
                    # ── Injecter les résultats d'outils dans l'historique ────────────────
                    # Exécuter les outils web/URL en parallèle
                    def _run_tool(task):
                        name, args = task
                        if name == "web_search":
                            return _web_search(args.get("query", ""))
                        elif name == "fetch_url":
                            return _fetch_url_content(args.get("url", ""), max_chars=CONSTANTS.AI_URL_MAX_CHARS)
                        return ""
                    with concurrent.futures.ThreadPoolExecutor() as _pool:
                        _web_tool_results = list(_pool.map(_run_tool, _tool_tasks))
                    _all_tool_results = _folder_tool_results + [
                        (_t_name, _result)
                        for (_t_name, _), _result in zip(_tool_tasks, _web_tool_results)
                    ]
                    if _screenshot_b64s:
                        messages.append({
                            "role": "user",
                            "content": "Voici la/les capture(s) d'écran demandée(s) :",
                            "images": _screenshot_b64s,
                        })
                    _image_tool_results_now = [
                        str(result)
                        for name, result in _all_tool_results
                        if name in ("generate_image", "edit_image")
                    ]
                    _image_tool_success_now = any(
                        "Image sauvegardée :" in result
                        for result in _image_tool_results_now
                    )
                    if _image_tool_success_now:
                        _saved_result = next(
                            (result for result in _image_tool_results_now if "Image sauvegardée :" in result),
                            "",
                        )
                        _saved_path = ""
                        if _saved_result:
                            _saved_path = _saved_result.split("Image sauvegardée :", 1)[1].splitlines()[0].strip()
                        _saved_name = os.path.basename(_saved_path) if _saved_path else ""
                        full_response = (
                            f"✅ Image générée et sauvegardée : {_saved_name}"
                            if _saved_name else
                            "✅ Image générée et sauvegardée."
                        )
                        _remove_loading()
                        _ai_add_bubble("assistant", full_response)
                        break
                    # Mémoriser le dernier résultat list_folder_contents pour
                    # l'auto-création si Gemma refuse d'appeler create_file.
                    for _t_name, _t_result in _all_tool_results:
                        if _t_name == "list_folder_contents":
                            _last_folder_listing = _t_result
                    if _text_parsed_tools:
                        # Gemma ne supporte pas role="tool" — injecter les résultats
                        # comme message user pour éviter HTTP 500 au deuxième appel.
                        _results_lines = [f"[{_tn}]: {_tr}" for _tn, _tr in _all_tool_results]
                        _injected_msg = "Résultats des outils :\n" + "\n\n".join(_results_lines)
                        # Ajouter la directive pour que Gemma continue (et appelle l'outil suivant
                        # si nécessaire) au lieu de répondre en texte avec les données.
                        _create_file_just_done_tp = any(
                            name == "create_file" for name, _ in _all_tool_results
                        )
                        _image_tool_results_tp = [
                            str(result)
                            for name, result in _all_tool_results
                            if name in ("generate_image", "edit_image")
                        ]
                        _image_tool_success_tp = any(
                            "Image sauvegardée :" in result
                            for result in _image_tool_results_tp
                        )
                        _image_tool_failed_tp = bool(_image_tool_results_tp) and not _image_tool_success_tp
                        if _create_file_just_done_tp:
                            _injected_msg += (
                                "\n\nLe fichier a été créé avec succès. "
                                "La tâche est terminée — réponds à l'utilisateur "
                                "pour confirmer ce qui a été fait, sans appeler d'autres outils."
                            )
                        elif _image_tool_success_tp:
                            _injected_msg += (
                                "\n\nL'image demandée a été générée/modifiée et sauvegardée avec succès. "
                                "La tâche est terminée — réponds à l'utilisateur "
                                "avec une confirmation très courte, sans répéter ni reformuler le prompt image, "
                                "et sans appeler d'autres outils."
                            )
                        elif _image_tool_failed_tp:
                            _injected_msg += (
                                "\n\nLa génération/modification d'image a échoué (aucun fichier image créé). "
                                "Réponds en expliquant clairement l'échec, sans prétendre qu'une image est disponible, "
                                "et propose une action de relance."
                            )
                        else:
                            _injected_msg += (
                                f"\n\nDemande à accomplir : \u00ab {_original_user_request} \u00bb\n"
                                "Les résultats des outils sont disponibles ci-dessus. "
                                "Si la tâche n'est pas encore terminée, appelle l'outil suivant "
                                "(ex. create_file si tu dois créer un fichier). "
                                "N'écris pas la réponse finale avant d'avoir utilisé tous les outils nécessaires."
                            )
                        messages.append({"role": "user", "content": _injected_msg})
                    else:
                        # Stratégie hybride pour Gemma 4 via Ollama :
                        # Round 0 → tool_responses dans le message assistant
                        #            (maintient Gemma en mode outil pour forcer le chaînage)
                        # Round 1+ → role:tool standard
                        #            (évite HTTP 500 causé par tool_responses multiples en historique)
                        if _tool_round == 0:
                            _last_assistant_idx = None
                            for _msg_idx in range(len(messages) - 1, -1, -1):
                                if messages[_msg_idx].get("role") == "assistant":
                                    _last_assistant_idx = _msg_idx
                                    break
                            if _last_assistant_idx is not None:
                                messages[_last_assistant_idx]["tool_responses"] = [
                                    {"name": _t_name, "response": _t_result}
                                    for _t_name, _t_result in _all_tool_results
                                ]
                        else:
                            # Rounds suivants : role:tool standard (évite HTTP 500)
                            for _t_name, _t_result in _all_tool_results:
                                messages.append({"role": "tool", "content": _t_result})
                        # Message user avec données réelles explicites + directive
                        # (garantit que Gemma voit les vraies données dans tous les cas)
                        _create_file_just_done = any(
                            name == "create_file" for name, _ in _all_tool_results
                        )
                        _image_tool_results = [
                            str(result)
                            for name, result in _all_tool_results
                            if name in ("generate_image", "edit_image")
                        ]
                        _image_tool_success = any(
                            "Image sauvegardée :" in result
                            for result in _image_tool_results
                        )
                        _image_tool_failed = bool(_image_tool_results) and not _image_tool_success
                        if _create_file_just_done:
                            messages.append({"role": "user", "content": (
                                "Le fichier a été créé avec succès. "
                                "La tâche est terminée — réponds à l'utilisateur "
                                "pour confirmer ce qui a été fait, sans appeler d'autres outils."
                            )})
                        elif _image_tool_success:
                            messages.append({"role": "user", "content": (
                                "L'image demandée a été générée/modifiée et sauvegardée avec succès. "
                                "La tâche est terminée — réponds à l'utilisateur "
                                "avec une confirmation très courte, sans répéter ni reformuler le prompt image, "
                                "et sans appeler d'autres outils."
                            )})
                        elif _image_tool_failed:
                            messages.append({"role": "user", "content": (
                                "La génération/modification d'image a échoué (aucun fichier image créé). "
                                "Réponds en expliquant clairement l'échec, sans prétendre qu'une image est disponible, "
                                "et propose une action de relance."
                            )})
                        else:
                            _results_lines = [f"[{_tn}]:\n{_tr}" for _tn, _tr in _all_tool_results]
                            messages.append({"role": "user", "content": (
                                "Résultats des outils :\n\n"
                                + "\n\n".join(_results_lines)
                                + f"\n\nDemande à accomplir : \u00ab {_original_user_request} \u00bb\n"
                                "Appelle maintenant l'outil suivant en utilisant EXACTEMENT "
                                "les données ci-dessus (ne pas les inventer ni les modifier). "
                                "Si la tâche est de créer un fichier, appelle create_file avec "
                                "le contenu recopié mot pour mot depuis les résultats."
                            )})
                    _remove_loading()
                else:
                    # 200 tours épuisés sans que Gemini ait terminé — cas exceptionnel
                    _ai_add_bubble("assistant", (
                        "*(Limite de 200 tours agentiques atteinte — "
                        "envoie un nouveau message pour continuer.)*"
                    ))
                    full_response = ""

                # Capturer le message original de l'utilisateur pour le log
                _last_user_text = next(
                    (m["content"] for m in reversed(ai_conversation) if m["role"] == "user"),
                    message_text
                )
                
                # Enrichir le log avec les pièces jointes si présentes
                _log_user_text = _last_user_text
                _attachments_info = []
                if "images_paths" in locals() and images_paths:
                    _attachments_info.append(f"🖼️ **Image(s) jointe(s) :** {', '.join(os.path.basename(p) for p in images_paths)}")
                if "files_to_inject" in locals() and files_to_inject:
                    _attachments_info.append(f"📁 **Document(s) joint(s) :** {', '.join(os.path.basename(p) for p in files_to_inject)}")
                
                if _attachments_info:
                    _attachments_str = "\n".join(_attachments_info)
                    if _log_user_text.strip():
                        _log_user_text = f"{_log_user_text.strip()}\n\n{_attachments_str}"
                    else:
                        _log_user_text = _attachments_str

                if full_response:
                    _entry = {"role": "assistant", "content": full_response}
                    if _thinking:
                        _entry["thinking"] = _thinking
                    if _turn_events:
                        _entry["events"] = _turn_events
                    ai_conversation.append(_entry)
                    _ai_save_history()
                    
                    # 📝 Log permanent dans le fichier Markdown
                    _log_exchange_to_md(_log_user_text, full_response, _thinking, _turn_events)
                else:
                    _fallback_response = "[Aucune réponse reçue]"
                    if _thinking:
                        # La réflexion a déjà été affichée — sauvegarder sans ajouter de bulle d'erreur par-dessus
                        _entry = {"role": "assistant", "content": _fallback_response, "thinking": _thinking}
                        if _turn_events:
                            _entry["events"] = _turn_events
                        ai_conversation.append(_entry)
                        _ai_save_history()
                    else:
                        if _turn_events:
                            ai_conversation.append({"role": "assistant", "content": _fallback_response, "events": _turn_events})
                            _ai_save_history()
                        _ai_add_bubble("assistant", _fallback_response)

                    # 📝 Log permanent même s'il n'y a pas eu de réponse
                    _log_exchange_to_md(_log_user_text, _fallback_response, _thinking, _turn_events)
                    
            except Exception as exc:
                _ai_add_bubble("assistant", f"[ERREUR] {exc}")
                full_response = ""
            finally:
                ai_streaming["value"] = False
                ai_progress_bar.visible = False
                ai_stop_button.icon_color = LIGHT_GREY
                ai_status_text.value = ""
                try:
                    page.update()
                except Exception:
                    pass
                if full_response and ai_tts_enabled["value"]:
                    threading.Thread(target=_speak_bubble, args=(full_response,), daemon=True).start()
                async def _refocus_after_response():
                    try:
                        await ai_input_field.focus()
                    except Exception:
                        pass
                page.run_task(_refocus_after_response)

        threading.Thread(target=_run, daemon=True).start()



    def _on_ai_submit():
        """Récupère le texte saisi, vide le champ et envoie le message à l'IA."""
        message_text = ai_input_field.value.strip()
        # Autoriser l'envoi sans texte si des images ou fichiers sont joints
        if not message_text and not ai_pending_images and not ai_pending_files:
            return
        # Ne pas effacer le champ si une réponse est déjà en cours
        if ai_streaming["value"]:
            return
        ai_input_field.value = ""
        ai_input_field.update()

        async def _refocus_ai():
            try:
                await ai_input_field.focus()
            except Exception:
                pass
        page.run_task(_refocus_ai)
        _send_ai_message(message_text)



    def _toggle_tts():
        """Active ou désactive la lecture vocale des réponses IA."""
        ai_tts_enabled["value"] = not ai_tts_enabled["value"]
        enabled = ai_tts_enabled["value"]
        ai_speaker_button.icon = ft.Icons.VOLUME_UP if enabled else ft.Icons.VOLUME_OFF
        ai_speaker_button.icon_color = BLUE if enabled else LIGHT_GREY
        ai_speaker_button.tooltip = "Désactiver la lecture vocale" if enabled else "Activer la lecture vocale"
        try:
            ai_speaker_button.update()
        except Exception:
            pass



    def _toggle_ai_image_size_mode():
        """Bascule entre envoi optimisé (1024px) et taille réelle pour les images IA."""
        ai_send_original_images["value"] = not ai_send_original_images["value"]
        use_original = ai_send_original_images["value"]
        ai_image_size_button.icon_color = GREEN if use_original else BLUE
        ai_image_mode_label.value = "REEL" if use_original else "1024"
        ai_image_mode_label.color = GREEN if use_original else BLUE
        ai_image_size_button.tooltip = (
            "Mode images IA en taille réelle (fichier original) — affecte uniquement les nouveaux fichiers joints"
            if use_original
            else "Mode images IA optimisé (1024px max) — affecte uniquement les nouveaux fichiers joints"
        )
        try:
            ai_image_size_button.update()
            ai_image_mode_label.update()
        except Exception:
            pass
        mode_label = "taille réelle" if use_original else "optimisé (1024px max)"
        log_to_terminal(f"[INFO] Envoi images IA : {mode_label}", LIGHT_GREY)



    def switch_to_terminal_mode():
        """Sauvegarde les notes/quitte l'IA et revient au terminal."""
        if note_mode["value"]:
            save_notes()
        note_mode["value"] = False
        ai_mode["value"]   = False
        terminal_output.visible   = True
        terminal_cmd_row.visible  = True
        update_overlay_visibility()
        terminal_output.update()
        terminal_cmd_row.update()

        async def _focus_term():
            try:
                await terminal_cmd_input.focus()
            except Exception:
                pass
        page.run_task(_focus_term)



    def on_terminal_command_submit(e):
        """Exécute la commande saisie dans le terminal intégré."""
        command_text = terminal_cmd_input.value.strip()
        if not command_text:
            return

        # ── Commandes internes (slash-commands) ───────────────────────
        if command_text.lower() == "/option":
            terminal_cmd_input.value = ""
            terminal_cmd_input.update()
            switch_to_options()
            return

        if not command_history or command_history[0] != command_text:
            command_history.insert(0, command_text)
        history_index["value"] = -1
        history_draft["value"] = ""
        terminal_cmd_input.value = ""
        terminal_cmd_input.update()

        async def _refocus():
            try:
                await terminal_cmd_input.focus()
            except Exception:
                pass

        page.run_task(_refocus)

        # Capturer le cwd maintenant, avant le lancement du thread,
        # pour éviter qu'une navigation simultanée ne le modifie.
        cwd = current_browse_folder["path"] or selected_folder.get("path") or app_directory
        log_to_terminal(f"> {command_text}", YELLOW)

        def _run():
            try:
                system = platform.system()
                if system == "Windows":
                    # PowerShell est plus capable que cmd.exe (ls, cat, pipeline…)
                    popen_kwargs = dict(
                        args=["powershell", "-NoProfile", "-NonInteractive", "-Command", command_text],
                        shell=False,
                    )
                else:
                    # zsh avec fallback bash si zsh n'est pas installé
                    shell_exe = "/bin/zsh" if os.path.exists("/bin/zsh") else "/bin/bash"
                    popen_kwargs = dict(
                        args=command_text,
                        shell=True,
                        executable=shell_exe,
                    )
                proc = subprocess.Popen(
                    **popen_kwargs,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    cwd=cwd,
                )
                killed_by_timeout = [False]

                def _kill_on_timeout():
                    if proc.poll() is None:
                        killed_by_timeout[0] = True
                        proc.kill()
                        log_to_terminal("[ERREUR] Commande interrompue (délai dépassé 30s)", RED)

                watchdog = threading.Timer(30.0, _kill_on_timeout)
                watchdog.daemon = True
                watchdog.start()
                try:
                    had_output = False
                    for line in iter(proc.stdout.readline, ""):
                        stripped = line.rstrip()
                        if stripped:
                            log_to_terminal(stripped)
                            had_output = True
                    proc.wait()
                    if not killed_by_timeout[0]:
                        if proc.returncode != 0:
                            log_to_terminal(f"[code retour {proc.returncode}]", RED)
                        elif not had_output:
                            log_to_terminal("[aucun résultat]", LIGHT_GREY)
                finally:
                    watchdog.cancel()
            except FileNotFoundError:
                log_to_terminal(f"[ERREUR] Dossier introuvable : {cwd}", RED)
            except Exception as error:
                log_to_terminal(f"[ERREUR] {error}", RED)

        threading.Thread(target=_run, daemon=True).start()



    # ================================================================ #
    #                   NAVIGATION & FICHIERS                          #
    # ================================================================ #
    def on_folder_path_submit(e):
        """Charge un dossier collé/saisi manuellement dans le champ."""
        raw = (folder_path.value or "").strip().strip('"').strip("'")
        raw = raw.replace("\\", os.sep).replace("/", os.sep)
        if raw and os.path.isdir(raw):
            folder_path.error_text = None
            navigate_to_folder(raw)
        else:
            folder_path.error_text = "Dossier introuvable"
            folder_path.value = _short_path(selected_folder.get("path", "") or "")
            folder_path.update()



    def on_folder_path_blur(e):
        """Restaure le chemin courant si le champ est laissé invalide."""
        folder_path.error_text = None
        folder_path.value = _short_path(selected_folder.get("path", "") or "")
        folder_path.update()



    def _rebuild_recent_folders_menu():
        """Reconstruit les items du bouton dossiers récents."""
        recent = _load_recent()
        if not recent:
            recent_folders_btn.items = [
                ft.PopupMenuItem(content=ft.Text("Aucun dossier récent"))
            ]
        else:
            recent_folders_btn.items = [
                ft.PopupMenuItem(
                    content=ft.Row([
                        ft.Icon(ft.Icons.FOLDER, size=16),
                        ft.Text(os.path.basename(recent_path) or recent_path),
                    ], spacing=8, tight=True),
                    on_click=lambda e, folder=recent_path: navigate_to_folder(folder),
                )
                for recent_path in recent
            ]
        try:
            recent_folders_btn.update()
        except Exception:
            pass



    def open_in_file_explorer(folder_path):
        """Ouvre le dossier dans l'explorateur de fichiers natif"""
        if not folder_path or not os.path.isdir(folder_path):
            log_to_terminal("Aucun dossier sélectionné", RED)
            return
        
        try:
            if platform.system() == "Windows":
                subprocess.Popen(f'explorer "{folder_path}"')
            elif platform.system() == "Darwin":  # macOS
                subprocess.Popen(["open", folder_path])
            else:  # Linux
                subprocess.Popen(["xdg-open", folder_path])
            log_to_terminal(f"[OK] Ouverture du dossier: {os.path.basename(folder_path)}", GREEN)
        except Exception as e:
            log_to_terminal(f"[ERREUR] Erreur lors de l'ouverture de l'explorateur: {e}", RED)



    def open_file_with_default_app(file_path):
        """Ouvre un fichier avec l'application par défaut du système en premier plan"""
        if not file_path or not os.path.isfile(file_path):
            return
        
        try:
            if platform.system() == "Windows":
                # Utilise 'start' avec '' pour ouvrir en premier plan
                subprocess.Popen(f'start "" "{file_path}"', shell=True)
            elif platform.system() == "Darwin":  # macOS
                subprocess.Popen(["open", file_path])
            else:  # Linux
                subprocess.Popen(["xdg-open", file_path])
        except Exception as e:
            log_to_terminal(f"[ERREUR] Erreur lors de l'ouverture du fichier: {e}", RED)



    def _print_files_with_default_app(files: list):
        """Utilise l'application par défaut de l'OS pour imprimer ou ouvrir les images."""
        files_to_print = list(files) if files else []

        # Si aucun fichier n'est sélectionné, on récupère tous les fichiers du dossier en cours
        if not files_to_print:
            active_dir = current_browse_folder["path"] or selected_folder["path"]
            if active_dir and os.path.isdir(active_dir):
                try:
                    files_to_print = [
                        os.path.join(active_dir, f)
                        for f in os.listdir(active_dir)
                        if os.path.isfile(os.path.join(active_dir, f))
                    ]
                except Exception as dir_err:
                    log_to_terminal(f"[ERREUR] Impossible de lister le dossier : {dir_err}", RED)
                    return False

        # On filtre STRICTEMENT sur les extensions d'images configurées dans CONSTANTS.py
        # Cela élimine d'office les PDF, dossiers, scripts, zip, etc. !
        image_files = [
            file_path for file_path in files_to_print
            if os.path.isfile(file_path)
            and os.path.splitext(file_path)[1].lower() in CONSTANTS.IMAGE_EXTS
        ]

        if not image_files:
            log_to_terminal("[ATTENTION] Aucune image à imprimer", ORANGE)
            return False

        if not _strip_state["active"]:
            _toggle_strip()

        printed_count = 0
        try:
            if platform.system() == "Windows":
                # Imports dynamiques pour le thread et COM
                try:
                    import win32com.client
                    import pythoncom
                except ImportError:
                    log_to_terminal(
                        "[ERREUR] Le module 'pywin32' est requis pour l'impression groupée sous Windows.\n"
                        "Lance : pip install pywin32", 
                        RED
                    )
                    return False

                # On définit la fonction qui va s'exécuter en tâche de fond
                def _run_wia_wizard(paths):
                    # Initialisation de COM obligatoire pour ce nouveau thread
                    pythoncom.CoInitialize()
                    try:
                        wia_dialog = win32com.client.Dispatch("WIA.CommonDialog")
                        wia_vector = win32com.client.Dispatch("WIA.Vector")
                        for file_path in paths:
                            abs_path = os.path.abspath(file_path)
                            wia_vector.Add(abs_path)
                        
                        # Cet appel va bloquer son propre thread, mais pas l'interface de l'app !
                        wia_dialog.ShowPhotoPrintingWizard(wia_vector)
                    except Exception as thread_err:
                        log_to_terminal(f"[ERREUR] Assistant d'impression : {thread_err}", RED)
                    finally:
                        # Libération propre des ressources COM pour ce thread
                        pythoncom.CoUninitialize()

                # On lance le thread en mode "daemon" pour qu'il ne bloque pas la fermeture de l'app
                print_thread = threading.Thread(
                    target=_run_wia_wizard, 
                    args=(image_files,), 
                    daemon=True
                )
                print_thread.start()
                printed_count = len(image_files)

            elif platform.system() == "Darwin":
                # Sous macOS, on ouvre toutes les images d'un coup dans la même instance (ex: Aperçu)
                import subprocess
                subprocess.call(["open"] + image_files)
                printed_count = len(image_files)

            else:
                # Pour Linux ou autre, on boucle sur ton helper existant
                for file_path in image_files:
                    open_file_with_default_app(file_path)
                printed_count = len(image_files)

        except Exception as err:
            log_to_terminal(f"[ERREUR] Impression : {err}", RED)

        if printed_count:
            if platform.system() == "Windows":
                log_to_terminal(f"[OK] Assistant d'impression lancé en arrière-plan pour {printed_count} image(s)", GREEN)
            else:
                log_to_terminal(f"[OK] Ouverture dans l'application par défaut pour {printed_count} image(s)", GREEN)
            return True
        return False



    # ── Ouvrir avec (menu clic-droit) ─────────────────────────────────
    def _load_open_with_programs() -> list:
        """Charge la liste des programmes depuis open_with.json."""
        try:
            with open(open_with_config_file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return [program for program in data if isinstance(program, dict) and "label" in program and "exe" in program]
        except Exception:
            return []



    def _resolve_exe_path(exe: str) -> str:
        """
        Si le chemin exe n'existe pas et contient 'WindowsApps' avec un numéro de version,
        cherche automatiquement la nouvelle version installée via glob.
        Retourne le chemin résolu (ou l'original si rien n'est trouvé).
        """
        if not exe or os.path.isfile(exe):
            return exe
        if "WindowsApps" not in exe:
            return exe
        import glob as _glob
        # Remplace x.x.x.x par * pour trouver la nouvelle version
        pattern = re.sub(r'_\d+\.\d+\.\d+\.\d+_', '_*_', exe)
        matches = _glob.glob(pattern)
        return matches[0] if matches else exe

    def _open_files_with(prog: dict, files: list):
        """Ouvre une liste de fichiers avec le programme spécifié."""
        exe = prog.get("exe", "")
        if not exe:
            return
        try:
            resolved = _resolve_exe_path(exe)
            if resolved != exe:
                # Mise à jour silencieuse du chemin dans open_with.json
                prog["exe"] = resolved
                all_progs = _load_open_with_programs()
                for p in all_progs:
                    if p.get("label") == prog.get("label") and p.get("exe") == exe:
                        p["exe"] = resolved
                        break
                _save_open_with_programs(all_progs)
                log_to_terminal(f"[INFO] Chemin mis à jour automatiquement pour {prog['label']}", ORANGE)
            if platform.system() == "Darwin":
                subprocess.Popen(["open", "-a", resolved] + files)
            else:
                subprocess.Popen([resolved] + files)
            display_names = ", ".join(os.path.basename(file_path) for file_path in files[:3])
            overflow_suffix = f" (+{len(files) - 3})" if len(files) > 3 else ""
            log_to_terminal(f"[OK] Ouvert avec {prog['label']}: {display_names}{overflow_suffix}", GREEN)
            return True
        except Exception as err:
            log_to_terminal(f"[ERREUR] {prog['label']}: {err}", RED)
            return False



    def _save_open_with_programs(programs: list):
        """Sauvegarde la liste des programmes dans open_with.json."""
        try:
            with open(open_with_config_file_path, "w", encoding="utf-8") as f:
                json.dump(programs, f, ensure_ascii=False, indent=2)
        except Exception as err:
            log_to_terminal(f"[ERREUR] Sauvegarde open_with.json : {err}", RED)



    _ROTATABLE_EXTS = CONSTANTS.ROTATABLE_EXTS



    def _zip_selection(items: list, zip_name: str):
        """Zippe la liste de fichiers et/ou dossiers dans le dossier courant sous zip_name."""
        if not items:
            return
        folder = os.path.dirname(items[0])
        base = zip_name if zip_name else "selection"
        if not base.lower().endswith(".zip"):
            base += ".zip"
        candidate = os.path.join(folder, base)
        archive_stem, archive_ext = os.path.splitext(candidate)
        name_counter = 1
        while os.path.exists(candidate):
            candidate = f"{archive_stem}_{name_counter}{archive_ext}"
            name_counter += 1
        try:
            with zipfile.ZipFile(candidate, "w", zipfile.ZIP_DEFLATED) as zip_archive:
                for item_path in items:
                    if os.path.isdir(item_path):
                        dir_name = os.path.basename(item_path)
                        for root, _dirs, files in os.walk(item_path):
                            for file_name in files:
                                full_path = os.path.join(root, file_name)
                                arcname = os.path.join(dir_name, os.path.relpath(full_path, item_path))
                                zip_archive.write(full_path, arcname=arcname)
                    else:
                        zip_archive.write(item_path, arcname=os.path.basename(item_path))
            log_to_terminal(f"[OK] Archive créée : {os.path.basename(candidate)}", YELLOW)
            page.pubsub.send_all_on_topic("deselect", None)
            refresh_preview()
        except Exception as ex:
            log_to_terminal(f"[ERREUR] Zip : {ex}", RED)



    def _prompt_and_zip_selection(e):
        """Affiche le dialog de nom puis zippe les éléments sélectionnés (fichiers et dossiers)."""
        items = list(selected_files)
        if not items:
            log_to_terminal("[ATTENTION] Aucun élément sélectionné à zipper", ORANGE)
            return
        default_name = ""
        if len(items) == 1 and os.path.isdir(items[0]):
            default_name = os.path.basename(items[0])
        zip_name_input = ft.TextField(
            label="Nom de l'archive",
            hint_text="Ex: selection",
            value=default_name,
            autofocus=True,
            width=320,
            bgcolor=DARK,
            border_color=GREY,
        )



        def _on_confirm_zip(ev):
            name = (zip_name_input.value or "").strip() or default_name or "selection"
            zip_dlg.open = False
            page.update()
            log_to_terminal(f"[ZIP] Création de {name}.zip en cours…", YELLOW)
            threading.Thread(target=_zip_selection, args=(items, name), daemon=True).start()

        def _on_cancel_zip(ev):
            zip_dlg.open = False
            page.update()

        zip_name_input.on_submit = _on_confirm_zip
        zip_dlg = ft.AlertDialog(
            modal=True,
            title=ft.Text("Nom de l'archive ZIP"),
            content=zip_name_input,
            actions=[
                ft.TextButton("Annuler", on_click=_on_cancel_zip),
                ft.TextButton("OK", on_click=_on_confirm_zip),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        page.overlay.append(zip_dlg)
        zip_dlg.open = True
        page.update()



    def _rotate_files(files, direction, clear_sel=False):
        """Pivote les images de 90° à gauche ou à droite en utilisant Pillow."""
        if _PILImage is None:
            log_to_terminal("[ERREUR] Pillow non installé — rotation impossible", RED)
            return
        rotated = 0
        timestamp = str(int(time.time() * 1000))
        for file_path in files:
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in _ROTATABLE_EXTS:
                continue
            try:
                with _PILImage.open(file_path) as img:
                    img_converted = img.copy()
                    if direction == "left":
                        result = img_converted.rotate(90, expand=True)
                    else:
                        result = img_converted.rotate(-90, expand=True)
                    save_kwargs = {}
                    image_format = "JPEG" if ext in (".jpg", ".jpeg") else "PNG"
                    if image_format == "JPEG":
                        save_kwargs["quality"] = 95
                        save_kwargs["subsampling"] = 0
                    # Sauvegarder le fichier original
                    result.save(file_path, **save_kwargs)
                    # Mettre à jour le cache base64 pour bypasser l'ancienne miniature
                    normalized_path = os.path.normpath(file_path)
                    _thumb_cache.pop(normalized_path, None)
                    new_b64 = thumb_cache.get_or_generate(file_path)
                    if new_b64:
                        _image_cache_busters[normalized_path] = new_b64
                        _thumb_cache[normalized_path] = new_b64
                    try:
                        _image_last_mtime[normalized_path] = os.stat(file_path).st_mtime
                    except OSError:
                        pass
                rotated += 1
            except Exception as ex:
                log_to_terminal(f"[ERREUR] Rotation {os.path.basename(file_path)}: {ex}", RED)
        if rotated:
            label = "gauche" if direction == "left" else "droite"
            log_to_terminal(f"[OK] {rotated} image(s) pivotée(s) vers {label}", YELLOW)
            if clear_sel:
                selected_files.clear()
                _update_select_toggle_button()
            refresh_preview(reset_page=False, force_reload=True)



    def _show_file_context_menu(files: list):
        """Menu contextuel clic-droit : rotation + liste Ouvrir avec intégrée."""
        image_files = [f for f in files if os.path.splitext(f)[1].lower() in _ROTATABLE_EXTS]
        has_images = bool(image_files)
        document_files = [
            f for f in files
            if os.path.splitext(f)[1].lower() in _AI_DOCUMENT_EXTS
        ]
        has_documents = bool(document_files)

        header_label = (
            os.path.basename(files[0]) if len(files) == 1
            else f"{len(files)} fichier(s) sélectionné(s)"
        )

        # ── Formulaire d'ajout de programme ──────────────────────────────
        add_label_field = ft.TextField(
            hint_text="Nom affiché (ex : Affinity)",
            border_color=BLUE, text_size=13, height=40,
            content_padding=ft.Padding(8, 4, 8, 4), expand=True,
        )
        add_exe_field = ft.TextField(
            hint_text="Chemin exe (ex : C:\\...\\Affinity.exe)",
            border_color=BLUE, text_size=13, height=40,
            content_padding=ft.Padding(8, 4, 8, 4), expand=True,
        )

        async def _browse_exe(e):
            result = await ft.FilePicker().pick_files(
                dialog_title="Choisir l'exécutable",
                allow_multiple=False,
            )
            if result:
                add_exe_field.value = result[0].path
                add_exe_field.update()

        browse_exe_btn = ft.IconButton(
            icon=ft.Icons.FOLDER_OPEN, icon_color=BLUE, icon_size=20,
            tooltip="Parcourir…", on_click=_browse_exe,
            style=ft.ButtonStyle(padding=ft.Padding.all(4)),
        )

        add_form = ft.Container(
            content=ft.Column([
                ft.Divider(height=8, color=GREY),
                ft.Text("Ajouter un programme", size=12, color=LIGHT_GREY),
                ft.Row([add_label_field], tight=True),
                ft.Row([add_exe_field, browse_exe_btn], tight=True),
            ], spacing=6, tight=True),
            visible=False,
        )

        programs_list_view = ft.ReorderableListView(padding=0, show_default_drag_handles=False)

        dlg = ft.AlertDialog(
            title=ft.Row([
                ft.Text(header_label, size=13, color=LIGHT_GREY,
                        max_lines=1, overflow=ft.TextOverflow.ELLIPSIS, expand=True),
                ft.Container(
                    content=ft.Icon(ft.Icons.ADD, size=13, color=DARK),
                    bgcolor=BLUE, border_radius=10,
                    padding=ft.Padding(3, 1, 3, 1),
                    tooltip="Ajouter un programme",
                    ink=True,
                    on_click=lambda e: _toggle_add_form(),
                ),
            ], spacing=8, tight=True, vertical_alignment=ft.CrossAxisAlignment.CENTER),
            content_padding=ft.Padding(12, 0, 12, 8),
        )



        def _close(e=None):
            dlg.open = False
            page.update()



        def _do_rotate(direction):
            _close()
            threading.Thread(target=_rotate_files, args=(image_files, direction, True), daemon=True).start()



        def _rebuild_items():
            programs = _load_open_with_programs()
            programs_list_view.controls.clear()
            if not programs:
                programs_list_view.controls.append(ft.Container(
                    key="empty",
                    content=ft.Text(
                        "Aucun programme configuré — cliquez + pour en ajouter.",
                        size=12, color=LIGHT_GREY,
                    ),
                    padding=ft.Padding(0, 6, 0, 6),
                ))
            else:
                for program_index, prog in enumerate(programs):
                    def _create_open_handler(program):
                        def _open_with_clicked(e):
                            if platform.system() == "Windows":
                                try:
                                    import ctypes
                                    hwnd = ctypes.windll.user32.GetForegroundWindow()
                                    if hwnd:
                                        ctypes.windll.user32.ShowWindow(hwnd, 6)  # SW_MINIMIZE
                                except Exception:
                                    pass
                            _close()
                            if _open_files_with(program, files):
                                clear_selection(None)
                        return _open_with_clicked
                    
                    def _create_delete_handler(program):
                        def _delete_program_clicked(e):
                            current_programs = _load_open_with_programs()
                            current_programs = [entry for entry in current_programs if entry != program]
                            _save_open_with_programs(current_programs)
                            _rebuild_items()
                            programs_list_view.update()
                        return _delete_program_clicked
                    programs_list_view.controls.append(ft.ListTile(
                        key=str(program_index),
                        leading=ft.ReorderableDragHandle(
                            content=ft.Icon(ft.Icons.DRAG_HANDLE, color=LIGHT_GREY, size=18),
                        ),
                        title=ft.Text(prog["label"], size=13, color=WHITE),
                        trailing=ft.IconButton(
                            icon=ft.Icons.CLOSE, icon_size=16,
                            icon_color=LIGHT_GREY, tooltip="Supprimer",
                            on_click=_create_delete_handler(prog),
                            style=ft.ButtonStyle(padding=ft.Padding.all(4)),
                        ),
                        on_click=_create_open_handler(prog),
                        hover_color=GREY, dense=True,
                        content_padding=ft.Padding(0, 0, 0, 0),
                    ))



        def _on_reorder(e: ft.OnReorderEvent):
            programs_list_view.controls.insert(e.new_index, programs_list_view.controls.pop(e.old_index))
            programs = _load_open_with_programs()
            programs.insert(e.new_index, programs.pop(e.old_index))
            _save_open_with_programs(programs)
            programs_list_view.update()

        programs_list_view.on_reorder = _on_reorder



        def _toggle_add_form():
            add_form.visible = not add_form.visible
            add_program_button.visible = add_form.visible
            if add_form.visible:
                add_label_field.value = ""
                add_exe_field.value = ""
            page.update()



        def _confirm_add(e):
            label = (add_label_field.value or "").strip()
            exe   = (add_exe_field.value or "").strip()
            if not label or not exe:
                add_label_field.error_text = "Requis" if not label else None
                add_exe_field.error_text   = "Requis" if not exe   else None
                page.update()
                return
            add_label_field.error_text = None
            add_exe_field.error_text   = None
            progs = _load_open_with_programs()
            progs.append({"label": label, "exe": exe})
            _save_open_with_programs(progs)
            add_form.visible = False
            add_program_button.visible = False
            _rebuild_items()
            page.update()

        # ── Assemblage du contenu ─────────────────────────────────────────
        content_rows = []
        if has_images:
            def _send_images_to_ai(e=None):
                _close()
                for image_path in image_files:
                    _ai_attach_image(image_path)
                switch_to_ai_mode()
                clear_selection(None)

            def _print_images(e=None):
                _close()
                if _print_files_with_default_app(image_files):
                    clear_selection(None)

            def _show_exif_data(e=None):
                _close()
                exif_path = image_files[0]
                rows = []
                try:
                    if _PILImage is None:
                        raise ImportError("PIL non disponible")
                    from PIL.ExifTags import TAGS
                    with _PILImage.open(exif_path) as img:
                        width, height = img.size
                        raw = img.getexif()
                    rows.append(
                        ft.Text(f"Résolution : {width} × {height} px", size=12, color=BLUE, selectable=True)
                    )
                    if raw:
                        for tag_id, value in raw.items():
                            tag_name = TAGS.get(tag_id, f"Tag {tag_id}")
                            if isinstance(value, bytes):
                                continue
                            rows.append(
                                ft.Text(f"{tag_name}: {value}", size=12, color=WHITE, selectable=True)
                            )
                    else:
                        rows.append(ft.Text("Aucune donnée EXIF.", size=12, color=LIGHT_GREY))
                except Exception as ex:
                    rows.append(ft.Text(f"Erreur : {ex}", size=12, color=RED))

                exif_dlg = ft.AlertDialog(
                    title=ft.Text(os.path.basename(exif_path), size=13, color=LIGHT_GREY),
                    content=ft.Column(
                        rows, spacing=2,
                        scroll=ft.ScrollMode.AUTO,
                        width=400, height=400,
                    ),
                )

                def _close_exif(e=None):
                    exif_dlg.open = False
                    page.update()

                exif_dlg.actions = [ft.TextButton("Fermer", on_click=_close_exif)]
                exif_dlg.actions_alignment = ft.MainAxisAlignment.END
                page.overlay.append(exif_dlg)
                exif_dlg.open = True
                page.update()

            icon_btns = [
                ft.IconButton(
                    icon=ft.Icons.SMART_TOY, icon_color=BLUE, icon_size=22,
                    tooltip=f"Envoyer à l'IA ({len(image_files)} image{'s' if len(image_files) > 1 else ''})",
                    on_click=_send_images_to_ai,
                ),
                ft.IconButton(
                    icon=ft.Icons.PRINT, icon_color=ORANGE, icon_size=22,
                    tooltip=f"Imprimer ({len(image_files)} image{'s' if len(image_files) > 1 else ''})",
                    on_click=_print_images,
                ),
                ft.IconButton(
                    icon=ft.Icons.ROTATE_LEFT, icon_color=HOVER_YELLOW, icon_size=22,
                    tooltip="Rotation gauche (−90°)",
                    on_click=lambda e: _do_rotate("left"),
                ),
                ft.IconButton(
                    icon=ft.Icons.ROTATE_RIGHT, icon_color=YELLOW, icon_size=22,
                    tooltip="Rotation droite (+90°)",
                    on_click=lambda e: _do_rotate("right"),
                ),
            ]
            if len(image_files) == 1:
                icon_btns.append(ft.IconButton(
                    icon=ft.Icons.INFO_OUTLINE, icon_color=LIGHT_GREY, icon_size=22,
                    tooltip="Voir les EXIF",
                    on_click=_show_exif_data,
                ))
            content_rows.append(ft.Row(icon_btns, spacing=0, alignment=ft.MainAxisAlignment.CENTER))
            content_rows.append(ft.Divider(height=8, color=GREY))
        if has_documents:
            def _send_docs_to_ai(e=None):
                _close()
                for doc_path in document_files:
                    _ai_attach_document_file(doc_path)
                switch_to_ai_mode()
                clear_selection(None)

            doc_label = (
                f"{len(document_files)} fichier{'s' if len(document_files) > 1 else ''}"
            )
            content_rows.append(ft.Row([
                ft.IconButton(
                    icon=ft.Icons.SMART_TOY, icon_color=VIOLET, icon_size=22,
                    tooltip=f"Envoyer à l'IA ({doc_label})",
                    on_click=_send_docs_to_ai,
                ),
            ], spacing=0, tight=True))
            content_rows.append(ft.Divider(height=8, color=GREY))
        content_rows.append(programs_list_view)
        content_rows.append(add_form)

        dlg.content = ft.Column(content_rows, spacing=0, tight=True, width=340)

        add_program_button = ft.TextButton("Ajouter", on_click=_confirm_add, visible=False)
        dlg.actions = [add_program_button, ft.TextButton("Fermer", on_click=_close)]
        dlg.actions_alignment = ft.MainAxisAlignment.END

        _rebuild_items()
        page.overlay.append(dlg)
        dlg.open = True
        page.update()



    def navigate_to_folder(new_path):
        """Navigue vers un dossier dans la preview"""
        if not new_path:
            return
        if note_mode["value"]:
            save_notes()
        new_path = _resolve_favorite_path(new_path)
        current_browse_folder["path"] = new_path
        selected_folder["path"] = new_path
        folder_path.value = _short_path(new_path)
        folder_path.update()
        selected_files.clear()
        selection_count_text.value = ""
        search_query["value"] = ""
        search_field.value = ""
        if show_only_selection["value"]:
            show_only_selection["value"] = False
            _update_filter_sel_btn()
        preview_page["value"] = 0
        _add_to_recent(new_path)
        _rebuild_recent_folders_menu()
        refresh_preview()



    def go_to_parent_folder(e):
        """Remonte au dossier parent"""
        if current_browse_folder["path"]:
            parent = os.path.dirname(current_browse_folder["path"])
            if parent and parent != current_browse_folder["path"]:
                navigate_to_folder(parent)



    def extract_zip(file_path):
        """Décompresse un .zip avec détection de dossier racine unique"""
        dest_dir = os.path.dirname(file_path)
        zip_name = os.path.splitext(os.path.basename(file_path))[0]
        
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                names = zf.namelist()
                top_levels = {n.split('/')[0] for n in names if n}
                # Si tout le contenu est sous un seul dossier racine, extraire directement
                if len(top_levels) == 1 and any('/' in n for n in names):
                    extract_to = dest_dir
                else:
                    extract_to = os.path.join(dest_dir, zip_name)
                    os.makedirs(extract_to, exist_ok=True)

                zf.extractall(extract_to)
            log_to_terminal(f"[OK] Décompressé: {os.path.basename(file_path)}", GREEN)
            if CONSTANTS.DELETE_ZIP_AFTER_EXTRACT:
                try:
                    os.remove(file_path)
                    log_to_terminal(f"[OK] ZIP supprimé: {os.path.basename(file_path)}", GREEN)
                except Exception as _ze:
                    log_to_terminal(f"[ERREUR] Suppression ZIP: {_ze}", RED)
                refresh_preview()
            else:
                def _confirm_del_zip(ev):
                    dlg_del_zip.open = False
                    page.update()
                    try:
                        os.remove(file_path)
                        log_to_terminal(f"[OK] ZIP supprimé: {os.path.basename(file_path)}", GREEN)
                        refresh_preview()
                    except Exception as _ze:
                        log_to_terminal(f"[ERREUR] Suppression ZIP: {_ze}", RED)
                def _cancel_del_zip(ev):
                    dlg_del_zip.open = False
                    page.update()
                dlg_del_zip = ft.AlertDialog(
                    title=ft.Text("Supprimer le fichier ZIP ?"),
                    content=ft.Text(f"Voulez-vous supprimer '{os.path.basename(file_path)}' ?"),
                    actions=[
                        ft.TextButton("Conserver", on_click=_cancel_del_zip),
                        ft.TextButton("Supprimer", on_click=_confirm_del_zip,
                                      style=ft.ButtonStyle(color=ft.Colors.RED)),
                    ],
                )
                page.overlay.append(dlg_del_zip)
                dlg_del_zip.open = True
                page.update()
                refresh_preview()
        except Exception as err:
            log_to_terminal(f"[ERREUR] Décompression: {err}", RED)



    def open_image_viewer(start_path):
        """Affiche un lecteur d'image avec PageView swipeable (support écran tactile)."""
        _blank_gif = "data:image/gif;base64,R0lGODlhAQABAAD/ACwAAAAAAQABAAACADs="

        def _resolve_viewer_image_src(path: str) -> str:
            """Retourne une source sûre pour la visionneuse plein écran.

            ``_image_cache_busters`` contient des bytes de miniature (thumb_cache) ou
            une str data-URL ou un chemin temp. Dans tous les cas non-str, on retourne
            le chemin original pour afficher l'image complète en plein écran.
            """
            normalized_path = os.path.normpath(path)
            cached_value = _image_cache_busters.get(normalized_path)
            # bytes = miniature PIL générée par thumb_cache → utiliser le fichier original
            if isinstance(cached_value, (bytes, bytearray)):
                return path
            # data-URL str = miniature encodée → utiliser le fichier original
            if isinstance(cached_value, str) and cached_value.startswith("data:image"):
                return path
            # Chemin temp (cache-buster) ou None
            return cached_value if cached_value else path

        entries = all_entries_data["list"]
        if show_only_selection["value"]:
            image_paths = [
                entry_path
                for (_name, entry_path, is_directory, is_image_file, _ext) in entries
                if is_image_file and not is_directory and entry_path in selected_files
            ]
        else:
            image_paths = [entry_path for (_, entry_path, is_directory, is_image_file, _ext) in entries if is_image_file and not is_directory]
        if not image_paths:
            image_paths = [start_path]
        try:
            initial_index = image_paths.index(start_path)
        except ValueError:
            initial_index = 0
            image_paths = [start_path]

        state = {"index": initial_index}
        previous_keyboard_handler = page.on_keyboard_event

        def _current_path() -> str:
            return image_paths[state["index"]] if image_paths else ""



        # ── Helpers ──────────────────────────────────────────────────────
        def _get_resolution(path):
            if _PILImage:
                try:
                    with _PILImage.open(path) as opened_image:
                        return f"{opened_image.width} × {opened_image.height}"
                except Exception:
                    pass
            return ""



        # ── Contrôles barre titre ─────────────────────────────────────────
        filename_text = ft.Text(
            os.path.basename(_current_path()),
            size=13,
            color=ft.Colors.WHITE,
            weight=ft.FontWeight.W_500,
            max_lines=1,
            overflow=ft.TextOverflow.ELLIPSIS,
        )
        counter_text = ft.Text(
            f"{state['index'] + 1} / {len(image_paths)}",
            size=12,
            color=ft.Colors.WHITE70,
        )
        resolution_text = ft.Text(
            "",
            size=12,
            color=ft.Colors.WHITE54,
        )
        viewer_checkbox = ft.Checkbox(
            value=_current_path() in selected_files,
            on_change=lambda e: on_checkbox_change(e, _current_path()),
        )



        # ── Chargement lazy des images ────────────────────────────────────
        page_image_controls: dict = {}
        pages_loaded: set = set()

        def _build_page_containers():
            containers = []
            win_w = page.window.width or 1280
            win_h = page.window.height or 800
            for idx in range(len(image_paths)):
                img_ctrl = ft.Image(
                    src=_blank_gif,
                    fit=ft.BoxFit.CONTAIN,
                    expand=True,
                    gapless_playback=True,
                    error_content=ft.Container(
                        content=ft.Icon(ft.Icons.BROKEN_IMAGE, color=ft.Colors.WHITE54, size=64),
                        alignment=ft.Alignment(0, 0),
                    ),
                )
                page_image_controls[idx] = img_ctrl
                # InteractiveViewer à l'intérieur de chaque page :
                # - Pinch-to-zoom / molette souris pour zoomer
                # - Swipe horizontal disponible quand zoom = 1× (PageView reprend la main)
                # - Pan disponible quand zoom > 1×
                viewer = ft.InteractiveViewer(
                    key=f"iv_{idx}",
                    content=img_ctrl,
                    min_scale=0.5,
                    max_scale=10.0,
                    pan_enabled=True,
                    scale_enabled=True,
                    constrained=True,
                    width=win_w,
                    height=win_h,
                    clip_behavior=ft.ClipBehavior.HARD_EDGE,
                )
                containers.append(
                    ft.Container(
                        content=viewer,
                        expand=True,
                        alignment=ft.Alignment(0, 0),
                        bgcolor="#1a1a1a",
                    )
                )
            return containers

        def _load_image_for_index(load_index: int) -> None:
            if load_index < 0 or load_index >= len(image_paths):
                return
            if load_index in pages_loaded:
                return
            path = image_paths[load_index]
            normalized = os.path.normpath(path)
            # Si le fichier a été modifié (signalé par _image_cache_busters), on lit
            # les bytes bruts du fichier pour contourner le cache URL de Flutter.
            if normalized in _image_cache_busters:
                try:
                    with open(path, "rb") as _f:
                        src: object = _f.read()
                except Exception:
                    src = _resolve_viewer_image_src(path)
            else:
                src = _resolve_viewer_image_src(path)
            if load_index in page_image_controls:
                page_image_controls[load_index].src = src
            pages_loaded.add(load_index)

            async def _apply():
                try:
                    page.update()
                except Exception:
                    pass

            page.run_task(_apply)

        def _load_pages_around(center: int) -> None:
            for offset in (0, 1, -1, 2, -2):
                target = center + offset
                if 0 <= target < len(image_paths):
                    threading.Thread(
                        target=_load_image_for_index,
                        args=(target,),
                        daemon=True,
                    ).start()



        # ── Mise à jour de la barre titre ─────────────────────────────────
        def _update_overlay_bar(new_index: int) -> None:
            state["index"] = new_index
            path = image_paths[new_index] if image_paths else ""
            filename_text.value = os.path.basename(path)
            counter_text.value = f"{new_index + 1} / {len(image_paths)}"
            viewer_checkbox.value = path in selected_files
            resolution_text.value = ""

            def _load_res():
                resolution_text.value = _get_resolution(path)
                try:
                    page.update()
                except Exception:
                    pass

            threading.Thread(target=_load_res, daemon=True).start()
            page.update()

        def on_page_change(e) -> None:
            new_index = int(e.data)
            _update_overlay_bar(new_index)
            _load_pages_around(new_index)



        # ── Viewer principal ──────────────────────────────────────────────
        _HAS_PAGE_VIEW = hasattr(ft, "PageView") and platform.system() != "Linux"
        if _HAS_PAGE_VIEW:
            images_page_view = ft.PageView(
                controls=_build_page_containers(),
                expand=True,
                horizontal=True,
                selected_index=initial_index,
                on_change=on_page_change,
            )
        else:
            # Fallback : une seule image visible à la fois (navigation par boutons/clavier)
            _fb_win_w = page.window.width or 1280
            _fb_win_h = page.window.height or 800
            _fb_img_ctrl = ft.Image(
                src=_blank_gif,
                fit=ft.BoxFit.CONTAIN,
                expand=True,
                gapless_playback=True,
                error_content=ft.Container(
                    content=ft.Icon(ft.Icons.BROKEN_IMAGE, color=ft.Colors.WHITE54, size=64),
                    alignment=ft.Alignment(0, 0),
                ),
            )
            page_image_controls[initial_index] = _fb_img_ctrl
            _fb_iv = ft.InteractiveViewer(
                key="iv_fb",
                content=_fb_img_ctrl,
                min_scale=0.5,
                max_scale=10.0,
                pan_enabled=True,
                scale_enabled=True,
                constrained=True,
                width=_fb_win_w,
                height=_fb_win_h,
                clip_behavior=ft.ClipBehavior.HARD_EDGE,
            )
            images_page_view = ft.Container(
                content=_fb_iv,
                expand=True,
                alignment=ft.Alignment(0, 0),
                bgcolor="#1a1a1a",
            )

            def _fb_navigate(new_idx: int) -> None:
                """Met à jour l'image affichée en mode fallback (sans PageView)."""
                old_idx = state["index"]
                path = image_paths[new_idx] if image_paths else ""
                _fb_img_ctrl.src = _resolve_viewer_image_src(path) if path else ""
                page_image_controls.clear()
                page_image_controls[new_idx] = _fb_img_ctrl
                pages_loaded.discard(old_idx)
                pages_loaded.add(new_idx)



        # ── Navigation ────────────────────────────────────────────────────
        async def navigate_prev(e) -> None:
            if not image_paths or state["index"] <= 0:
                return
            if _HAS_PAGE_VIEW:
                await images_page_view.previous_page(  # type: ignore[union-attr]
                    animation_curve=ft.AnimationCurve.EASE_IN_OUT_CUBIC_EMPHASIZED,
                    animation_duration=ft.Duration(milliseconds=300),
                )
            else:
                new_idx = state["index"] - 1
                _fb_navigate(new_idx)
                _update_overlay_bar(new_idx)

        async def navigate_next(e) -> None:
            if not image_paths or state["index"] >= len(image_paths) - 1:
                return
            if _HAS_PAGE_VIEW:
                await images_page_view.next_page(  # type: ignore[union-attr]
                    animation_curve=ft.AnimationCurve.EASE_IN_OUT_CUBIC_EMPHASIZED,
                    animation_duration=ft.Duration(milliseconds=300),
                )
            else:
                new_idx = state["index"] + 1
                _fb_navigate(new_idx)
                _update_overlay_bar(new_idx)



        # ── Fermeture ─────────────────────────────────────────────────────
        def close_viewer(e) -> None:
            page.on_keyboard_event = previous_keyboard_handler
            if preview_overlay in page.overlay:
                page.overlay.remove(preview_overlay)
            # Restaurer la page preview sur l'image courante
            current_path = _current_path()
            try:
                entry_index = image_paths.index(current_path)
                preview_page["value"] = entry_index // PAGE_SIZE
            except ValueError:
                pass
            refresh_preview(reset_page=False)
            page.update()



        # ── Suppression ───────────────────────────────────────────────────
        def delete_current_image(e) -> None:
            path = _current_path()
            fname = os.path.basename(path)

            def _confirm(e2):
                page.on_keyboard_event = on_key
                delete_confirmation_dialog.open = False
                page.update()
                try:
                    os.remove(path)
                    log_to_terminal(f"[OK] Supprimé: {fname}", GREEN)
                except Exception as err:
                    log_to_terminal(f"[ERREUR] {err}", RED)
                    return

                # Retirer de la sélection si présent
                if path in selected_files:
                    selected_files.remove(path)
                    selection_count_text.value = _selection_label()
                    selection_count_text.update()
                # Retirer l'entrée de all_entries_data pour cohérence
                all_entries_data["list"] = [
                    e for e in all_entries_data["list"] if e[1] != path
                ]

                cur_idx = state["index"]
                image_paths.pop(cur_idx)
                if _HAS_PAGE_VIEW:
                    images_page_view.controls.pop(cur_idx)
                page_image_controls.pop(cur_idx, None)

                if not image_paths:
                    close_viewer(None)
                    return

                # Décaler les clés de page_image_controls après cur_idx
                shifted = {}
                for k, v in page_image_controls.items():
                    shifted[k if k < cur_idx else k - 1] = v
                page_image_controls.clear()
                page_image_controls.update(shifted)

                # Vider pages_loaded pour forcer le rechargement propre de tout
                pages_loaded.clear()

                new_idx = min(cur_idx, len(image_paths) - 1)

                if _HAS_PAGE_VIEW:
                    images_page_view.selected_index = new_idx
                else:
                    _fb_navigate(new_idx)

                _update_overlay_bar(new_idx)
                _load_pages_around(new_idx)
                page.update()

                # Rafraîchir la grille en arrière-plan sans changer de page
                threading.Thread(
                    target=lambda: refresh_preview(reset_page=False),
                    daemon=True,
                ).start()

            def _cancel(e2):
                page.on_keyboard_event = on_key
                delete_confirmation_dialog.open = False
                page.update()

            def _on_key_dialog(e2: ft.KeyboardEvent):
                if e2.key == "Escape":
                    _cancel(None)
                elif e2.key == "Enter":
                    _confirm(None)

            delete_confirmation_dialog = ft.AlertDialog(
                title=ft.Text("Supprimer l'image ?"),
                content=ft.Text(f"'{fname}' sera définitivement supprimé."),
                actions=[
                    ft.TextButton("Annuler", on_click=_cancel),
                    ft.TextButton("Supprimer", on_click=_confirm, style=ft.ButtonStyle(color=ft.Colors.RED)),
                ],
            )
            page.overlay.append(delete_confirmation_dialog)
            delete_confirmation_dialog.open = True
            page.on_keyboard_event = _on_key_dialog
            page.update()



        # ── Rotation ──────────────────────────────────────────────────────
        def _rotate_current(direction: str) -> None:
            path = _current_path()
            if not path:
                return
            if os.path.splitext(path)[1].lower() not in _ROTATABLE_EXTS:
                return

            def _do_rotate():
                _rotate_files([path], direction)
                src = _resolve_viewer_image_src(path)
                cur_idx = state["index"]
                if cur_idx in page_image_controls:
                    page_image_controls[cur_idx].src = src
                    pages_loaded.discard(cur_idx)
                try:
                    page.update()
                except Exception:
                    pass

            threading.Thread(target=_do_rotate, daemon=True).start()



        # ── Clavier ───────────────────────────────────────────────────────
        def on_key(e: ft.KeyboardEvent) -> None:
            if e.key in ("Arrow Left", "ArrowLeft"):
                page.run_task(navigate_prev, e)
            elif e.key in ("Arrow Right", "ArrowRight"):
                page.run_task(navigate_next, e)
            elif e.key == "Escape":
                close_viewer(None)
            elif e.key in ("Delete", "Backspace"):
                delete_current_image(None)
            elif e.key == "[":
                _rotate_current("left")
            elif e.key == "]":
                _rotate_current("right")

        page.on_keyboard_event = on_key



        # ── UI ────────────────────────────────────────────────────────────
        button_style = ft.ButtonStyle(
            overlay_color=ft.Colors.with_opacity(0.15, ft.Colors.WHITE),
        )

        overlay_bar_color = ft.Colors.with_opacity(0.72, GREY)

        top_bar = ft.Row(
            [
                ft.Container(
                    content=ft.Column(
                        [
                            filename_text,
                            ft.Row(
                                [counter_text, ft.Text("·", size=12, color=ft.Colors.WHITE38), resolution_text],
                                alignment=ft.MainAxisAlignment.CENTER,
                                spacing=6,
                            ),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=2,
                    ),
                    bgcolor=overlay_bar_color,
                    padding=ft.Padding.symmetric(horizontal=24, vertical=10),
                    border_radius=16,
                    width=320,
                ),
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=8,
        )

        close_btn_top = ft.Container(
            content=ft.IconButton(
                icon=ft.Icons.CLOSE_ROUNDED,
                icon_color=ft.Colors.WHITE,
                icon_size=24,
                tooltip="Fermer (Échap)",
                on_click=close_viewer,
                style=button_style,
            ),
            bgcolor=overlay_bar_color,
            border_radius=20,
        )

        navigation_bar = ft.Container(
            content=ft.Row(
                [
                    ft.IconButton(
                        icon=ft.Icons.ARROW_BACK_IOS_NEW_ROUNDED,
                        icon_color=ft.Colors.WHITE,
                        icon_size=26,
                        tooltip="Image précédente",
                        on_click=navigate_prev,
                        style=button_style,
                    ),
                    ft.Container(
                        content=viewer_checkbox,
                        padding=ft.Padding.symmetric(horizontal=4, vertical=0),
                    ),
                    ft.IconButton(
                        icon=ft.Icons.ROTATE_LEFT,
                        icon_color=ft.Colors.WHITE70,
                        icon_size=22,
                        tooltip="Pivoter à gauche ([)",
                        on_click=lambda e: _rotate_current("left"),
                        style=button_style,
                    ),
                    ft.IconButton(
                        icon=ft.Icons.DELETE_ROUNDED,
                        icon_color=ft.Colors.RED_300,
                        icon_size=22,
                        tooltip="Supprimer l'image (Suppr / ⌫)",
                        on_click=delete_current_image,
                        style=button_style,
                    ),
                    ft.IconButton(
                        icon=ft.Icons.ROTATE_RIGHT,
                        icon_color=ft.Colors.WHITE70,
                        icon_size=22,
                        tooltip="Pivoter à droite (])",
                        on_click=lambda e: _rotate_current("right"),
                        style=button_style,
                    ),
                    ft.IconButton(
                        icon=ft.Icons.ARROW_FORWARD_IOS_ROUNDED,
                        icon_color=ft.Colors.WHITE,
                        icon_size=26,
                        tooltip="Image suivante",
                        on_click=navigate_next,
                        style=button_style,
                    ),
                ],
                alignment=ft.MainAxisAlignment.CENTER,
                vertical_alignment=ft.CrossAxisAlignment.CENTER,
                tight=True,
            ),
            bgcolor=overlay_bar_color,
            padding=ft.Padding.symmetric(horizontal=8, vertical=6),
            border_radius=16,
        )

        preview_overlay = ft.Stack([
            # Viewer principal (fond plein)
            images_page_view,
            # Titre/compteur centré en haut
            ft.Container(
                content=top_bar,
                top=8,
                left=0,
                right=0,
                alignment=ft.Alignment(0, 0),
            ),
            # Bouton fermer en haut à droite
            ft.Container(
                content=close_btn_top,
                top=8,
                right=8,
            ),
            # Barre de navigation inférieure flottante
            ft.Container(
                content=ft.Row(
                    [navigation_bar],
                    alignment=ft.MainAxisAlignment.CENTER,
                ),
                bottom=16,
                left=0,
                right=0,
            ),
        ], expand=True)
        page.overlay.append(preview_overlay)
        page.update()

        # Résolution initiale en arrière-plan
        def _load_initial_res():
            resolution_text.value = _get_resolution(_current_path())
            try:
                page.update()
            except Exception:
                pass

        threading.Thread(target=_load_initial_res, daemon=True).start()

        # Chargement des images autour de l'index initial
        _load_pages_around(initial_index)



    def _open_json_in_side_panel(file_path):
        """Lance Side Panel avec le fichier JSON pré-chargé dans l'onglet Liste."""
        log_to_terminal(
            f"[OK] Ouverture dans Side Panel → Liste : {os.path.basename(file_path)}",
            VIOLET,
        )
        _launch_side_panel({"SELECTEUR_JSON_PATH": file_path})



    def on_file_click(file_path, is_dir):
        """
        Gère le clic sur un élément de la preview.

        - Dossier      → navigation
        - ZIP          → extraction
        - Image        → visionneuse plein écran
        - JSON         → ouverture dans Side Panel
        - Autre        → application par défaut du système
        """
        if is_dir:
            navigate_to_folder(file_path)
        elif os.path.splitext(file_path)[1].lower() == ".zip":
            log_to_terminal(f"Extraction: {os.path.splitext(os.path.basename(file_path))[0]}", YELLOW)
            extract_zip(file_path)
        elif os.path.splitext(file_path)[1].lower() in _NOTEPAD_EXTS:
            open_file_in_notepad(file_path)
        elif os.path.splitext(file_path)[1].lower() in _IMAGE_VIEWER_EXTS:
            open_image_viewer(file_path)
        elif os.path.splitext(file_path)[1].lower() == ".json":
            _open_json_in_side_panel(file_path)
        else:
            open_file_with_default_app(file_path)



    # ================================================================ #
    #                  OPÉRATIONS SUR FICHIERS                         #
    # ================================================================ #
    def _rename_item(file_path):
        """Renomme un fichier ou un dossier via une boîte de dialogue."""
        current_name = os.path.basename(file_path)
        parent_dir = os.path.dirname(file_path)
        if current_name.lower() == CONSTANTS.THUMB_CACHE_DB_NAME.lower():
            log_to_terminal("[INFO] Le fichier .thumbcache.db est exclu du renommage.", BLUE)
            return
        stem, ext = os.path.splitext(current_name)

        name_input = ft.TextField(
            value=stem if ext else current_name,
            suffix=ft.Text(ext, color=LIGHT_GREY) if ext else None,
            autofocus=True,
            width=360,
            bgcolor=DARK,
            border_color=GREY,
            on_submit=lambda e: _do_rename(e),
        )

        def _do_rename(e):
            new_stem = name_input.value.strip() if name_input.value else ""
            if not new_stem:
                return
            new_name = new_stem + ext
            new_path = os.path.join(parent_dir, new_name)
            rename_dialog.open = False
            page.update()
            if new_name == current_name:
                _resume_keyboard_shortcuts()
                return
            try:
                os.rename(file_path, new_path)
                log_to_terminal(f"[OK] Renommé: {current_name} → {new_name}", GREEN)
                refresh_preview(reset_page=False)
                _resume_keyboard_shortcuts()
            except Exception as err:
                log_to_terminal(f"[ERREUR] Renommage: {err}", RED)
                _resume_keyboard_shortcuts()

        def _cancel_rename(e):
            rename_dialog.open = False
            page.update()
            _resume_keyboard_shortcuts()

        rename_dialog = ft.AlertDialog(
            title=ft.Text("Renommer"),
            content=name_input,
            actions=[
                ft.TextButton("Annuler", on_click=_cancel_rename),
                ft.TextButton("Renommer", on_click=_do_rename),
            ],
        )
        page.overlay.append(rename_dialog)
        _suspend_keyboard_shortcuts()
        rename_dialog.open = True
        page.update()



    def delete_item(file_path):
        """Supprime un fichier ou dossier avec confirmation"""
        def confirm_delete(e):
            """
            Exécute la suppression confirmée du fichier ou dossier.

            Appelée par le bouton « Supprimer » de la boîte de dialogue
            de confirmation. Ferme la dialog et rafraîchit la preview.
            """
            try:
                if os.path.isdir(file_path):
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                # Retirer de la sélection si présent
                if file_path in selected_files:
                    selected_files.remove(file_path)
                    selection_count_text.value = _selection_label()
                dialog.open = False
                page.update()
                refresh_preview()
                log_to_terminal(f"[OK] Supprimé: {os.path.basename(file_path)}", GREEN)
            except Exception as err:
                log_to_terminal(f"[ERREUR] Erreur lors de la suppression: {err}", RED)
                dialog.open = False
                page.update()
        
        def cancel_delete(e):
            """Annule la suppression et ferme la boîte de dialogue."""
            dialog.open = False
            page.update()
        
        item_type = "dossier" if os.path.isdir(file_path) else "fichier"
        item_name = os.path.basename(file_path)
        
        dialog = ft.AlertDialog(
            title=ft.Text(f"Supprimer {item_type}?"),
            content=ft.Text(f"Voulez-vous vraiment supprimer '{item_name}'?"),
            actions=[
                ft.TextButton("Annuler", on_click=cancel_delete),
                ft.TextButton("Supprimer", on_click=confirm_delete, style=ft.ButtonStyle(color=ft.Colors.RED)),
            ],
        )
        page.overlay.append(dialog)
        dialog.open = True
        page.update()



    def create_new_folder(e):
        """Crée un nouveau dossier dans le dossier actuel"""
        target_folder = current_browse_folder["path"] or selected_folder["path"]
        if not target_folder:
            log_to_terminal("[ERREUR] Aucun dossier sélectionné", RED)
            return
        
        def confirm_create(e):
            """
            Valide la création du nouveau dossier.

            Lit le nom saisi, crée le dossier via ``os.makedirs``,
            ferme la dialog et rafraîchit la preview.
            """
            folder_name = folder_name_input.value.strip()
            if not folder_name:
                log_to_terminal("[ERREUR] Le nom du dossier ne peut pas être vide", RED)
                return
            
            new_folder_path = os.path.join(target_folder, folder_name)
            try:
                if os.path.exists(new_folder_path):
                    log_to_terminal(f"[ERREUR] Le dossier '{folder_name}' existe déjà", RED)
                    dialog.open = False
                    page.update()
                    _resume_keyboard_shortcuts()
                else:
                    os.makedirs(new_folder_path)
                    log_to_terminal(f"[OK] Dossier créé: {folder_name}", BLUE)
                    dialog.open = False
                    page.update()
                    _resume_keyboard_shortcuts()
                    navigate_to_folder(new_folder_path)
            except Exception as err:
                log_to_terminal(f"[ERREUR] Erreur lors de la création du dossier: {err}", RED)
                dialog.open = False
                page.update()
                _resume_keyboard_shortcuts()
        
        def cancel_create(e):
            """Annule la création du dossier et ferme la boîte de dialogue."""
            dialog.open = False
            page.update()
            _resume_keyboard_shortcuts()
        
        folder_name_input = ft.TextField(
            label="Nom du dossier",
            autofocus=True,
            on_submit=confirm_create,
        )
        
        dialog = ft.AlertDialog(
            title=ft.Text("Créer un nouveau dossier"),
            content=folder_name_input,
            actions=[
                ft.TextButton("Annuler", on_click=cancel_create),
                ft.TextButton("Créer", on_click=confirm_create),
            ],
        )
        page.overlay.append(dialog)
        _suspend_keyboard_shortcuts()
        dialog.open = True
        page.update()



    def copy_selected_files(e):
        """Copie les fichiers sélectionnés dans le presse-papiers"""
        if not selected_files:
            log_to_terminal("[ATTENTION] Aucun fichier sélectionné", ORANGE)
            return
        clipboard["files"] = list(selected_files)
        clipboard["cut"] = False
        count = len(clipboard["files"])
        log_to_terminal(f"[OK] {count} élément(s) copié(s)", BLUE)
        clear_selection(None)



    def copy_to_selection_folder(e):
        """Copie les fichiers sélectionnés dans SELECTION/ et navigue vers ce dossier."""
        if not selected_files:
            log_to_terminal("[ATTENTION] Aucun fichier sélectionné", ORANGE)
            return
        folder = current_browse_folder["path"] or selected_folder["path"]
        if not folder:
            log_to_terminal("[ERREUR] Aucun dossier ouvert", RED)
            return
        selection_folder = os.path.join(folder, "SELECTION")
        os.makedirs(selection_folder, exist_ok=True)
        copied = 0
        errors = 0
        for src in list(selected_files):
            if not os.path.isfile(src):
                continue
            dst = os.path.join(selection_folder, os.path.basename(src))
            try:
                shutil.copy2(src, dst)
                copied += 1
            except Exception as ex:
                log_to_terminal(f"[ERREUR] {os.path.basename(src)} : {ex}", RED)
                errors += 1
        if copied:
            log_to_terminal(f"[OK] {copied} fichier(s) copié(s) dans SELECTION/", BLUE)
        if errors:
            log_to_terminal(f"[ATTENTION] {errors} erreur(s)", ORANGE)
        navigate_to_folder(selection_folder)



    def cut_selected_files(e):
        """Coupe les fichiers sélectionnés (déplacement à la destination)"""
        if not selected_files:
            log_to_terminal("[ATTENTION] Aucun fichier sélectionné", ORANGE)
            return
        clipboard["files"] = list(selected_files)
        clipboard["cut"] = True
        count = len(clipboard["files"])
        log_to_terminal(f"[OK] {count} élément(s) coupé(s) — collé avec Ctrl+V", ORANGE)
        clear_selection(None)



    def select_by_filter(e):
        """Sélectionne tous les fichiers visibles (en tenant compte de la recherche active)."""
        entries = all_entries_data["list"]
        if search_query["value"]:
            query_lower = search_query["value"].lower()
            entries = [entry for entry in entries if query_lower in entry[0].lower()]
        added = 0
        for _name, fpath, is_dir, _is_img, _ext in entries:
            if not is_dir and fpath not in selected_files:
                selected_files.append(fpath)
                added += 1
        if show_only_selection["value"]:
            _render_preview_page()
        else:
            _update_visible_checkboxes()
        if added:
            log_to_terminal(f"[OK] {added} fichier(s) sélectionné(s)", BLUE)
        else:
            log_to_terminal("[ATTENTION] Aucun fichier à sélectionner", ORANGE)



    def _update_select_toggle_button():
        """Met à jour l'apparence du bouton sélectionner/désélectionner."""
        if search_query["value"]:
            query_lower = search_query["value"].lower()
            filtered_paths = {
                fpath for (_name, fpath, is_dir, _is_img, _ext) in all_entries_data["list"]
                if not is_dir and query_lower in _name.lower()
            }
            all_filtered_selected = bool(filtered_paths) and filtered_paths.issubset(set(selected_files))
        else:
            all_filtered_selected = bool(selected_files)

        if all_filtered_selected:
            select_toggle_button.icon = ft.Icons.DESELECT
            select_toggle_button.icon_color = ORANGE
            select_toggle_button.tooltip = "Désélectionner tout"
        else:
            select_toggle_button.icon = ft.Icons.SELECT_ALL
            select_toggle_button.icon_color = VIOLET
            select_toggle_button.tooltip = "Tout sélectionner"
        try:
            select_toggle_button.update()
        except Exception:
            pass



    def _update_filter_sel_btn():
        """Met à jour l'apparence du bouton 'afficher uniquement la sélection'."""
        if show_only_selection["value"]:
            filter_sel_btn.icon_color = BLUE
            filter_sel_btn.tooltip = "Afficher tous les fichiers"
        else:
            filter_sel_btn.icon_color = LIGHT_GREY
            filter_sel_btn.tooltip = "Afficher uniquement la sélection"
        try:
            filter_sel_btn.update()
        except Exception:
            pass



    def _toggle_show_only_selection(e):
        """Active/désactive le filtre 'uniquement la sélection'."""
        show_only_selection["value"] = not show_only_selection["value"]
        _update_filter_sel_btn()
        _render_preview_page()



    def toggle_select_all(e):
        """Sélectionne tout si rien sélectionné, sinon désélectionne tout.
        Si une recherche est active : sélectionne d'abord les fichiers filtrés,
        puis désélectionne tout au second appui."""
        if search_query["value"]:
            query_lower = search_query["value"].lower()
            filtered_paths = {
                fpath for (_name, fpath, is_dir, _is_img, _ext) in all_entries_data["list"]
                if not is_dir and query_lower in _name.lower()
            }
            if not filtered_paths.issubset(set(selected_files)):
                # Premier appuis : sélectionner les fichiers filtrés
                select_by_filter(e)
            else:
                # Deuxième appui : tout désélectionner
                clear_selection(e)
        else:
            if selected_files:
                clear_selection(e)
            else:
                select_by_filter(e)



    def invert_selection(e):
        """Inverse la sélection : sélectionne les non-sélectionnés, désélectionne les sélectionnés."""
        entries = all_entries_data["list"]
        if search_query["value"]:
            query_lower = search_query["value"].lower()
            entries = [entry for entry in entries if query_lower in entry[0].lower()]
        for _name, fpath, is_dir, _is_img, _ext in entries:
            if is_dir:
                continue
            if fpath in selected_files:
                selected_files.remove(fpath)
            else:
                selected_files.append(fpath)
        if show_only_selection["value"]:
            _render_preview_page()
        else:
            _update_visible_checkboxes()
        log_to_terminal(f"[OK] Sélection inversée — {len(selected_files)} fichier(s) sélectionné(s)", BLUE)



    def _get_image_reference_date(file_path: str):
        """Retourne la date de modification du fichier."""
        try:
            return datetime.date.fromtimestamp(os.path.getmtime(file_path))
        except OSError:
            return None



    def select_same_date(e):
        """Sélectionne tous les fichiers du dossier pris à la même date que le fichier sélectionné."""
        log_to_terminal("[CMD] select_same_date()", BLUE)
        if not selected_files:
            log_to_terminal("[ATTENTION] Aucun fichier sélectionné comme référence", ORANGE)
            return
        ref_path = selected_files[-1]
        ref_date = _get_image_reference_date(ref_path)
        if ref_date is None:
            log_to_terminal("[ERREUR] Impossible de lire la date du fichier de référence", RED)
            return
        added = 0
        for _name, fpath, is_dir, _is_img, _ext in all_entries_data["list"]:
            if is_dir:
                continue

            fdate = _get_image_reference_date(fpath)
            if fdate is None:
                continue
            if fdate == ref_date and fpath not in selected_files:
                selected_files.append(fpath)
                added += 1
        if show_only_selection["value"]:
            _render_preview_page()
        else:
            _update_visible_checkboxes()
        _update_select_toggle_button()
        log_to_terminal(f"[OK] {len(selected_files)} fichier(s) du {ref_date.strftime('%d/%m/%Y')} sélectionné(s) (+{added} ajouté(s))", BLUE)



    def paste_files(e):
        """Colle les fichiers du presse-papiers dans le dossier actuel"""
        target_folder = current_browse_folder["path"] or selected_folder["path"]
        if not target_folder:
            log_to_terminal("[ERREUR] Aucun dossier de destination sélectionné", RED)
            return
        
        if not clipboard["files"]:
            log_to_terminal("[ATTENTION] Presse-papiers vide", ORANGE)
            return

        # Snapshot du presse-papiers avant de lancer le thread
        files_to_paste = list(clipboard["files"])
        is_cut = clipboard["cut"]
        count = len(files_to_paste)
        action_label = "déplacement" if is_cut else "copie"
        log_to_terminal(f"[...] {action_label.capitalize()} de {count} élément(s) en cours…", ORANGE)

        def _do_paste():
            copied_count = 0
            errors = []

            for source_path in files_to_paste:
                if not os.path.exists(source_path):
                    errors.append(f"{os.path.basename(source_path)}: fichier source introuvable")
                    continue

                dest_path = os.path.join(target_folder, os.path.basename(source_path))

                # Si le fichier existe déjà, ajouter un suffixe
                if os.path.exists(dest_path):
                    base_name = os.path.basename(source_path)
                    name, ext = os.path.splitext(base_name)
                    counter = 1
                    while os.path.exists(dest_path):
                        new_name = f"{name} ({counter}){ext}"
                        dest_path = os.path.join(target_folder, new_name)
                        counter += 1

                try:
                    if os.path.isdir(source_path):
                        shutil.copytree(source_path, dest_path)
                    else:
                        shutil.copy2(source_path, dest_path)
                    copied_count += 1

                    # Si mode couper : supprimer la source après copie réussie
                    if is_cut:
                        try:
                            if os.path.isdir(source_path):
                                shutil.rmtree(source_path)
                            else:
                                os.remove(source_path)
                            if source_path in selected_files:
                                selected_files.remove(source_path)
                        except Exception as del_err:
                            errors.append(f"Suppression source {os.path.basename(source_path)}: {del_err}")
                except Exception as err:
                    errors.append(f"{os.path.basename(source_path)}: {err}")

            if copied_count > 0:
                action = "déplacé" if is_cut else "collé"
                log_to_terminal(f"[OK] {copied_count} élément(s) {action}(s)", BLUE)
                if is_cut:
                    clipboard["files"] = []
                    clipboard["cut"] = False
                    selection_count_text.value = _selection_label()

            if errors:
                for error in errors:
                    log_to_terminal(f"[ERREUR] {error}", RED)

            app_progress_bar.visible = False
            try:
                page.update()
            except Exception:
                pass
            refresh_preview()

        app_progress_bar.visible = True
        try:
            page.update()
        except Exception:
            pass
        threading.Thread(target=_do_paste, daemon=True).start()



    def _apply_print_rename(file_path, new_count):
        basename = os.path.basename(file_path)
        folder = os.path.dirname(file_path)

        if basename.lower() == CONSTANTS.THUMB_CACHE_DB_NAME.lower():
            return
        
        if not os.path.exists(file_path):
            return
            
        print_prefix_match = re.match(r'^(\d+)X_', basename, re.IGNORECASE)
        print_prefix_pattern = re.compile(r'^\d+X_', re.IGNORECASE)
        
        if print_prefix_match:
            current_count = int(print_prefix_match.group(1))
            clean_basename = re.sub(r'^\d+X_', '', basename, flags=re.IGNORECASE)
        else:
            current_count = 0
            clean_basename = basename
            
        if new_count > 0:
            new_name = f"{new_count}X_{clean_basename}"
        else:
            new_name = clean_basename
            
        new_path = os.path.join(folder, new_name)
        
        norm_file_path = os.path.normpath(file_path).lower()
        norm_new_path = os.path.normpath(new_path).lower()
        
        if norm_new_path != norm_file_path:
            try:
                os.rename(file_path, new_path)
                log_to_terminal(f"[Impressions] {basename} → {new_name}", GREEN)
                
                # Mettre à jour dans selected_files
                if file_path in selected_files:
                    selected_files[selected_files.index(file_path)] = new_path
                    
                # Mettre à jour le dictionnaire de live counts
                _live_print_counts[new_path] = new_count
                _live_print_counts.pop(file_path, None)
                
                # Mettre à jour les références UI refs si elles existent
                if file_path in _print_count_text_refs:
                    _print_count_text_refs[new_path] = _print_count_text_refs.pop(file_path)
                if file_path in _print_minus_btn_refs:
                    _print_minus_btn_refs[new_path] = _print_minus_btn_refs.pop(file_path)
                    
                    # Mettre à jour l'action du bouton Moins pour le nouveau chemin
                    minus_ref = _print_minus_btn_refs[new_path]
                    minus_ref.on_click = (lambda e, p=new_path: _decrement_print_count(p)) if new_count > 0 else None
                    try:
                        minus_ref.update()
                    except Exception:
                        pass  # Ignorer si le bouton n'est plus monté dans la page
                
                page.pubsub.send_all_on_topic("refresh", None)
            except Exception as err:
                log_to_terminal(f"[ERREUR] {err}", RED)
                return
                
        # 1. Si le fichier est repassé à 0, on remet TOUS les autres fichiers du dossier à 0 aussi
        if new_count <= 0:
            renamed_count = 0
            for file_name in os.listdir(folder):
                if file_name.startswith(".") or file_name.lower() in _OS_JUNK:
                    continue
                entry_path = os.path.join(folder, file_name)
                if not os.path.isfile(entry_path) or os.path.normpath(entry_path).lower() == norm_new_path:
                    continue
                if not print_prefix_pattern.match(file_name):
                    continue
                clean_other_basename = re.sub(r'^\d+X_', '', file_name, flags=re.IGNORECASE)
                clean_entry_path = os.path.join(folder, clean_other_basename)
                try:
                    os.rename(entry_path, clean_entry_path)
                    if entry_path in selected_files:
                        selected_files[selected_files.index(entry_path)] = clean_entry_path
                    
                    _live_print_counts[clean_entry_path] = 0
                    _live_print_counts.pop(entry_path, None)
                    
                    if entry_path in _print_count_text_refs:
                        txt_ref = _print_count_text_refs[entry_path]
                        txt_ref.value = "·"
                        txt_ref.color = LIGHT_GREY
                        try:
                            txt_ref.update()
                        except Exception:
                            pass  # Ignorer si le texte n'est plus monté
                        _print_count_text_refs[clean_entry_path] = _print_count_text_refs.pop(entry_path)
                        
                    if entry_path in _print_minus_btn_refs:
                        minus_ref = _print_minus_btn_refs[entry_path]
                        minus_ref.content.color = LIGHT_GREY
                        minus_ref.bgcolor = GREY
                        minus_ref.on_click = None
                        minus_ref.ink = False
                        try:
                            minus_ref.update()
                        except Exception:
                            pass  # Ignorer si le bouton n'est plus monté
                        _print_minus_btn_refs[clean_entry_path] = _print_minus_btn_refs.pop(entry_path)
                        
                    renamed_count += 1
                except Exception as err:
                    log_to_terminal(f"[ERREUR] {file_name}: {err}", RED)
            if renamed_count:
                log_to_terminal(f"[OK] Préfixe retiré de {renamed_count} fichier(s)", GREEN)
                page.pubsub.send_all_on_topic("refresh", None)
                
        # 2. Gérer l'auto-préfixe 1X_ si besoin
        elif current_count <= 0 and new_count > 0:
            renamed_count = 0
            for file_name in os.listdir(folder):
                if file_name.startswith(".") or file_name.lower() in _OS_JUNK:
                    continue
                entry_path = os.path.join(folder, file_name)
                if not os.path.isfile(entry_path) or os.path.normpath(entry_path).lower() == norm_new_path:
                    continue
                if print_prefix_pattern.match(file_name):
                    continue
                    
                new_file_name = f"1X_{file_name}"
                new_entry_path = os.path.join(folder, new_file_name)
                try:
                    os.rename(entry_path, new_entry_path)
                    if entry_path in selected_files:
                        selected_files[selected_files.index(entry_path)] = new_entry_path
                    
                    _live_print_counts[new_entry_path] = 1
                    _live_print_counts.pop(entry_path, None)
                    
                    if entry_path in _print_count_text_refs:
                        txt_ref = _print_count_text_refs[entry_path]
                        txt_ref.value = "1"
                        txt_ref.color = YELLOW
                        try:
                            txt_ref.update()
                        except Exception:
                            pass  # Ignorer si le texte n'est plus monté
                        _print_count_text_refs[new_entry_path] = _print_count_text_refs.pop(entry_path)
                        
                    if entry_path in _print_minus_btn_refs:
                        minus_ref = _print_minus_btn_refs[entry_path]
                        minus_ref.content.color = DARK
                        minus_ref.bgcolor = ORANGE
                        minus_ref.on_click = lambda e, p=new_entry_path: _decrement_print_count(p)
                        minus_ref.ink = True
                        try:
                            minus_ref.update()
                        except Exception:
                            pass  # Ignorer si le bouton n'est plus monté
                        _print_minus_btn_refs[new_entry_path] = _print_minus_btn_refs.pop(entry_path)
                        
                    renamed_count += 1
                except Exception as err:
                    log_to_terminal(f"[ERREUR] {file_name}: {err}", RED)
            if renamed_count:
                log_to_terminal(f"[OK] {renamed_count} fichier(s) renommé(s) avec le préfixe 1X_", GREEN)
                page.pubsub.send_all_on_topic("refresh", None)



    def _increment_print_count(file_path):
        """Incrémente le compteur d'impressions de façon réactive et débouncée."""
        if file_path in _live_print_counts:
            current_count = _live_print_counts[file_path]
        else:
            basename = os.path.basename(file_path)
            print_prefix_match = re.match(r'^(\d+)X_', basename, re.IGNORECASE)
            current_count = int(print_prefix_match.group(1)) if print_prefix_match else 0
            _live_print_counts[file_path] = current_count
            
        new_count = current_count + 1
        _live_print_counts[file_path] = new_count
        
        # Mettre à jour l'UI instantanément
        if file_path in _print_count_text_refs:
            txt_ref = _print_count_text_refs[file_path]
            txt_ref.value = str(new_count) if new_count > 0 else "·"
            txt_ref.color = YELLOW if new_count > 0 else LIGHT_GREY
            txt_ref.update()
            
        if file_path in _print_minus_btn_refs:
            minus_ref = _print_minus_btn_refs[file_path]
            minus_ref.content.color = DARK if new_count > 0 else LIGHT_GREY
            minus_ref.bgcolor = ORANGE if new_count > 0 else GREY
            minus_ref.on_click = (lambda e, p=file_path: _decrement_print_count(p)) if new_count > 0 else None
            minus_ref.ink = (new_count > 0)
            minus_ref.update()
            
        # Annuler l'ancien timer et démarrer le nouveau
        if file_path in _print_rename_timers:
            _print_rename_timers[file_path].cancel()
            
        timer = threading.Timer(1.5, _apply_print_rename, args=(file_path, new_count))
        _print_rename_timers[file_path] = timer
        timer.start()


    def _decrement_print_count(file_path):
        """Décrémente le compteur d'impressions de façon réactive et débouncée."""
        if file_path in _live_print_counts:
            current_count = _live_print_counts[file_path]
        else:
            basename = os.path.basename(file_path)
            print_prefix_match = re.match(r'^(\d+)X_', basename, re.IGNORECASE)
            current_count = int(print_prefix_match.group(1)) if print_prefix_match else 0
            _live_print_counts[file_path] = current_count
            
        if current_count <= 0:
            return
            
        new_count = current_count - 1
        _live_print_counts[file_path] = new_count
        
        # Mettre à jour l'UI instantanément
        if file_path in _print_count_text_refs:
            txt_ref = _print_count_text_refs[file_path]
            txt_ref.value = str(new_count) if new_count > 0 else "·"
            txt_ref.color = YELLOW if new_count > 0 else LIGHT_GREY
            txt_ref.update()
            
        if file_path in _print_minus_btn_refs:
            minus_ref = _print_minus_btn_refs[file_path]
            minus_ref.content.color = DARK if new_count > 0 else LIGHT_GREY
            minus_ref.bgcolor = ORANGE if new_count > 0 else GREY
            minus_ref.on_click = (lambda e, p=file_path: _decrement_print_count(p)) if new_count > 0 else None
            minus_ref.ink = (new_count > 0)
            minus_ref.update()
            
        # Annuler l'ancien timer et démarrer le nouveau
        if file_path in _print_rename_timers:
            _print_rename_timers[file_path].cancel()
            
        timer = threading.Timer(1.5, _apply_print_rename, args=(file_path, new_count))
        _print_rename_timers[file_path] = timer
        timer.start()



    # ================================================================ #
    #                          SÉLECTION                               #
    # ================================================================ #
    def _selection_label():
        """Retourne le libellé de sélection affiché dans la barre d'état."""
        selected_count = len(selected_files)
        if selected_count == 0:
            return ""
        plural_suffix = "s" if selected_count > 1 else ""
        return f"{selected_count} fichier{plural_suffix} sélectionné{plural_suffix}"



    def on_checkbox_change(e, file_path):
        """Gère le changement d'état d'une checkbox"""
        if e.control.value:
            if file_path not in selected_files:
                selected_files.append(file_path)
        else:
            if file_path in selected_files:
                selected_files.remove(file_path)
        selection_count_text.value = _selection_label()
        _update_select_toggle_button()
        selection_count_text.update()



    def _update_visible_checkboxes():
        """Met à jour les cases à cocher en place sans reconstruire les widgets image."""
        for file_path, checkbox_widget in _checkbox_refs.items():
            checkbox_widget.value = file_path in selected_files
        selection_count_text.value = _selection_label()
        _update_select_toggle_button()
        page.update()



    def clear_selection(e):
        """Désélectionne tous les fichiers et rafraîchit la preview."""
        selected_files.clear()
        if show_only_selection["value"]:
            show_only_selection["value"] = False
            _update_filter_sel_btn()
            _render_preview_page()  # la liste filtrée change, re-rendu nécessaire
        else:
            _update_visible_checkboxes()  # seulement les cases à cocher changent
        log_to_terminal("[OK] Sélection effacée", GREEN)



    def delete_selected_files(e):
        """Supprime tous les fichiers et dossiers sélectionnés avec confirmation"""
        if not selected_files:
            log_to_terminal("[ATTENTION] Aucun élément sélectionné", ORANGE)
            return
        
        def confirm_delete_multiple(e):
            """
            Supprime tous les éléments sélectionnés après confirmation.

            Itère sur ``selected_files``, supprime fichiers et dossiers,
            reporte le nombre de suppressions réussies dans le terminal.
            """
            deleted_files = 0
            deleted_folders = 0
            errors = []
            for item_path in list(selected_files):
                try:
                    if os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                        deleted_folders += 1
                    elif os.path.isfile(item_path):
                        os.remove(item_path)
                        deleted_files += 1
                except Exception as err:
                    errors.append(f"{os.path.basename(item_path)}: {err}")
            
            selected_files.clear()
            selection_count_text.value = _selection_label()
            dialog.open = False
            page.update()
            refresh_preview()
            
            if deleted_files > 0 or deleted_folders > 0:
                msg_parts = []
                if deleted_files > 0:
                    msg_parts.append(f"{deleted_files} fichier(s)")
                if deleted_folders > 0:
                    msg_parts.append(f"{deleted_folders} dossier(s)")
                log_to_terminal(f"[OK] Supprimé: {' et '.join(msg_parts)}", GREEN)
            if errors:
                for error in errors:
                    log_to_terminal(f"[ERREUR] {error}", RED)
        
        def cancel_delete_multiple(e):
            """Annule la suppression multiple et ferme la boîte de dialogue."""
            dialog.open = False
            page.update()
        
        item_count = len(selected_files)
        
        dialog = ft.AlertDialog(
            title=ft.Text(f"Supprimer {item_count} élément(s)?"),
            content=ft.Text(f"Voulez-vous vraiment supprimer les {item_count} élément(s) sélectionné(s)?"),
            actions=[
                ft.TextButton("Annuler", on_click=cancel_delete_multiple),
                ft.TextButton("Supprimer", on_click=confirm_delete_multiple, style=ft.ButtonStyle(color=ft.Colors.RED)),
            ],
        )
        page.overlay.append(dialog)
        dialog.open = True
        page.update()



    def apply_selected_files_by_name(selected_names_str):
        """Sélectionne dans la preview les éléments correspondant aux noms fournis"""
        folder_to_display = current_browse_folder["path"] or selected_folder["path"]
        if not folder_to_display or not os.path.isdir(folder_to_display):
            return

        names_ordered = [name for name in selected_names_str.split("|") if name]
        names_to_select = set(names_ordered)

        selected_files.clear()



        # Utilise all_entries_data pour garantir que les chemins dans selected_files
        # sont identiques (même objet str) à ceux utilisés par _render_current_page(),
        # évitant toute divergence de normalisation Unicode NFD/NFC sur macOS.
        entries = all_entries_data["list"]
        if names_to_select and entries:
            name_to_path = {file_name: file_path for file_name, file_path, is_dir, is_image, ext in entries if file_name in names_to_select}
            for name in names_ordered:
                if name in name_to_path:
                    selected_files.append(name_to_path[name])
        elif names_to_select:
            # Fallback si entries pas encore peuplées
            existing = {item_name: os.path.join(folder_to_display, item_name) for item_name in os.listdir(folder_to_display)}
            for name in names_ordered:
                if name in existing:
                    selected_files.append(existing[name])

        selection_count_text.value = _selection_label()



        # Naviguer vers la première page contenant au moins un fichier sélectionné
        # (évite que les fichiers au-delà de PAGE_SIZE soient invisibles)
        if selected_files and entries:
            for entry_index, (file_name, file_path, is_dir, is_image, ext) in enumerate(entries):
                if file_path in selected_files:
                    preview_page["value"] = entry_index // PAGE_SIZE
                    break



        # Re-rendu immédiat sans nouveau scan (le script ne modifie aucun fichier)
        _render_preview_page()

        if names_to_select and not selected_files:
            log_to_terminal("[ATTENTION] Aucun fichier correspondant trouvé dans la preview", ORANGE)



    # ================================================================ #
    #                    FILTRAGE & PÉRIPHÉRIQUES                      #
    # ================================================================ #
    # ── Recherche dans la preview ──────────────────────────────────────
    def _on_search_change(e):
        """Met à jour la requête de recherche et re-rend la preview.
        Si le champ est vidé, referme la barre et restaure l'icône seule."""
        typed_text = (search_field.value or "").strip()
        if not typed_text:
            _clear_search(e)
            return
        if show_only_selection["value"]:
            show_only_selection["value"] = False
            _update_filter_sel_btn()
        search_query["value"] = typed_text
        preview_page["value"] = 0
        _render_preview_page()



    def _clear_search(e):
        """Efface la recherche et restaure tous les fichiers."""
        search_query["value"] = ""
        search_field.value = ""
        _render_preview_page()
        page.update()



    def _get_removable_drives():
        """Détecte les périphériques amovibles (cross-platform, sans dépendance externe)."""
        drives = []
        try:
            if platform.system() == "Darwin":
                macos_system_volumes = {
                    "Macintosh HD", "Macintosh HD - Data",
                    "com.apple.TimeMachine.localsnapshots",
                    "Recovery", "Preboot", "VM", "Update",
                }
                for entry in os.scandir("/Volumes"):
                    if entry.is_dir() and os.path.ismount(entry.path) and entry.name not in macos_system_volumes and not entry.name.startswith("."):
                        drives.append((entry.name, entry.path))
                # macOS crée parfois NAME et NAME-1 simultanément (stub APFS + nouveau montage).
                # Dans ce cas, NAME est le stub vide et NAME-1 est le vrai volume :
                # on garde l'entrée NAME mais on la fait pointer vers NAME-1.
                _sfx_re = re.compile(r'^(.*?)-(\d+)$')
                _drive_names = {n for n, _ in drives}
                _replacements = {}
                for _dn, _dp in drives:
                    _m = _sfx_re.match(_dn)
                    if _m:
                        _base, _n = _m.group(1), int(_m.group(2))
                        if _base in _drive_names:
                            _prev = _replacements.get(_base)
                            if _prev is None or _n > _prev[0]:
                                _replacements[_base] = (_n, _dp)
                if _replacements:
                    _skip = {_dn for _dn, _ in drives
                             if _sfx_re.match(_dn) and _sfx_re.match(_dn).group(1) in _replacements}
                    drives = [(_dn, _replacements[_dn][1]) if _dn in _replacements else (_dn, _dp)
                              for _dn, _dp in drives if _dn not in _skip]

            elif platform.system() == "Windows":
                import ctypes
                DRIVE_TYPE_REMOVABLE, DRIVE_TYPE_CDROM = 2, 5
                volume_label_buffer = ctypes.create_unicode_buffer(261)
                for letter in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
                    path = f"{letter}:\\"
                    drive_type = ctypes.windll.kernel32.GetDriveTypeW(path)
                    if drive_type in (DRIVE_TYPE_REMOVABLE, DRIVE_TYPE_CDROM) and os.path.exists(path):
                        ctypes.windll.kernel32.GetVolumeInformationW(
                            path, volume_label_buffer, 261, None, None, None, None, 0)
                        label = volume_label_buffer.value or letter
                        drives.append((f"{label} ({letter}:)", path))

            else:  # Linux
                for base in ("/media", "/run/media"):
                    if not os.path.isdir(base):
                        continue
                    for entry in os.scandir(base):
                        if not entry.is_dir():
                            continue
                        if os.path.ismount(entry.path):
                            drives.append((entry.name, entry.path))
                        else:
                            try:
                                for sub in os.scandir(entry.path):
                                    if sub.is_dir() and os.path.ismount(sub.path):
                                        drives.append((sub.name, sub.path))
                            except PermissionError:
                                pass
        except Exception:
            pass
        return drives



    def _rebuild_favorites_panel():
        """Reconstruit la liste des favoris dans le conteneur."""
        favorites_list = _load_favorites()
        favorites_list_view.controls.clear()
        if not favorites_list:
            favorites_list_view.controls.append(
                ft.Text("Aucun favori — cliquez sur + pour ajouter", size=10,
                        color=LIGHT_GREY, italic=True)
            )
        else:
            for i, fav in enumerate(favorites_list):
                p = fav["path"]
                display_name = fav["label"] or os.path.basename(p) or p
                def _nav(e, path=p):
                    resolved = _resolve_favorite_path(path)
                    if os.path.isdir(resolved):
                        navigate_to_folder(resolved)
                    else:
                        log_to_terminal(f"[ERREUR] Dossier introuvable : {path}", RED)
                def _remove(e, path=p):
                    updated_favorites = _load_favorites()
                    updated_favorites = [f for f in updated_favorites if f["path"] != path]
                    _save_favorites(updated_favorites)
                    _rebuild_favorites_panel()
                    try:
                        favorites_list_view.update()
                    except Exception:
                        pass
                def _rename(e, path=p, current_label=fav["label"]):
                    _show_rename_favorite_dialog(path, current_label)
                favorites_list_view.controls.append(
                    ft.Row([
                        ft.ReorderableDragHandle(
                            content=ft.Icon(ft.Icons.DRAG_INDICATOR, size=16, color=LIGHT_GREY),
                            mouse_cursor=ft.MouseCursor.GRAB,
                        ),
                        ft.Icon(ft.Icons.FOLDER, size=16, color=BLUE),
                        ft.Container(
                            content=ft.Text(display_name, size=16, color=WHITE,
                                            overflow=ft.TextOverflow.ELLIPSIS, max_lines=1),
                            expand=True,
                            on_click=_nav,
                            tooltip=p,
                            ink=True,
                        ),
                        ft.Container(
                            content=ft.Icon(ft.Icons.EDIT_OUTLINED, size=15, color=LIGHT_GREY),
                            on_click=_rename,
                            tooltip="Renommer le raccourci",
                            ink=True,
                            border_radius=8,
                            padding=ft.Padding(3, 2, 3, 2),
                        ),
                        ft.Container(
                            content=ft.Icon(ft.Icons.CLOSE, size=16, color=LIGHT_GREY),
                            on_click=_remove,
                            tooltip="Retirer des favoris",
                            ink=True,
                            border_radius=8,
                            padding=ft.Padding(3, 2, 3, 2),
                        ),
                    ], spacing=4, tight=True, height=32, vertical_alignment=ft.CrossAxisAlignment.CENTER)
                )
        try:
            favorites_list_view.update()
        except Exception:
            pass



    def _show_rename_favorite_dialog(path: str, current_label: str):
        """Ouvre un dialog pour renommer le raccourci d'un favori."""
        label_field = ft.TextField(
            value=current_label or os.path.basename(path),
            hint_text=os.path.basename(path),
            border_color=BLUE, text_size=13, height=40,
            content_padding=ft.Padding(8, 4, 8, 4),
            autofocus=True,
        )
        dlg = ft.AlertDialog(
            title=ft.Text("Renommer le raccourci", size=13, color=WHITE),
            content=ft.Column([
                ft.Text(path, size=10, color=LIGHT_GREY,
                        overflow=ft.TextOverflow.ELLIPSIS, max_lines=1),
                label_field,
            ], spacing=6, tight=True, width=280),
            actions_alignment=ft.MainAxisAlignment.END,
        )

        def _confirm(e):
            new_label = (label_field.value or "").strip()
            favorites = _load_favorites()
            for fav in favorites:
                if fav["path"] == path:
                    fav["label"] = new_label
                    break
            _save_favorites(favorites)
            dlg.open = False
            page.update()
            _rebuild_favorites_panel()

        def _cancel(e):
            dlg.open = False
            page.update()

        dlg.actions = [
            ft.TextButton("OK", on_click=_confirm),
            ft.TextButton("Annuler", on_click=_cancel),
        ]
        page.overlay.append(dlg)
        dlg.open = True
        page.update()



    def _add_favorite_current():
        """Ajoute le dossier courant aux favoris (avec choix du nom du raccourci)."""
        path = current_browse_folder.get("path") or selected_folder.get("path")
        if not path or not os.path.isdir(path):
            log_to_terminal("[ATTENTION] Aucun dossier sélectionné à ajouter en favori", ORANGE)
            return
        path = os.path.normpath(path)
        favorites_list = _load_favorites()
        if any(favorite["path"] == path for favorite in favorites_list):
            log_to_terminal("[INFO] Ce dossier est déjà dans les favoris", LIGHT_GREY)
            return

        default_name = os.path.basename(path)
        label_field = ft.TextField(
            value=default_name,
            hint_text=default_name,
            border_color=BLUE, text_size=13, height=40,
            content_padding=ft.Padding(8, 4, 8, 4),
            autofocus=True,
        )
        dlg = ft.AlertDialog(
            title=ft.Text("Ajouter aux favoris", size=13, color=WHITE),
            content=ft.Column([
                ft.Text("Nom du raccourci :", size=11, color=LIGHT_GREY),
                label_field,
            ], spacing=6, tight=True, width=280),
            actions_alignment=ft.MainAxisAlignment.END,
        )

        def _confirm(e):
            label = (label_field.value or "").strip()
            dlg.open = False
            page.update()
            favorites_list = _load_favorites()
            if not any(favorite["path"] == path for favorite in favorites_list):
                favorites_list.insert(0, {"path": path, "label": label})
                _save_favorites(favorites_list)
                _rebuild_favorites_panel()
                log_to_terminal(f"[OK] Favori ajouté : {label or default_name}", YELLOW)

        def _cancel(e):
            dlg.open = False
            page.update()

        dlg.actions = [
            ft.TextButton("Ajouter", on_click=_confirm),
            ft.TextButton("Annuler", on_click=_cancel),
        ]
        page.overlay.append(dlg)
        dlg.open = True
        page.update()



    def _eject_drive(path):
        """Démonte/éjecte un périphérique amovible selon le système d'exploitation.
        Réessaie automatiquement jusqu'à 3 fois si l'éjection n'est pas confirmée."""
        log_to_terminal(f"[...] Éjection en cours : {path}", VIOLET)

        def _run():
            sys_name = platform.system()
            max_attempts = 3
            for attempt in range(1, max_attempts + 1):
                try:
                    if sys_name == "Windows":
                        drive_letter = os.path.splitdrive(path)[0]  # ex: "E:"
                        powershell_eject_command = (
                            f"(New-Object -comObject Shell.Application)"
                            f".Namespace(17).ParseName('{drive_letter}').InvokeVerb('Eject')"
                        )
                        subprocess.run(["powershell", "-Command", powershell_eject_command],
                                       creationflags=subprocess.CREATE_NO_WINDOW,
                                       timeout=10)
                        # InvokeVerb est asynchrone : on attend que le lecteur disparaisse
                        time.sleep(1.5)
                        if not os.path.exists(path):
                            log_to_terminal(f"[OK] Éjecté : {path}", VIOLET)
                            return
                    elif sys_name == "Darwin":
                        result = subprocess.run(
                            ["diskutil", "eject", path],
                            capture_output=True, text=True, timeout=10,
                        )
                        if result.returncode == 0:
                            log_to_terminal(f"[OK] Éjecté : {path}", VIOLET)
                            return
                    else:  # Linux
                        result = subprocess.run(
                            ["umount", path],
                            capture_output=True, text=True, timeout=10,
                        )
                        if result.returncode == 0:
                            log_to_terminal(f"[OK] Éjecté : {path}", VIOLET)
                            return
                except subprocess.TimeoutExpired:
                    pass
                except Exception as ex:
                    log_to_terminal(f"[ERREUR] Éjection impossible : {ex}", RED)
                    return

                if attempt < max_attempts:
                    log_to_terminal(f"[...] Tentative {attempt}/{max_attempts} échouée — nouvel essai…", ORANGE)
                    time.sleep(1)

            log_to_terminal(f"[ATTENTION] Éjection non confirmée après {max_attempts} tentatives : {path}", ORANGE)

        threading.Thread(target=_run, daemon=True).start()



    def _rebuild_drives_panel(drives):
        """Met à jour la section périphériques (appelée depuis le callback pubsub)."""
        drives_list_view.controls.clear()
        for name, path in drives:
            def _nav(e, p=path):
                if os.path.isdir(p):
                    navigate_to_folder(p)
                else:
                    log_to_terminal(f"[ERREUR] Périphérique introuvable : {p}", RED)
            def _eject(e, p=path):
                _eject_drive(p)
            drives_list_view.controls.append(
                ft.Row([
                    ft.ReorderableDragHandle(
                        content=ft.Icon(ft.Icons.DRAG_INDICATOR, size=16, color=LIGHT_GREY),
                        mouse_cursor=ft.MouseCursor.GRAB,
                    ),
                    ft.Icon(ft.Icons.STORAGE, size=16, color=VIOLET),
                    ft.Container(
                        content=ft.Text(name, size=16, color=WHITE,
                                        overflow=ft.TextOverflow.ELLIPSIS, max_lines=1),
                        expand=True,
                        on_click=_nav,
                        tooltip=path,
                        ink=True,
                    ),
                    ft.Container(
                        content=ft.Icon(ft.Icons.EJECT, size=16, color=LIGHT_GREY),
                        on_click=_eject,
                        tooltip="Éjecter le périphérique",
                        ink=True,
                        border_radius=8,
                        padding=ft.Padding(3, 2, 3, 2),
                    ),
                ], spacing=4, tight=True, height=32, vertical_alignment=ft.CrossAxisAlignment.CENTER)
            )
        drives_panel.visible = bool(drives)
        try:
            drives_panel.update()
            page.update()
        except Exception:
            pass



    def _on_drives_changed(topic, drives):
        """Callback pubsub : met à jour l'UI périphériques depuis le thread de fond."""
        _rebuild_drives_panel(drives)

    page.pubsub.subscribe_topic("drives_changed", _on_drives_changed)



    def _poll_removable_drives():
        """Thread de fond : détecte les changements de périphériques toutes les 3 s."""
        prev_drives = []
        ordered_drives = []
        while True:
            time.sleep(3)
            try:
                drives = _get_removable_drives()
                if drives != prev_drives:
                    new_drives = [drive for drive in drives if drive not in prev_drives]
                    existing_drives = [drive for drive in ordered_drives if drive in drives]
                    ordered_drives = new_drives + existing_drives
                    prev_drives = drives
                    removable_drives_state["list"] = ordered_drives
                    page.pubsub.send_all_on_topic("drives_changed", ordered_drives)
            except Exception:
                pass



    # ================================================================ #
    #                           PREVIEW                                #
    # ================================================================ #
    def _render_preview_page():
        """
        Construit et affiche les contrôles ListView pour la page courante.
        Appelée depuis on_preview_ready et go_to_page (thread UI uniquement).
        Les images sont affichées d'abord avec une icône placeholder, puis les miniatures
        sont générées en arrière-plan et mises à jour sans bloquer l'UI.
        """
        try:
            _checkbox_refs.clear()
            _pending_thumb_refs.clear()
            entries = all_entries_data["list"]
            # Appliquer la recherche textuelle
            if search_query["value"]:
                query_lower = search_query["value"].lower()
                entries = [entry for entry in entries if query_lower in entry[0].lower()]
            # Afficher uniquement la sélection si actif
            if show_only_selection["value"]:
                entries = [entry for entry in entries if entry[1] in selected_files]
            total = len(entries)
            total_pages = max(1, (total + PAGE_SIZE - 1) // PAGE_SIZE)
            current_pg = min(preview_page["value"], total_pages - 1)
            preview_page["value"] = current_pg

            new_controls = []

            if all_entries_data["error"]:
                new_controls.append(ft.Text(all_entries_data["error"], color="red"))
            elif not entries:
                folder_to_display = current_browse_folder["path"] or selected_folder["path"]
                if folder_to_display:
                    new_controls.append(ft.Text("(dossier vide)", color=GREY))
            else:
                start = current_pg * PAGE_SIZE
                end = min(start + PAGE_SIZE, total)
                for list_idx, (file, file_path, is_dir, is_image, ext) in enumerate(entries[start:end]):
                    if is_dir:
                        icon = ft.Icons.FOLDER
                        icon_color = ft.Colors.AMBER_400
                    elif is_image:
                        icon = ft.Icons.IMAGE
                        icon_color = ft.Colors.GREEN_400
                    elif ext in [".pdf"]:
                        icon = ft.Icons.PICTURE_AS_PDF
                        icon_color = ft.Colors.RED_400
                    elif ext == ".zip":
                        icon = ft.Icons.FOLDER_ZIP
                        icon_color = ORANGE
                    elif ext in [".txt", ".md", ".log"]:
                        icon = ft.Icons.DESCRIPTION
                        icon_color = ft.Colors.BLUE_GREY_400
                    elif ext in [".af", ".afphoto", ".afdesign", ".afpub", ".psd", ".psb", ".svg", ".eps", ".ai"]:
                        icon = ft.Icons.ADOBE
                        icon_color = GREEN
                    else:
                        icon = ft.Icons.INSERT_DRIVE_FILE
                        icon_color = ft.Colors.BLUE_GREY_400

                    checkbox = ft.Checkbox(
                        border_side=ft.BorderSide(color=BLUE),
                        value=file_path in selected_files,
                        on_change=lambda e, path=file_path: on_checkbox_change(e, path),
                    )
                    _checkbox_refs[file_path] = checkbox

                    if is_image:
                        norm_path = os.path.normpath(file_path)
                        cached_b64 = _thumb_cache.get(norm_path) or _image_cache_busters.get(norm_path)
                        _ts = CONSTANTS.DASHBOARD_THUMB_SIZE
                        if cached_b64:
                            visual = ft.Image(
                                src=cached_b64,
                                width=_ts, height=_ts,
                                fit=ft.BoxFit.COVER,
                                border_radius=ft.BorderRadius.all(4),
                            )
                        else:
                            img_ref = ft.Container(
                                bgcolor=GREY,
                                width=_ts, height=_ts,
                                border_radius=ft.BorderRadius.all(4),
                            )
                            visual = img_ref
                            _pending_thumb_refs[norm_path] = (img_ref, file_path, icon, icon_color)
                    else:
                        visual = ft.Icon(icon, color=icon_color, size=22)

                    delete_button = ft.IconButton(
                        icon=ft.Icons.DELETE_OUTLINE, icon_size=18,
                        icon_color=ft.Colors.RED_300, tooltip="Supprimer",
                        on_click=lambda e, path=file_path: delete_item(path),
                        style=ft.ButtonStyle(padding=ft.Padding.all(4)),
                    )
                    rename_button = ft.IconButton(
                        icon=ft.Icons.EDIT_OUTLINED, icon_size=17,
                        icon_color=LIGHT_GREY, tooltip="Renommer",
                        on_click=lambda e, path=file_path: _rename_item(path),
                        style=ft.ButtonStyle(padding=ft.Padding.all(4)),
                    )
                    if is_dir:
                        trailing = ft.Row(
                            [rename_button, delete_button],
                            spacing=0, tight=True,
                            vertical_alignment=ft.CrossAxisAlignment.CENTER,
                        )
                    else:
                        if file_path in _live_print_counts:
                            print_count = _live_print_counts[file_path]
                        else:
                            print_prefix_match = re.match(r'^(\d+)X_', file)
                            print_count = int(print_prefix_match.group(1)) if print_prefix_match else 0
                            _live_print_counts[file_path] = print_count

                        minus_btn = ft.Container(
                            content=ft.Text("−", size=13, color=DARK if print_count > 0 else LIGHT_GREY,
                                            text_align=ft.TextAlign.CENTER, weight=ft.FontWeight.BOLD),
                            width=26, height=26,
                            bgcolor=ORANGE if print_count > 0 else GREY,
                            border_radius=ft.BorderRadius(4, 0, 0, 4),
                            alignment=ft.Alignment(0, 0),
                            on_click=(lambda e, p=file_path: _decrement_print_count(p)) if print_count > 0 else None,
                            ink=print_count > 0,
                            tooltip="Réduire les impressions" if print_count > 0 else "",
                        )
                        _print_minus_btn_refs[file_path] = minus_btn

                        count_text_widget = ft.Text(
                            str(print_count) if print_count > 0 else "·",
                            size=11, color=YELLOW if print_count > 0 else LIGHT_GREY,
                            text_align=ft.TextAlign.CENTER, weight=ft.FontWeight.BOLD,
                        )
                        _print_count_text_refs[file_path] = count_text_widget

                        count_display = ft.Container(
                            content=count_text_widget,
                            width=22, height=26,
                            bgcolor=DARK,
                            alignment=ft.Alignment(0, 0),
                        )
                        plus_btn = ft.Container(
                            content=ft.Text("+", size=13, color=DARK,
                                            text_align=ft.TextAlign.CENTER, weight=ft.FontWeight.BOLD),
                            width=26, height=26,
                            bgcolor=GREEN,
                            border_radius=ft.BorderRadius(0, 4, 4, 0),
                            alignment=ft.Alignment(0, 0),
                            on_click=lambda e, p=file_path: _increment_print_count(p),
                            ink=True,
                            tooltip="Augmenter les impressions",
                        )
                        print_controls = ft.Row(
                            [minus_btn, count_display, plus_btn],
                            spacing=0, tight=True,
                        )
                        trailing = ft.Row(
                            [print_controls, rename_button, delete_button],
                            spacing=4, tight=True,
                            vertical_alignment=ft.CrossAxisAlignment.CENTER,
                        )

                    def _create_right_click_handler(fp, d):
                        def _on_right_click(e):
                            if d:
                                return  # pas de menu pour les dossiers
                            if fp in selected_files and len(selected_files) > 1:
                                files_to_open = list(selected_files)
                            else:
                                files_to_open = [fp]
                            _show_file_context_menu(files_to_open)
                        return _on_right_click

                    new_controls.append(
                        ft.GestureDetector(
                            on_secondary_tap_up=_create_right_click_handler(file_path, is_dir),
                            content=ft.ListTile(
                                leading=checkbox,
                                title=ft.Row(
                                    [visual, ft.Text(file, size=16, color=WHITE, expand=True)],
                                    spacing=12,
                                    vertical_alignment=ft.CrossAxisAlignment.CENTER,
                                ),
                                trailing=trailing,
                                on_click=lambda e, path=file_path, d=is_dir: on_file_click(path, d),
                                hover_color=GREY, dense=False,
                                content_padding=ft.Padding(left=8, top=2, right=8, bottom=2),
                                min_leading_width=0,
                            ),
                        )
                    )

            preview_list.controls.clear()
            preview_list.controls.extend(new_controls)
            preview_list.on_scroll = None
            preview_loading.visible = False



            # Contrôles de pagination
            if total > PAGE_SIZE:
                prev_page_btn.visible = current_pg > 0
                next_page_btn.visible = current_pg < total_pages - 1
                page_indicator_text.value = f"{current_pg * PAGE_SIZE + 1}-{min((current_pg + 1) * PAGE_SIZE, total)}/{total}"
            else:
                prev_page_btn.visible = False
                next_page_btn.visible = False
                page_indicator_text.value = ""



            _update_select_toggle_button()
            page.update()
            # Lancer le chargement des miniatures en arrière-plan après le rendu initial
            if _pending_thumb_refs:
                _start_thumb_loader()
        except Exception as ex:
            log_to_terminal(f"[ERREUR] Rendu preview: {ex}", RED)



    def refresh_preview(reset_page=True, force_reload=False):
        """
        Déclenche un rafraîchissement asynchrone du panneau de prévisualisation.

        Incrémente le jeton de rafraîchissement (annulant tout rafraîchissement
        précédent en cours), vide la liste courante, affiche un indicateur de
        chargement, puis lance un thread de fond ``_bg()`` qui scanne le dossier
        courant et envoie les nouveaux contrôles via PubSub.

        Parameters
        ----------
        reset_page : bool
            Si True (défaut), revient à la page 0. Si False, conserve la page courante.
        force_reload : bool
            Si True, vérifie les mtimes et régénère les caches miniatures depuis le disque.
            Ne passer True que sur le bouton Rafraîchir explicite ou après une rotation.
        """
        if reset_page:
            preview_page["value"] = 0
        preview_refresh_token["value"] += 1
        current_refresh_token = preview_refresh_token["value"]
        preview_list.on_scroll = None
        preview_list.controls.clear()
        file_count_text.value = ""
        folder_to_display = current_browse_folder["path"] or selected_folder["path"]
        preview_loading.visible = bool(folder_to_display)
        page.update()

        def _background_folder_scan():
            """
            Thread de fond : scanne le dossier et stocke des tuples de données brutes.

            Ne crée aucun widget Flet — la construction des contrôles est déléguée
            à _render_current_page() (thread UI), ce qui évite de sérialiser des
            centaines de widgets en un seul page.update().
            """
            entries_data = []
            new_file_count_text = ""
            error_text = ""

            if folder_to_display:
                try:
                    with os.scandir(folder_to_display) as directory_scanner:
                        raw_entries = [dir_entry for dir_entry in directory_scanner if not _is_os_junk(dir_entry)]

                    file_count = sum(1 for dir_entry in raw_entries if not dir_entry.name.startswith(".") and not dir_entry.is_dir())
                    new_file_count_text = f"({file_count} fichier{'s' if file_count > 1 else ''})"

                    if raw_entries:
                        if sort_mode["value"] == 2:
                            sorted_entries = sorted(raw_entries, key=lambda entry: (not entry.is_dir(), -entry.stat().st_mtime))
                        elif sort_mode["value"] == 1:
                            sorted_entries = sorted(raw_entries, key=lambda entry: (not entry.is_dir(), entry.name.lower()), reverse=True)
                        else:
                            sorted_entries = sorted(raw_entries, key=lambda entry: (not entry.is_dir(), entry.name.lower()))

                        for entry in sorted_entries:
                            name = entry.name
                            path = entry.path
                            is_dir = entry.is_dir()
                            ext = os.path.splitext(name)[1].lower()
                            is_image = ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".ico", ".tiff", ".tif"]
                            entries_data.append((name, path, is_dir, is_image, ext))

                            # Mémoriser le mtime à chaque scan pour détecter les modifications externes.
                            # Sur rafraîchissement explicite (force_reload), invalider systématiquement
                            # les miniatures en mémoire pour refléter les remplacements de fichiers
                            # sous le même nom dès le prochain rendu.
                            if is_image and not is_dir:
                                normalized_path = os.path.normpath(path)
                                try:
                                    current_mtime = entry.stat().st_mtime
                                except OSError:
                                    continue
                                if force_reload:
                                    # Invalider le cache visuel local puis régénérer.
                                    _thumb_cache.pop(normalized_path, None)
                                    _image_cache_busters.pop(normalized_path, None)
                                    new_b64 = thumb_cache.get_or_generate(path)
                                    if new_b64:
                                        _image_cache_busters[normalized_path] = new_b64
                                        _thumb_cache[normalized_path] = new_b64
                                _image_last_mtime[normalized_path] = current_mtime

                except PermissionError:
                    error_text = "⚠️ Accès refusé à ce dossier"
                except Exception as ex:
                    error_text = f"⚠️ Erreur: {str(ex)}"

            # Envoyer tuples de données brutes — le rendu des widgets se fait sur le thread UI
            page.pubsub.send_all_on_topic("preview_ready", (current_refresh_token, entries_data, new_file_count_text, error_text))

        threading.Thread(target=_background_folder_scan, daemon=True).start()
    


    def on_sort_change(e):
        """Change le mode de tri et rafraîchit la preview"""
        sort_mode["value"] = e.control.selected_index
        refresh_preview()



    def go_to_page(delta):
        """Navigue de ±1 page dans la preview sans rescanner le dossier."""
        entries = all_entries_data["list"]
        total_pages = max(1, (len(entries) + PAGE_SIZE - 1) // PAGE_SIZE)
        new_pg = max(0, min(preview_page["value"] + delta, total_pages - 1))
        if new_pg == preview_page["value"]:
            return
        preview_page["value"] = new_pg
        _render_preview_page()
        async def _scroll_top():
            await preview_list.scroll_to(offset=0, duration=0)
        page.run_task(_scroll_top)



    # ================================================================ #
    #                LANCEMENT D'APPLICATIONS                          #
    # ================================================================ #
    def _ask_text_before_launch(dialog_title: str, field_label: str, field_hint: str,
                                app_name: str, app_path: str, is_local: bool):
        """Affiche un dialog de saisie texte, puis relance launch_app avec le texte saisi.

        Utilisé pour recueillir un paramètre (nom de série, nom de PDF) avant
        de lancer un script qui en a besoin.
        """
        text_input = ft.TextField(
            label=field_label,
            hint_text=field_hint,
            autofocus=True,
            width=320,
            bgcolor=DARK,
            border_color=GREY,
        )

        def _on_confirm(e):
            name = text_input.value.strip() if text_input.value else ""
            param_dialog.open = False
            page.update()
            launch_app(app_name, app_path, is_local, series_name=name)

        def _on_cancel(e):
            param_dialog.open = False
            page.update()

        text_input.on_submit = _on_confirm
        param_dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text(dialog_title),
            content=text_input,
            actions=[
                ft.TextButton("Annuler", on_click=_on_cancel),
                ft.TextButton("OK", on_click=_on_confirm),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        page.overlay.append(param_dialog)
        param_dialog.open = True
        page.update()



    def launch_app(app_name, app_path, is_local, series_name=None):
        """
        Lance un script Data/ en tant que sous-processus Python avec les variables
        d'environnement appropriées.

        Pour certains scripts (``Renommer séquence.py``, ``Images en PDF.py``),
        affiche d'abord une boîte de dialogue pour recueillir un paramètre
        (nom de série / nom du PDF) avant de relancer la fonction.

        Pour les scripts « locaux » (is_local=True, ex. Kiosk, Transfert),
        lit stdout/stderr en temps réel, interprète les commandes spéciales
        ``NAVIGATE_TO:<path>`` et filtre les messages de fermeture Flet.
        Pour les scripts « dossier » (is_local=False), injecte ``FOLDER_PATH``,
        ``SELECTED_FILES``, ``RESIZE_SIZE``, etc. et rafraîchit la preview
        à la fin du processus.

        Parameters
        ----------
        app_name : str
            Nom du fichier script (ex. ``"Recadrage manuel.pyw"``).
        app_path : str
            Chemin absolu vers le script.
        is_local : bool
            ``True`` si le script fonctionne sans dossier utilisateur sélectionné
            (ex. Kiosk, Transfert vers TEMP).
        series_name : str or None, optional
            Paramètre textuel à transmettre via ``SERIES_NAME`` ou ``PDF_NAME``
            (renseigné par le dialog affiché au premier appel).
        """



        # Pour Renommer séquence.py, demander le nom de la série avant de lancer
        if app_name == "Copyright.py" and series_name is None:
            copyright_custom_field = ft.TextField(
                prefix="© ",
                border_color=BLUE,
                text_size=13,
                height=40,
                content_padding=ft.Padding(8, 4, 8, 4),
                expand=True,
                visible=False,
            )
            copyright_custom_row = ft.AnimatedSwitcher(
                content=ft.Container(height=0),
                duration=150,
            )

            selected_copyright_mode = {"value": None}

            def _make_copyright_option(mode, icon, label, description):
                def _on_click(e):
                    selected_copyright_mode["value"] = mode
                    copyright_custom_field.visible = (mode == "custom")
                    # Highlight sélectionné
                    for btn in copyright_options_row.controls:
                        btn.border = ft.Border.all(2, BLUE if btn.data == mode else GREY)
                    page.update()
                btn = ft.Container(
                    data=mode,
                    content=ft.Column([
                        ft.Icon(icon, size=22, color=BLUE),
                        ft.Text(label, size=12, color=WHITE, text_align=ft.TextAlign.CENTER, weight=ft.FontWeight.BOLD),
                        ft.Text(description, size=10, color=LIGHT_GREY, text_align=ft.TextAlign.CENTER),
                    ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=4, tight=True),
                    bgcolor=DARK,
                    border=ft.Border.all(2, GREY),
                    border_radius=8,
                    padding=ft.Padding(10, 10, 10, 10),
                    width=105,
                    height=100,
                    on_click=_on_click,
                    ink=True,
                )
                return btn

            copyright_options_row = ft.Row(
                [
                    _make_copyright_option("date",   ft.Icons.CALENDAR_TODAY,    "Date",          "de prise de vue"),
                    _make_copyright_option("filename", ft.Icons.INSERT_DRIVE_FILE, "Nom du fichier",   ""),
                    _make_copyright_option("custom",  ft.Icons.EDIT,              "Personnalisé",  "© Votre nom"),
                ],
                spacing=8,
                alignment=ft.MainAxisAlignment.CENTER,
            )

            def _on_confirm_copyright(e):
                mode = selected_copyright_mode["value"]
                if mode is None:
                    return
                custom_text = copyright_custom_field.value.strip() if mode == "custom" else ""
                if mode == "custom" and not custom_text:
                    copyright_custom_field.error_text = "Requis"
                    page.update()
                    return
                copyright_dlg.open = False
                page.update()
                launch_app(app_name, app_path, is_local, series_name=f"{mode}:© {custom_text}")

            def _on_cancel_copyright(e):
                copyright_dlg.open = False
                page.update()

            copyright_dlg = ft.AlertDialog(
                modal=True,
                title=ft.Text("Ajouter Copyright"),
                content=ft.Column([
                    ft.Text("Quel texte afficher sur les images ?", size=13, color=LIGHT_GREY),
                    ft.Container(height=4),
                    copyright_options_row,
                    ft.Container(height=4),
                    copyright_custom_field,
                ], tight=True, spacing=4, width=350),
                actions=[
                    ft.TextButton("Annuler", on_click=_on_cancel_copyright),
                    ft.TextButton("Appliquer", on_click=_on_confirm_copyright),
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )
            page.overlay.append(copyright_dlg)
            copyright_dlg.open = True
            page.update()
            return

        if app_name == "Images en PDF.py" and series_name is None:
            _ask_text_before_launch("Nom du PDF", "Nom du PDF", "Ex: Album_Mariage",
                                    app_name, app_path, is_local)
            return

        if app_name == "2 en 1.py" and series_name is None:
            _TWO_IN_ONE_FORMATS = CONSTANTS.TWO_IN_ONE_FORMATS
            two_in_one_dialog = ft.AlertDialog(
                modal=True,
                title=ft.Text("Format 2 en 1"),
                content=None,
            )

            def _pick_two_in_one(val):
                two_in_one_dialog.open = False
                page.update()
                launch_app(app_name, app_path, is_local, series_name=val)

            def on_cancel_two_in_one(e):
                two_in_one_dialog.open = False
                page.update()

            two_in_one_buttons = ft.Column(
                controls=[
                    ft.Container(
                        content=ft.Text(label, size=14, color=HOVER_YELLOW, text_align=ft.TextAlign.CENTER),
                        bgcolor=GREY,
                        border=ft.Border.all(1, HOVER_YELLOW),
                        border_radius=4,
                        padding=ft.Padding(12, 10, 12, 10),
                        width=280,
                        alignment=ft.Alignment(0, 0),
                        ink=True,
                        on_click=lambda e, v=val: _pick_two_in_one(v),
                    )
                    for label, val in _TWO_IN_ONE_FORMATS
                ],
                spacing=6,
                tight=True,
            )
            two_in_one_dialog.content = two_in_one_buttons
            two_in_one_dialog.actions = [ft.TextButton("Annuler", on_click=on_cancel_two_in_one)]
            two_in_one_dialog.actions_alignment = ft.MainAxisAlignment.END
            page.overlay.append(two_in_one_dialog)
            two_in_one_dialog.open = True
            page.update()
            return
        
        if app_name == "Fit 203.py" and series_name is None:
            _FIT_203_FORMATS = CONSTANTS.FIT_203_FORMATS
            fit_203_dialog = ft.AlertDialog(
                modal=True,
                title=ft.Text("Format Fit 203"),
                content=None,
            )

            def _pick_fit_203(val):
                fit_203_dialog.open = False
                page.update()
                launch_app(app_name, app_path, is_local, series_name=val)

            def on_cancel_fit_203(e):
                fit_203_dialog.open = False
                page.update()

            fit_203_buttons = ft.Column(
                controls=[
                    ft.Container(
                        content=ft.Text(label, size=14, color=HOVER_YELLOW, text_align=ft.TextAlign.CENTER),
                        bgcolor=GREY,
                        border=ft.Border.all(1, HOVER_YELLOW),
                        border_radius=4,
                        padding=ft.Padding(12, 10, 12, 10),
                        width=280,
                        alignment=ft.Alignment(0, 0),
                        ink=True,
                        on_click=lambda e, v=f"{crop}|{canvas}|{label.split(' sur ')[-1].replace('×', 'x')}": _pick_fit_203(v),
                    )
                    for label, crop, canvas in _FIT_203_FORMATS
                ],
                spacing=6,
                tight=True,
            )
            fit_203_dialog.content = fit_203_buttons
            fit_203_dialog.actions = [ft.TextButton("Annuler", on_click=on_cancel_fit_203)]
            fit_203_dialog.actions_alignment = ft.MainAxisAlignment.END
            page.overlay.append(fit_203_dialog)
            fit_203_dialog.open = True
            page.update()
            return

        if app_name == "Recadrage automatique.py" and series_name is None:
            _format_items = list(CONSTANTS.FORMATS.items())
            _default_format = "10x15" if "10x15" in CONSTANTS.FORMATS else (_format_items[0][0] if _format_items else "10x15")
            _default_mm = CONSTANTS.FORMATS.get(_default_format, (102, 152))
            _auto_scope_value = "selected" if selected_files else "all"

            force_mode_state = {"manual": False}

            force_format_dropdown = ft.Dropdown(
                value=_default_format,
                options=[ft.dropdown.Option(name) for name, _ in _format_items],
                width=260,
                text_size=13,
                border_color=GREEN,
                bgcolor=DARK,
            )
            force_manual_width = ft.TextField(
                label="Largeur (mm)",
                value=str(_default_mm[0]),
                width=125,
                text_size=13,
                keyboard_type=ft.KeyboardType.NUMBER,
                border_color=GREEN,
                bgcolor=DARK,
                disabled=True,
            )
            force_manual_height = ft.TextField(
                label="Hauteur (mm)",
                value=str(_default_mm[1]),
                width=125,
                text_size=13,
                keyboard_type=ft.KeyboardType.NUMBER,
                border_color=GREEN,
                bgcolor=DARK,
                disabled=True,
            )
            force_manual_switch = ft.Switch(
                label="Saisie manuelle (mm)",
                value=False,
                active_color=GREEN,
            )
            force_scope_info = ft.Text(
                "Portée auto : sélection en cours" if _auto_scope_value == "selected" else "Portée auto : tout le dossier",
                size=12,
                color=LIGHT_GREY,
                text_align=ft.TextAlign.CENTER,
            )
            force_error_text = ft.Text("", size=12, color=RED, text_align=ft.TextAlign.CENTER)

            force_fit_switch = ft.Switch(
                label="Fit 100% (sans rognage)",
                value=False,
                active_color=GREEN,
            )
            force_white_border_switch = ft.Switch(
                label="Bord blanc 5mm",
                value=False,
                active_color=GREEN,
            )

            def _update_force_mode_ui():
                _manual = force_mode_state["manual"]
                force_format_dropdown.disabled = _manual
                force_manual_width.disabled = not _manual
                force_manual_height.disabled = not _manual
                page.update()

            def _on_force_switch_change(e):
                force_mode_state["manual"] = bool(e.control.value)
                _update_force_mode_ui()

            def _on_force_confirm(e):
                try:
                    if force_mode_state["manual"]:
                        _w = int((force_manual_width.value or "").strip())
                        _h = int((force_manual_height.value or "").strip())
                        if _w <= 0 or _h <= 0:
                            raise ValueError()
                        _size_value = f"{_w}x{_h}"
                    else:
                        _fmt_name = (force_format_dropdown.value or "").strip()
                        _fmt_mm = CONSTANTS.FORMATS.get(_fmt_name)
                        if not _fmt_mm:
                            raise ValueError()
                        _size_value = f"{_fmt_mm[0]}x{_fmt_mm[1]}"

                    _scope_value = "selected" if selected_files else "all"
                    _fit_value = "1" if force_fit_switch.value else "0"
                    _wb_value = "1" if force_white_border_switch.value else "0"

                    force_crop_dialog.open = False
                    page.update()
                    launch_app(app_name, app_path, is_local, series_name=f"{_size_value}|{_scope_value}|{_fit_value}|{_wb_value}")
                except Exception:
                    force_error_text.value = "Taille invalide. Utilise des entiers en mm (ex: 102 et 152)."
                    page.update()

            def _on_force_cancel(e):
                force_crop_dialog.open = False
                page.update()

            force_manual_switch.on_change = _on_force_switch_change

            force_crop_dialog = ft.AlertDialog(
                modal=True,
                title=ft.Text("Recadrage automatique - format", text_align=ft.TextAlign.CENTER, color=GREEN),
                content=ft.Column(
                    [
                        ft.Text(
                            "Choisis un format (CONSTANTS) ou saisis une taille manuelle en mm.",
                            size=13,
                            color=LIGHT_GREY,
                            text_align=ft.TextAlign.CENTER,
                        ),
                        force_format_dropdown,
                        force_manual_switch,
                        ft.Row(
                            [force_manual_width, force_manual_height],
                            spacing=8,
                            tight=True,
                            alignment=ft.MainAxisAlignment.CENTER,
                        ),
                        force_fit_switch,
                        force_white_border_switch,
                        force_scope_info,
                        force_error_text,
                    ],
                    tight=True,
                    spacing=8,
                    width=380,
                    horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                ),
                actions=[
                    ft.TextButton("Annuler", on_click=_on_force_cancel, style=ft.ButtonStyle(color=GREEN)),
                    ft.TextButton("Lancer", on_click=_on_force_confirm, style=ft.ButtonStyle(color=GREEN)),
                ],
                actions_alignment=ft.MainAxisAlignment.CENTER,
            )
            page.overlay.append(force_crop_dialog)
            force_crop_dialog.open = True
            _update_force_mode_ui()
            return

        if app_name == "Renommer séquence.py" and series_name is None:
            _ask_text_before_launch("Renommer la série", "Nom de la série", "Ex: Mariage_Martin",
                                    app_name, app_path, is_local)
            return

        if app_name == "Débruiter.py" and series_name is None:
            _dn_fields = {
                "h": ft.TextField(
                    label="Force luminance (h)",
                    value=str(CONSTANTS.DENOISE_H),
                    hint_text="1 léger · 4 standard · 10 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "h_color": ft.TextField(
                    label="Force couleur (hColor)",
                    value=str(CONSTANTS.DENOISE_H_COLOR),
                    hint_text="1 léger · 2 standard · 6 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "template": ft.TextField(
                    label="Fenêtre comparaison (impair)",
                    value=str(CONSTANTS.DENOISE_TEMPLATE_WINDOW),
                    hint_text="5 · 7 standard · 11",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "search": ft.TextField(
                    label="Fenêtre recherche (impair)",
                    value=str(CONSTANTS.DENOISE_SEARCH_WINDOW),
                    hint_text="11 rapide · 21 standard · 35 lent",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _dn_error = ft.Text("", size=12, color=RED, text_align=ft.TextAlign.CENTER)

            def _on_dn_confirm(e):
                try:
                    h    = int((_dn_fields["h"].value or "").strip())
                    hc   = int((_dn_fields["h_color"].value or "").strip())
                    tmpl = int((_dn_fields["template"].value or "").strip())
                    srch = int((_dn_fields["search"].value or "").strip())
                    if h <= 0 or hc <= 0 or tmpl <= 0 or srch <= 0:
                        raise ValueError()
                    if tmpl % 2 == 0 or srch % 2 == 0:
                        _dn_error.value = "Les fenêtres doivent être impaires."
                        page.update()
                        return
                except Exception:
                    _dn_error.value = "Valeurs invalides — entiers positifs impairs requis."
                    page.update()
                    return
                token = f"{h}|{hc}|{tmpl}|{srch}"
                _dn_dlg.open = False
                page.update()
                launch_app(app_name, app_path, is_local, series_name=token)

            def _on_dn_cancel(e):
                _dn_dlg.open = False
                page.update()

            _dn_dlg = ft.AlertDialog(
                modal=True,
                title=ft.Text("Débruiter — paramètres NLM", text_align=ft.TextAlign.CENTER, color=BLUE),
                content=ft.Column(
                    [
                        ft.Text(
                            "Les valeurs par défaut viennent de CONSTANTS.py (section 12.1).",
                            size=11, color=LIGHT_GREY, text_align=ft.TextAlign.CENTER,
                        ),
                        ft.Row([_dn_fields["h"], _dn_fields["h_color"]], spacing=8),
                        ft.Row([_dn_fields["template"], _dn_fields["search"]], spacing=8),
                        _dn_error,
                    ],
                    tight=True,
                    spacing=10,
                    width=420,
                    horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                ),
                actions=[
                    ft.TextButton("Annuler", on_click=_on_dn_cancel, style=ft.ButtonStyle(color=BLUE)),
                    ft.TextButton("Lancer", on_click=_on_dn_confirm, style=ft.ButtonStyle(color=BLUE)),
                ],
                actions_alignment=ft.MainAxisAlignment.CENTER,
            )
            page.overlay.append(_dn_dlg)
            _dn_dlg.open = True
            page.update()
            return

        if app_name == "Grain pellicule.py" and series_name is None:

            def _grain_group(label, color, amount_val, size_val, color_val, shadow_val, chroma_shift_val):
                """Retourne (container, dict_of_fields) pour un groupe de 5 champs."""
                fields = {
                    "amount": ft.TextField(
                        label="Intensité (amount)",
                        value=str(amount_val),
                        hint_text="0.03 fin · 0.10 ISO 400 · 0.20 ISO 1600",
                        text_size=12,
                        keyboard_type=ft.KeyboardType.NUMBER,
                        border=ft.InputBorder.OUTLINE,
                        border_color=color,
                        focused_border_color=color,
                        bgcolor=DARK,
                        height=56,
                        expand=True,
                    ),
                    "size": ft.TextField(
                        label="Taille (% min. dim.)",
                        value=str(size_val),
                        hint_text="0.1 fin · 0.3 moyen · 0.6 gros",
                        text_size=12,
                        keyboard_type=ft.KeyboardType.NUMBER,
                        border=ft.InputBorder.OUTLINE,
                        border_color=color,
                        focused_border_color=color,
                        bgcolor=DARK,
                        height=56,
                        expand=True,
                    ),
                    "color": ft.TextField(
                        label="Couleur (color_ratio)",
                        value=str(color_val),
                        hint_text="0.0 mono · 0.3 subtil · 1.0 plein",
                        text_size=12,
                        keyboard_type=ft.KeyboardType.NUMBER,
                        border=ft.InputBorder.OUTLINE,
                        border_color=color,
                        focused_border_color=color,
                        bgcolor=DARK,
                        height=56,
                        expand=True,
                    ),
                    "shadow": ft.TextField(
                        label="Concentration mi-tons",
                        value=str(shadow_val),
                        hint_text="1.0 large · 2.0 centré · 3.0 serré",
                        text_size=12,
                        keyboard_type=ft.KeyboardType.NUMBER,
                        border=ft.InputBorder.OUTLINE,
                        border_color=color,
                        focused_border_color=color,
                        bgcolor=DARK,
                        height=56,
                        expand=True,
                    ),
                    "chroma_shift": ft.TextField(
                        label="Décalage inter-canal (%)",
                        value=str(chroma_shift_val),
                        hint_text="0 = désactivé · 0.1 subtil · 0.3 prononcé",
                        text_size=12,
                        keyboard_type=ft.KeyboardType.NUMBER,
                        border=ft.InputBorder.OUTLINE,
                        border_color=color,
                        focused_border_color=color,
                        bgcolor=DARK,
                        height=56,
                        expand=True,
                    ),
                }
                container = ft.Container(
                    content=ft.Column(
                        [
                            ft.Text(label, size=12, color=color, weight=ft.FontWeight.BOLD),
                            ft.Row([fields["amount"], fields["size"]], spacing=8),
                            ft.Row([fields["color"], fields["shadow"]], spacing=8),
                            ft.Row([fields["chroma_shift"]], spacing=8),
                        ],
                        spacing=6,
                        tight=True,
                    ),
                    border=ft.Border.all(1, color),
                    border_radius=6,
                    padding=ft.Padding(10, 8, 10, 8),
                )
                return container, fields

            _grain1_enabled = {"value": True}
            _g1_container, _g1 = _grain_group(
                "Couche 1",
                ORANGE,
                CONSTANTS.GRAIN_AMOUNT,
                CONSTANTS.GRAIN_SIZE,
                CONSTANTS.GRAIN_COLOR_RATIO,
                CONSTANTS.GRAIN_SHADOW_BOOST,
                CONSTANTS.GRAIN_CHROMA_SHIFT,
            )
            _grain1_switch = ft.Switch(value=True, active_color=ORANGE)

            def _on_grain1_toggle(e):
                _grain1_enabled["value"] = bool(e.control.value)
                _g1_container.opacity = 1.0 if _grain1_enabled["value"] else 0.3
                for f in _g1.values():
                    f.disabled = not _grain1_enabled["value"]
                page.update()

            _grain1_switch.on_change = _on_grain1_toggle

            _grain2_enabled = {"value": True}
            _g2_container, _g2 = _grain_group(
                "Couche 2",
                LIGHT_GREY,
                CONSTANTS.GRAIN2_AMOUNT,
                CONSTANTS.GRAIN2_SIZE,
                CONSTANTS.GRAIN2_COLOR_RATIO,
                CONSTANTS.GRAIN2_SHADOW_BOOST,
                CONSTANTS.GRAIN2_CHROMA_SHIFT,
            )

            _grain2_switch = ft.Switch(
                value=True,
                active_color=ORANGE,
            )

            def _on_grain2_toggle(e):
                _grain2_enabled["value"] = bool(e.control.value)
                _g2_container.opacity = 1.0 if _grain2_enabled["value"] else 0.3
                for f in _g2.values():
                    f.disabled = not _grain2_enabled["value"]
                page.update()

            _grain2_switch.on_change = _on_grain2_toggle

            # ── Halation ──────────────────────────────────────────────────────
            _halation_enabled = {"value": True}
            _halation_fields = {
                "threshold": ft.TextField(
                    label="Seuil (threshold)",
                    value=str(CONSTANTS.HALATION_THRESHOLD),
                    hint_text="0.55 large · 0.65 standard · 0.80 éclats seuls",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=RED, focused_border_color=RED,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "radius": ft.TextField(
                    label="Rayon (% image)",
                    value=str(CONSTANTS.HALATION_RADIUS),
                    hint_text="1 discret · 5 standard · 15 prononcé",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=RED, focused_border_color=RED,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "intensity": ft.TextField(
                    label="Intensité",
                    value=str(CONSTANTS.HALATION_INTENSITY),
                    hint_text="0.1 léger · 0.6 standard · 1.0 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=RED, focused_border_color=RED,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "red_shift": ft.TextField(
                    label="Décalage rouge",
                    value=str(CONSTANTS.HALATION_RED_SHIFT),
                    hint_text="0.0 neutre · 0.5 chaud · 1.0 rouge vif",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=RED, focused_border_color=RED,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _halation_container = ft.Container(
                content=ft.Column(
                    [
                        ft.Text("Halation", size=12, color=RED, weight=ft.FontWeight.BOLD),
                        ft.Row([_halation_fields["threshold"], _halation_fields["radius"]], spacing=8),
                        ft.Row([_halation_fields["intensity"], _halation_fields["red_shift"]], spacing=8),
                    ],
                    spacing=6, tight=True,
                ),
                border=ft.Border.all(1, RED),
                border_radius=6,
                padding=ft.Padding(10, 8, 10, 8),
            )
            _halation_switch = ft.Switch(value=True, active_color=RED)

            def _on_halation_toggle(e):
                _halation_enabled["value"] = bool(e.control.value)
                _halation_container.opacity = 1.0 if _halation_enabled["value"] else 0.3
                for f in _halation_fields.values():
                    f.disabled = not _halation_enabled["value"]
                page.update()

            _halation_switch.on_change = _on_halation_toggle

            # ── Bloom ─────────────────────────────────────────────────────────
            _bloom_enabled = {"value": True}
            _bloom_fields = {
                "radius": ft.TextField(
                    label="Rayon (% image)",
                    value=str(CONSTANTS.BLOOM_RADIUS),
                    hint_text="2 discret · 10 standard · 20 prononcé",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "intensity": ft.TextField(
                    label="Intensité",
                    value=str(CONSTANTS.BLOOM_INTENSITY),
                    hint_text="0.1 léger · 0.3 standard · 0.6 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=BLUE, focused_border_color=BLUE,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _bloom_container = ft.Container(
                content=ft.Column(
                    [
                        ft.Text("Bloom (Soft Light)", size=12, color=BLUE, weight=ft.FontWeight.BOLD),
                        ft.Row([_bloom_fields["radius"], _bloom_fields["intensity"]], spacing=8),
                    ],
                    spacing=6, tight=True,
                ),
                border=ft.Border.all(1, BLUE),
                border_radius=6,
                padding=ft.Padding(10, 8, 10, 8),
            )
            _bloom_switch = ft.Switch(value=True, active_color=BLUE)

            def _on_bloom_toggle(e):
                _bloom_enabled["value"] = bool(e.control.value)
                _bloom_container.opacity = 1.0 if _bloom_enabled["value"] else 0.3
                for f in _bloom_fields.values():
                    f.disabled = not _bloom_enabled["value"]
                page.update()

            _bloom_switch.on_change = _on_bloom_toggle

            # ── Courbe tonale argentique ───────────────────────────────────────
            _curve_enabled = {"value": True}
            _curve_fields = {
                "shoulder_start": ft.TextField(
                    label="Épaulement — seuil",
                    value=str(CONSTANTS.CURVE_SHOULDER_START),
                    hint_text="0.70 large · 0.80 standard · 0.90 conservateur",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=GREEN, focused_border_color=GREEN,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "shoulder_strength": ft.TextField(
                    label="Épaulement — force",
                    value=str(CONSTANTS.CURVE_SHOULDER_STRENGTH),
                    hint_text="0.2 doux · 0.5 standard · 1.5 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=GREEN, focused_border_color=GREEN,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "toe_start": ft.TextField(
                    label="Pied — seuil",
                    value=str(CONSTANTS.CURVE_TOE_START),
                    hint_text="0.03–0.12 (luma)",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=GREEN, focused_border_color=GREEN,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "toe_lift": ft.TextField(
                    label="Pied — relèvement",
                    value=str(CONSTANTS.CURVE_TOE_LIFT),
                    hint_text="0 = aucun · 0.08 subtil · 0.20 prononcé",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=GREEN, focused_border_color=GREEN,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _curve_container = ft.Container(
                content=ft.Column(
                    [
                        ft.Text("Courbe tonale argentique", size=12, color=GREEN, weight=ft.FontWeight.BOLD),
                        ft.Row([_curve_fields["shoulder_start"], _curve_fields["shoulder_strength"]], spacing=8),
                        ft.Row([_curve_fields["toe_start"],      _curve_fields["toe_lift"]],          spacing=8),
                    ],
                    spacing=6, tight=True,
                ),
                border=ft.Border.all(1, GREEN),
                border_radius=6,
                padding=ft.Padding(10, 8, 10, 8),
            )
            _curve_switch = ft.Switch(value=True, active_color=GREEN)

            def _on_curve_toggle(e):
                _curve_enabled["value"] = bool(e.control.value)
                _curve_container.opacity = 1.0 if _curve_enabled["value"] else 0.3
                for f in _curve_fields.values():
                    f.disabled = not _curve_enabled["value"]
                page.update()

            _curve_switch.on_change = _on_curve_toggle

            # ── Désaturation des extrêmes ──────────────────────────────────────
            _desat_enabled = {"value": True}
            _desat_fields = {
                "shadow_threshold": ft.TextField(
                    label="Seuil ombres",
                    value=str(CONSTANTS.DESAT_SHADOW_THRESHOLD),
                    hint_text="0.15 fort · 0.25 standard · 0.40 large",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=VIOLET, focused_border_color=VIOLET,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "shadow_intensity": ft.TextField(
                    label="Intensité ombres",
                    value=str(CONSTANTS.DESAT_SHADOW_INTENSITY),
                    hint_text="0.3 subtil · 0.6 standard · 1.0 gris pur",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=VIOLET, focused_border_color=VIOLET,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "highlight_threshold": ft.TextField(
                    label="Seuil hautes lumières",
                    value=str(CONSTANTS.DESAT_HIGHLIGHT_THRESHOLD),
                    hint_text="0.75 large · 0.85 standard · 0.95 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=VIOLET, focused_border_color=VIOLET,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "highlight_intensity": ft.TextField(
                    label="Intensité hautes lumières",
                    value=str(CONSTANTS.DESAT_HIGHLIGHT_INTENSITY),
                    hint_text="0.2 subtil · 0.5 standard · 1.0 gris pur",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=VIOLET, focused_border_color=VIOLET,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "midtone_boost": ft.TextField(
                    label="Boost saturation mi-tons",
                    value=str(CONSTANTS.DESAT_MIDTONE_BOOST),
                    hint_text="0 = aucun · 0.15 subtil · 0.30 prononcé",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=VIOLET, focused_border_color=VIOLET,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _desat_container = ft.Container(
                content=ft.Column(
                    [
                        ft.Row([_desat_fields["shadow_threshold"],   _desat_fields["shadow_intensity"]], spacing=8),
                        ft.Row([_desat_fields["highlight_threshold"], _desat_fields["highlight_intensity"]], spacing=8),
                        ft.Row([_desat_fields["midtone_boost"]], spacing=8),
                    ],
                    spacing=6, tight=True,
                ),
                border=ft.Border.all(1, VIOLET),
                border_radius=6,
                padding=ft.Padding(10, 8, 10, 8),
            )
            _desat_switch = ft.Switch(value=True, active_color=VIOLET)

            def _on_desat_toggle(e):
                _desat_enabled["value"] = bool(e.control.value)
                _desat_container.opacity = 1.0 if _desat_enabled["value"] else 0.3
                for f in _desat_fields.values():
                    f.disabled = not _desat_enabled["value"]
                page.update()

            _desat_switch.on_change = _on_desat_toggle

            # ── Aberrations chromatiques ──────────────────────────────────────
            _ca_enabled = {"value": True}
            _ca_fields = {
                "strength": ft.TextField(
                    label="Intensité (% diagonale)",
                    value=str(CONSTANTS.CA_STRENGTH),
                    hint_text="0.3 subtil · 1.0 prononcé · 2.0 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=YELLOW, focused_border_color=YELLOW,
                    bgcolor=DARK, height=56, expand=True,
                ),
                "axial_ratio": ft.TextField(
                    label="Composante axiale",
                    value=str(CONSTANTS.CA_AXIAL_RATIO),
                    hint_text="0 = purement radial · 0.15 subtil · 0.4 fort",
                    text_size=12, keyboard_type=ft.KeyboardType.NUMBER,
                    border=ft.InputBorder.OUTLINE, border_color=YELLOW, focused_border_color=YELLOW,
                    bgcolor=DARK, height=56, expand=True,
                ),
            }
            _ca_container = ft.Container(
                content=ft.Column(
                    [
                        ft.Row([_ca_fields["strength"], _ca_fields["axial_ratio"]], spacing=8),
                    ],
                    spacing=6, tight=True,
                ),
                border=ft.Border.all(1, YELLOW),
                border_radius=6,
                padding=ft.Padding(10, 8, 10, 8),
            )
            _ca_switch = ft.Switch(value=True, active_color=YELLOW)

            def _on_ca_toggle(e):
                _ca_enabled["value"] = bool(e.control.value)
                _ca_container.opacity = 1.0 if _ca_enabled["value"] else 0.3
                for f in _ca_fields.values():
                    f.disabled = not _ca_enabled["value"]
                page.update()

            _ca_switch.on_change = _on_ca_toggle

            _grain_error_text = ft.Text("", size=12, color=RED, text_align=ft.TextAlign.CENTER)

            def _parse_grain_group(fields):
                amount       = float((fields["amount"].value or "").replace(",", ".").strip())
                size         = float((fields["size"].value or "").replace(",", ".").strip())
                color        = float((fields["color"].value or "").replace(",", ".").strip())
                shadow       = float((fields["shadow"].value or "").replace(",", ".").strip())
                chroma_shift = float((fields["chroma_shift"].value or "0").replace(",", ".").strip())
                if amount < 0 or size <= 0 or not (0.0 <= color <= 1.0) or shadow < 0 or chroma_shift < 0:
                    raise ValueError()
                return amount, size, color, shadow, chroma_shift

            def _on_grain_confirm(e):
                try:
                    a1, s1, c1, sh1, cs1 = _parse_grain_group(_g1)
                    a2, s2, c2, sh2, cs2 = _parse_grain_group(_g2)
                    ht   = float((_halation_fields["threshold"].value or "").replace(",", ".").strip())
                    hr   = float((_halation_fields["radius"].value or "").replace(",", ".").strip())
                    hi   = float((_halation_fields["intensity"].value or "").replace(",", ".").strip())
                    hred = float((_halation_fields["red_shift"].value or "").replace(",", ".").strip())
                    br   = float((_bloom_fields["radius"].value or "").replace(",", ".").strip())
                    bi   = float((_bloom_fields["intensity"].value or "").replace(",", ".").strip())
                    dst  = float((_desat_fields["shadow_threshold"].value or "").replace(",", ".").strip())
                    dsi  = float((_desat_fields["shadow_intensity"].value or "").replace(",", ".").strip())
                    dht  = float((_desat_fields["highlight_threshold"].value or "").replace(",", ".").strip())
                    dhi  = float((_desat_fields["highlight_intensity"].value or "").replace(",", ".").strip())
                    dmb  = float((_desat_fields["midtone_boost"].value or "").replace(",", ".").strip())
                    css  = float((_curve_fields["shoulder_start"].value or "").replace(",", ".").strip())
                    cstr = float((_curve_fields["shoulder_strength"].value or "").replace(",", ".").strip())
                    cts  = float((_curve_fields["toe_start"].value or "").replace(",", ".").strip())
                    ctl  = float((_curve_fields["toe_lift"].value or "").replace(",", ".").strip())
                    if not (0.0 <= ht <= 1.0) or hr <= 0 or not (0.0 <= hi <= 1.0) or not (0.0 <= hred <= 1.0):
                        raise ValueError()
                    if br <= 0 or not (0.0 <= bi <= 1.0):
                        raise ValueError()
                    if not (0.0 <= dst <= 1.0) or not (0.0 <= dsi <= 1.0):
                        raise ValueError()
                    if not (0.0 <= dht <= 1.0) or not (0.0 <= dhi <= 1.0):
                        raise ValueError()
                    if dmb < 0:
                        raise ValueError()
                    if not (0.0 <= css <= 1.0) or cstr < 0 or cts < 0 or ctl < 0:
                        raise ValueError()
                    ca_s = float((_ca_fields["strength"].value or "0").replace(",", ".").strip())
                    ca_ax = float((_ca_fields["axial_ratio"].value or "0").replace(",", ".").strip())
                    if ca_s < 0 or ca_ax < 0:
                        raise ValueError()
                except Exception:
                    _grain_error_text.value = "Valeurs invalides. Vérifie les champs."
                    page.update()
                    return
                g1 = 1 if _grain1_enabled["value"] else 0
                g2 = 1 if _grain2_enabled["value"] else 0
                h  = 1 if _halation_enabled["value"] else 0
                b  = 1 if _bloom_enabled["value"] else 0
                d  = 1 if _desat_enabled["value"] else 0
                cv = 1 if _curve_enabled["value"] else 0
                ca = 1 if _ca_enabled["value"] else 0
                token = f"{g1}|{a1}|{s1}|{c1}|{sh1}|{g2}|{a2}|{s2}|{c2}|{sh2}|{h}|{ht}|{hr}|{hi}|{hred}|{b}|{br}|{bi}|{d}|{dst}|{dsi}|{dht}|{dhi}|{cv}|{css}|{cstr}|{cts}|{ctl}|{dmb}|{cs1}|{cs2}|{ca}|{ca_s}|{ca_ax}"
                _grain_dlg.open = False
                page.update()
                launch_app(app_name, app_path, is_local, series_name=token)

            def _on_grain_cancel(e):
                _grain_dlg.open = False
                page.update()

            _grain_dlg = ft.AlertDialog(
                modal=True,
                title=ft.Text("Grain pellicule — paramètres", text_align=ft.TextAlign.CENTER, color=ORANGE),
                content=ft.Column(
                    [
                        ft.Text(
                            "Les valeurs par défaut viennent de CONSTANTS.py (section 12).",
                            size=11,
                            color=LIGHT_GREY,
                            text_align=ft.TextAlign.CENTER,
                        ),
                        ft.ExpansionTile(
                            leading=_grain1_switch,
                            title=ft.Text("Grain — Couche 1", color=ORANGE, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_g1_container],
                        ),
                        ft.ExpansionTile(
                            leading=_grain2_switch,
                            title=ft.Text("Grain — Couche 2", color=ORANGE, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_g2_container],
                        ),
                        ft.ExpansionTile(
                            leading=_halation_switch,
                            title=ft.Text("Halation", color=RED, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_halation_container],
                        ),
                        ft.ExpansionTile(
                            leading=_bloom_switch,
                            title=ft.Text("Bloom (Soft Light)", color=BLUE, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_bloom_container],
                        ),
                        ft.ExpansionTile(
                            leading=_desat_switch,
                            title=ft.Text("Désaturation des extrêmes", color=VIOLET, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_desat_container],
                        ),
                        ft.ExpansionTile(
                            leading=_curve_switch,
                            title=ft.Text("Courbe tonale", color=GREEN, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_curve_container],
                        ),
                        ft.ExpansionTile(
                            leading=_ca_switch,
                            title=ft.Text("Aberrations chromatiques", color=YELLOW, weight=ft.FontWeight.BOLD, size=13),
                            expanded=False,
                            maintain_state=True,
                            tile_padding=ft.Padding(8, 0, 8, 0),
                            controls=[_ca_container],
                        ),
                        _grain_error_text,
                    ],
                    tight=True,
                    spacing=4,
                    width=500,
                    height=460,
                    scroll=ft.ScrollMode.AUTO,
                    horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                ),
                actions=[
                    ft.TextButton("Annuler", on_click=_on_grain_cancel, style=ft.ButtonStyle(color=ORANGE)),
                    ft.TextButton("Lancer", on_click=_on_grain_confirm, style=ft.ButtonStyle(color=ORANGE)),
                ],
                actions_alignment=ft.MainAxisAlignment.CENTER,
            )
            page.overlay.append(_grain_dlg)
            _grain_dlg.open = True
            page.update()
            return

        if not is_local and not (current_browse_folder["path"] or selected_folder["path"]):
            log_to_terminal("[ERREUR] Veuillez sélectionner un dossier avant de lancer cette application", RED)
            return

        try:
            display_name = app_name[:-4] if app_name.endswith(".pyw") else app_name[:-3]
            log_to_terminal(f"▶ Lancement de {display_name}...", BLUE)
            
            if is_local:
                # Préparer l'environnement pour les apps locales
                env = os.environ.copy()
                env["PYTHONIOENCODING"] = "utf-8"
                env["DATA_PATH"] = os.path.join(app_directory, "Data")
                
                # Naviguer vers le PATH pour order_it gauche/droite (après fin du processus)
                kiosk_target_path = None
                if app_name == "Kiosk gauche.py":
                    kiosk_target_path = CONSTANTS.KIOSK_GAUCHE_DEST
                    if not os.path.isdir(kiosk_target_path):
                        log_to_terminal(f"[AVERTISSEMENT] Le dossier {kiosk_target_path} n'est pas accessible", ORANGE)
                        kiosk_target_path = None

                elif app_name == "Kiosk droite.py":
                    kiosk_target_path = CONSTANTS.KIOSK_DROITE_DEST
                    if not os.path.isdir(kiosk_target_path):
                        log_to_terminal(f"[AVERTISSEMENT] Le dossier {kiosk_target_path} n'est pas accessible", ORANGE)
                        kiosk_target_path = None



                # Ajouter le dossier destination pour Transfert vers TEMP.py
                if app_name == "Transfert vers TEMP.py":
                    if platform.system() == "Windows":
                        env["DEST_FOLDER"] = "\\\\diskstation\\travaux en cours\\Z2026\\TEMP"
                    else:
                        env["DEST_FOLDER"] = "/Volumes/TRAVAUX EN COURS/Z2026/TEMP"
                    env["LAUNCHED_FROM_DASHBOARD"] = "1"
                    if selected_files:
                        env["SOURCE_FILES"] = "|".join(str(f) for f in selected_files)
                        
                        # Afficher un dialog de confirmation de suppression AVANT de lancer
                        def _launch_transfer_with_delete_choice(delete_after: bool):
                            env["DELETE_AFTER_TRANSFER"] = "1" if delete_after else "0"
                            transfer_confirm_dialog.open = False
                            page.update()
                            
                            # Lancer le subprocess maintenant
                            process = subprocess.Popen(
                                [sys.executable, "-u", app_path],
                                env=env,
                                stdout=subprocess.PIPE,
                                stderr=subprocess.PIPE,
                                text=True,
                                encoding="utf-8",
                                errors="replace",
                                bufsize=1,
                                universal_newlines=True
                            )
                            
                            # Minimise Dashboard pour laisser la place à l'app lancée
                            if app_path.endswith(".pyw"):
                                page.window.minimized = True
                                page.update()

                                def _watch_local(proc=process):
                                    proc.wait()
                                    page.window.minimized = False
                                    page.window.maximized = True
                                    page.run_task(page.window.to_front)
                                    page.update()

                                threading.Thread(target=_watch_local, daemon=True).start()
                            
                            # Lire la sortie en temps réel (même code que le subprocess normal)
                            def read_output(pipe, color):
                                for line in iter(pipe.readline, ''):
                                    if not line:
                                        break
                                    if 'Session closed by remote host' in line or 'Segmentation fault' in line:
                                        continue
                                    if line.startswith("NAVIGATE_TO:"):
                                        nav_path = line.replace("NAVIGATE_TO:", "").strip()
                                        if os.path.isdir(nav_path):
                                            page.pubsub.send_all_on_topic("navigate", nav_path)
                                    else:
                                        log_to_terminal(line.rstrip(), color)
                            
                            threading.Thread(target=read_output, args=(process.stdout, GREEN), daemon=True).start()
                            threading.Thread(target=read_output, args=(process.stderr, RED), daemon=True).start()
                            return
                        
                        def _confirm_delete(e):
                            _launch_transfer_with_delete_choice(True)
                        
                        def _skip_delete(e):
                            _launch_transfer_with_delete_choice(False)
                        
                        transfer_confirm_dialog = ft.AlertDialog(
                            title=ft.Text("Supprimer les fichiers après transfert ?"),
                            content=ft.Text(
                                f"{len(selected_files)} fichier(s) sélectionné(s) seront transférés.\n\n"
                                "Supprimer les fichiers source après la copie réussie ?"
                            ),
                            actions=[
                                ft.TextButton("Conserver", on_click=_skip_delete),
                                ft.TextButton("Supprimer", on_click=_confirm_delete, style=ft.ButtonStyle(color=ft.Colors.RED)),
                            ],
                        )
                        page.overlay.append(transfer_confirm_dialog)
                        transfer_confirm_dialog.open = True
                        page.update()
                        return
                
                process = subprocess.Popen(
                    [sys.executable, "-u", app_path],
                    env=env,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    bufsize=1,
                    universal_newlines=True
                )

                # Minimise Dashboard pour laisser la place à l'app lancée
                if app_path.endswith(".pyw"):
                    page.window.minimized = True
                    page.update()

                    def _watch_local(proc=process, nav_path=kiosk_target_path):
                        proc.wait()
                        page.window.minimized = False
                        page.window.maximized = True
                        page.run_task(page.window.to_front)
                        page.update()
                        if nav_path and os.path.isdir(nav_path):
                            navigate_to_folder(nav_path)

                    threading.Thread(target=_watch_local, daemon=True).start()
                elif kiosk_target_path:
                    def _watch_kiosk(proc=process, nav_path=kiosk_target_path):
                        proc.wait()
                        if os.path.isdir(nav_path):
                            navigate_to_folder(nav_path)

                    threading.Thread(target=_watch_kiosk, daemon=True).start()



                # Lire la sortie en temps réel
                def read_output(pipe, color):
                    """
                    Lit en temps réel une pipe stdout ou stderr du sous-processus
                    local (Flet) et envoie chaque ligne non vide au terminal.

                    Filtre les messages de fermeture de session Flet et interprète
                    les commandes ``NAVIGATE_TO:<path>`` pour naviguer dans la preview.

                    Parameters
                    ----------
                    pipe : IO
                        Pipe stdout ou stderr du subprocess.
                    color : str
                        Couleur hexadécimale pour l'affichage dans le terminal.
                    """
                    try:
                        for line in iter(pipe.readline, ''):
                            if line:
                                line_stripped = line.rstrip()
                                # Ignorer les messages de fermeture de session Flet
                                if "Session closed" in line_stripped or "session" in line_stripped.lower() and "closed" in line_stripped.lower():
                                    continue
                                # Détecter la commande de navigation
                                if line_stripped.startswith("NAVIGATE_TO:"):
                                    folder_to_navigate = line_stripped[12:]  # Enlever "NAVIGATE_TO:"
                                    if os.path.isdir(folder_to_navigate):
                                        # Utiliser pubsub pour mettre à jour l'UI depuis un thread
                                        page.pubsub.send_all_on_topic("navigate", folder_to_navigate)
                                else:
                                    log_to_terminal(line_stripped, color)
                    except:
                        pass
                    finally:
                        pipe.close()
                
                threading.Thread(target=read_output, args=(process.stdout, WHITE), daemon=True).start()
                threading.Thread(target=read_output, args=(process.stderr, RED), daemon=True).start()
                
            else:
                # Préparer l'environnement avec le chemin du dossier Data
                env = os.environ.copy()
                env["PYTHONIOENCODING"] = "utf-8"
                env["DATA_PATH"] = os.path.join(app_directory, "Data")
                env["FOLDER_PATH"] = current_browse_folder["path"] or selected_folder["path"]

                cycle_options = _get_dashboard_window_cycle_options(app_name) or {}
                should_apply_dashboard_window_cycle = bool(cycle_options)
                previous_maximized_state_for_cycle = (
                    bool(page.window.maximized)
                    if should_apply_dashboard_window_cycle and bool(cycle_options.get("restore_previous_maximized_state", False))
                    else None
                )



                # Ajouter les chemins ImageMagick pour Wand (Homebrew sur macOS)
                if platform.system() == "Darwin":
                    # Chemins Homebrew pour Apple Silicon et Intel
                    homebrew_paths = ["/opt/homebrew", "/usr/local"]
                    for brew_path in homebrew_paths:
                        magick_lib = os.path.join(brew_path, "lib")
                        if os.path.exists(magick_lib):
                            env["MAGICK_HOME"] = brew_path
                            env["DYLD_LIBRARY_PATH"] = magick_lib + ":" + env.get("DYLD_LIBRARY_PATH", "")
                            env["PATH"] = os.path.join(brew_path, "bin") + ":" + env.get("PATH", "")
                            break



                # Ajouter la taille de redimensionnement pour Redimensionner.py
                if app_name == "Redimensionner.py":
                    env["RESIZE_SIZE"] = resize_size["value"]



                # Ajouter la taille de redimensionnement avec watermark pour Redimensionner filigrane.py
                if app_name == "Redimensionner filigrane.py":
                    env["RESIZE_WATERMARK_SIZE"] = resize_watermark_size["value"]



                # Ajouter le dossier pour Transfert vers TEMP.py
                if app_name == "Transfert vers TEMP.py":
                    if platform.system() == "Windows":
                        env["DEST_FOLDER"] = "Z:/temp"
                    else:
                        env["DEST_FOLDER"] = "/Volumes/TRAVAUX EN COURS/Z2026/TEMP"



                # Ajouter le nom de la série pour Renommer séquence.py
                if app_name == "Renommer séquence.py" and series_name:
                    env["SERIES_NAME"] = series_name



                # Ajouter le nom du PDF pour Images en PDF.py
                if app_name == "Images en PDF.py" and series_name:
                    env["PDF_NAME"] = series_name



                # Ajouter les dimensions pour 2 en 1.py
                if app_name == "2 en 1.py" and series_name:
                    parts = series_name.split("x")
                    if len(parts) == 2:
                        env["TWO_IN_ONE_WIDTH"] = parts[0]
                        env["TWO_IN_ONE_HEIGHT"] = parts[1]


                # Ajouter les dimensions pour Fit 203.py
                if app_name == "Fit 203.py" and series_name:
                    parts = series_name.split("|")
                    if len(parts) >= 2:
                        env["FIT_203_CROP_SIZE"] = parts[0]
                        env["FIT_203_PRINT_SIZE"] = parts[1]
                    if len(parts) == 3:
                        env["FIT_203_OUTPUT_FOLDER"] = parts[2]


                # Ajouter les dimensions/portée pour Recadrage automatique.py
                if app_name == "Recadrage automatique.py" and series_name:
                    parts = series_name.split("|")
                    if len(parts) >= 1 and parts[0]:
                        env["FORCE_CROP_SIZE"] = parts[0]
                    if len(parts) >= 2 and parts[1] in ("selected", "all"):
                        env["FORCE_CROP_SCOPE"] = parts[1]
                    _fit_token = parts[2].strip().lower() if len(parts) >= 3 else "0"
                    env["FORCE_CROP_FIT"] = "1" if _fit_token in ("1", "fit", "true", "yes", "on") else "0"
                    _wb_token = parts[3].strip().lower() if len(parts) >= 4 else "0"
                    env["FORCE_CROP_WHITE_BORDER"] = "1" if _wb_token in ("1", "true", "yes", "on") else "0"

                # Paramètres Copyright
                if app_name == "Copyright.py" and series_name:
                    mode_part, _, custom_part = series_name.partition(":")
                    env["COPYRIGHT_MODE"] = mode_part
                    if mode_part == "custom" and custom_part:
                        env["COPYRIGHT_CUSTOM"] = custom_part

                # Paramètres Débruiter
                if app_name == "Débruiter.py" and series_name:
                    parts = series_name.split("|")
                    if len(parts) >= 4:
                        env["DENOISE_H"]               = parts[0]
                        env["DENOISE_H_COLOR"]         = parts[1]
                        env["DENOISE_TEMPLATE_WINDOW"] = parts[2]
                        env["DENOISE_SEARCH_WINDOW"]   = parts[3]

                # Paramètres Grain pellicule
                if app_name == "Grain pellicule.py" and series_name:
                    parts = series_name.split("|")
                    if len(parts) >= 18:
                        # Format 31 parties : G1|a1|s1|c1|sh1|G2|a2|s2|c2|sh2|H|ht|hr|hi|hred|B|br|bi|D|dst|dsi|dht|dhi|cv|css|cstr|cts|ctl|dmb|cs1|cs2
                        env["GRAIN1_ENABLED"]     = parts[0]
                        env["GRAIN_AMOUNT"]       = parts[1]
                        env["GRAIN_SIZE"]         = parts[2]
                        env["GRAIN_COLOR_RATIO"]  = parts[3]
                        env["GRAIN_SHADOW_BOOST"] = parts[4]
                        if parts[5] == "1":
                            env["GRAIN2_AMOUNT"]       = parts[6]
                            env["GRAIN2_SIZE"]         = parts[7]
                            env["GRAIN2_COLOR_RATIO"]  = parts[8]
                            env["GRAIN2_SHADOW_BOOST"] = parts[9]
                        env["HALATION_ENABLED"]   = parts[10]
                        env["HALATION_THRESHOLD"] = parts[11]
                        env["HALATION_RADIUS"]    = parts[12]
                        env["HALATION_INTENSITY"] = parts[13]
                        env["HALATION_RED_SHIFT"] = parts[14]
                        env["BLOOM_ENABLED"]      = parts[15]
                        env["BLOOM_RADIUS"]       = parts[16]
                        env["BLOOM_INTENSITY"]    = parts[17]
                        if len(parts) >= 23:
                            env["DESAT_ENABLED"]             = parts[18]
                            env["DESAT_SHADOW_THRESHOLD"]    = parts[19]
                            env["DESAT_SHADOW_INTENSITY"]    = parts[20]
                            env["DESAT_HIGHLIGHT_THRESHOLD"] = parts[21]
                            env["DESAT_HIGHLIGHT_INTENSITY"] = parts[22]
                        if len(parts) >= 28:
                            env["CURVE_ENABLED"]           = parts[23]
                            env["CURVE_SHOULDER_START"]    = parts[24]
                            env["CURVE_SHOULDER_STRENGTH"] = parts[25]
                            env["CURVE_TOE_START"]         = parts[26]
                            env["CURVE_TOE_LIFT"]          = parts[27]
                        if len(parts) >= 29:
                            env["DESAT_MIDTONE_BOOST"]     = parts[28]
                        if len(parts) >= 31:
                            env["GRAIN_CHROMA_SHIFT"]  = parts[29]
                            env["GRAIN2_CHROMA_SHIFT"] = parts[30]
                        if len(parts) >= 33:
                            env["CA_ENABLED"]  = parts[31]
                            env["CA_STRENGTH"] = parts[32]
                        if len(parts) >= 34:
                            env["CA_AXIAL_RATIO"] = parts[33]


                # (si aucun n'est sélectionné, la variable sera vide)
                if selected_files:
                    env["SELECTED_FILES"] = "|".join(os.path.basename(f) for f in selected_files)
                
                process = subprocess.Popen(
                    [sys.executable, "-u", app_path],
                    cwd=os.path.join(app_directory, "Data"),
                    env=env,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    bufsize=1,
                )
                if platform.system() == "Windows":
                    try:
                        import ctypes
                        ctypes.windll.user32.AllowSetForegroundWindow(process.pid)
                    except Exception:
                        pass

                if should_apply_dashboard_window_cycle:
                    page.window.minimized = True
                    page.update()



                # Lire la sortie en temps réel
                def read_output(pipe, color):
                    """
                    Lit en temps réel une pipe stdout ou stderr du sous-processus
                    et envoie chaque ligne au terminal.

                    Détecte les lignes préfixées par ``SELECTED_FILES:`` pour
                    resélectionner des fichiers dans la preview après exécution
                    du script.

                    Parameters
                    ----------
                    pipe : IO
                        Pipe stdout ou stderr du subprocess.
                    color : str
                        Couleur hexadécimale pour l'affichage dans le terminal.
                    """
                    try:
                        for line in iter(pipe.readline, ''):
                            if line:
                                line_stripped = line.rstrip()
                                if line_stripped.startswith(selected_files_prefix):
                                    selected_names = line_stripped[len(selected_files_prefix):]
                                    # Stocker pour que on_preview_ready l'applique
                                    # avec les données fraîches du prochain scan.
                                    pending_file_selection["names"] = selected_names
                                elif line_stripped.startswith("NAVIGATE_TO:"):
                                    nav_path = line_stripped[len("NAVIGATE_TO:"):].strip()
                                    if os.path.isdir(nav_path):
                                        page.pubsub.send_all_on_topic("navigate", nav_path)
                                else:
                                    log_to_terminal(line_stripped, color)
                    except Exception as read_err:
                        log_to_terminal(f"[ERREUR] Lecture sortie script: {read_err}", RED)
                    finally:
                        pipe.close()
                
                stdout_reader_thread = threading.Thread(target=read_output, args=(process.stdout, WHITE), daemon=True)
                stderr_reader_thread = threading.Thread(target=read_output, args=(process.stderr, RED), daemon=True)
                stdout_reader_thread.start()
                stderr_reader_thread.start()



                # Attendre la fin et rafraîchir la preview
                def done():
                    """
                    Attend la fin du sous-processus ET la lecture complète des pipes,
                    puis journalise le résultat et demande un rafraîchissement de la preview.
                    On attend les threads de lecture pour s'assurer que SELECTED_FILES:
                    a bien été traité avant de déclencher le refresh.
                    """
                    stdout_reader_thread.join()
                    stderr_reader_thread.join()
                    process.wait()
                    if should_apply_dashboard_window_cycle:
                        _restore_dashboard_window(previous_maximized_state_for_cycle)
                    app_progress_bar.visible = False
                    try:
                        page.update()
                    except Exception:
                        pass
                    log_to_terminal(f"[OK] {display_name} terminé", GREEN)
                    # Désélectionner les fichiers traités (sauf si le script en a sélectionné de nouveaux)
                    if selected_files and pending_file_selection["names"] is None:
                        page.pubsub.send_all_on_topic("deselect", None)
                    # Rafraîchir la preview pour afficher les nouveaux dossiers/fichiers créés
                    request_refresh()
                
                app_progress_bar.visible = True
                try:
                    page.update()
                except Exception:
                    pass
                threading.Thread(target=done, daemon=True).start()
        except Exception as err:
            log_to_terminal(f"[ERREUR] Erreur lors du lancement: {err}", RED)



    # ── Handlers des champs Redimensionner ───────────────────────────
    def on_resize_input_change(e):
        """Met à jour la taille de redimensionnement cible en pixels."""
        resize_size["value"] = e.control.value



    def launch_resize(e):
        """Lance Redimensionner.py avec la taille saisie dans resize_input."""
        app_path = os.path.join(app_directory, "Data", "Redimensionner.py")
        if os.path.exists(app_path):
            launch_app("Redimensionner.py", app_path, False)



    def on_resize_watermark_input_change(e):
        """Met à jour la taille de redimensionnement+filigrane cible en pixels."""
        resize_watermark_size["value"] = e.control.value



    def launch_resize_watermark(e):
        """Lance Redimensionner filigrane.py avec la taille saisie dans resize_watermark_input."""
        app_path = os.path.join(app_directory, "Data", "Redimensionner filigrane.py")
        if os.path.exists(app_path):
            launch_app("Redimensionner filigrane.py", app_path, False)



    def refresh_apps():
        """
        Reconstruit la grille des applications disponibles.

        Filtre les scripts présents sur disque, crée des widgets spéciaux
        (champ numérique + bouton) pour ``Redimensionner.py`` et
        ``Redimensionner filigrane.py``, et des boutons simples pour les autres.
        """
        items = []

        for app_name, app_config in apps.items():
            is_local = app_config[0]
            app_color = app_config[1]
            app_path = app_config[2] if len(app_config) > 2 else os.path.join(app_directory, "Data", app_name)
            is_special_gemini_panel = app_name == "IA / Bloc-notes"
            if not is_special_gemini_panel and not os.path.exists(app_path):
                continue

            if app_name == "Redimensionner.py":
                items.append(
                    ft.Container(
                        content=ft.Column([
                            ft.Text("Redimensionner", size=13, color=app_color, weight=ft.FontWeight.W_500, text_align=ft.TextAlign.CENTER),
                            resize_input,
                            ft.Text("px", size=11, color=LIGHT_GREY, text_align=ft.TextAlign.CENTER),
                        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, alignment=ft.MainAxisAlignment.CENTER, spacing=3),
                        expand=True,
                        alignment=ft.Alignment(0, 0),
                        bgcolor=GREY,
                        border=ft.Border.all(1, app_color),
                        padding=ft.Padding(5, 8, 5, 8),
                        border_radius=4,
                        on_click=launch_resize,
                        ink=True,
                    )
                )
            elif app_name == "Redimensionner filigrane.py":
                items.append(
                    ft.Container(
                        content=ft.Column([
                            ft.Text("Redimensionner + filigrane", size=12, color=app_color, weight=ft.FontWeight.W_500, text_align=ft.TextAlign.CENTER),
                            resize_watermark_input,
                            ft.Text("px", size=11, color=LIGHT_GREY, text_align=ft.TextAlign.CENTER),
                        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, alignment=ft.MainAxisAlignment.CENTER, spacing=3),
                        expand=True,
                        alignment=ft.Alignment(0, 0),
                        bgcolor=GREY,
                        border=ft.Border.all(1, app_color),
                        padding=ft.Padding(5, 8, 5, 8),
                        border_radius=4,
                        on_click=launch_resize_watermark,
                        ink=True,
                    )
                )
            elif is_special_gemini_panel:
                open_panels_title = ft.Text(
                    "Outils IA",
                    size=11,
                    color=BLUE,
                    text_align=ft.TextAlign.CENTER,
                    weight=ft.FontWeight.W_500,
                )
                open_panels_caption = ft.Text(
                    "IA / Bloc-notes",
                    size=12,
                    color=BLUE,
                    text_align=ft.TextAlign.CENTER,
                    weight=ft.FontWeight.W_500,
                )
                open_panels_card = ft.Container(
                    content=ft.Column(
                        [
                            open_panels_title,
                            open_panels_button,
                            open_panels_caption,
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=3,
                    ),
                    expand=True,
                    alignment=ft.Alignment(0, 0),
                    bgcolor=GREY,
                    border=ft.Border.all(1, RED if (ai_mode["value"] or note_mode["value"]) else BLUE),
                    padding=ft.Padding(10, 10, 10, 10),
                    border_radius=4,
                    on_click=lambda e: toggle_panels_open(),
                    ink=True,
                )
                open_panels_card_state["control"] = open_panels_card
                open_panels_card_state["title"] = open_panels_title
                open_panels_card_state["caption"] = open_panels_caption
                items.append(
                    open_panels_card
                )
            else:
                if app_name == "SidePanel.pyw":
                    on_click_handler = lambda e: _launch_side_panel()
                elif app_name == "Comparaison.pyw":
                    on_click_handler = lambda e: _launch_comparaison()
                else:
                    on_click_handler = lambda e, name=app_name, path=app_path, local=is_local: launch_app(name, path, local)
                display_name = (
                    "Side Panel" if app_name == "SidePanel.pyw"
                    else "Recadrage automatique" if app_name == "Recadrage automatique.py"
                    else "Recadrage manuel" if app_name == "Recadrage manuel.pyw"
                    else (app_name[:-4] if app_name.endswith(".pyw") else app_name[:-3])
                )
                items.append(
                    ft.Container(
                        content=ft.Text(
                            display_name,
                            size=14,
                            color=app_color,
                            text_align=ft.TextAlign.CENTER,
                            weight=ft.FontWeight.W_500,
                            max_lines=3,
                        ),
                        expand=True,
                        alignment=ft.Alignment(0, 0),
                        on_click=on_click_handler,
                        bgcolor=GREY,
                        border=ft.Border.all(1, app_color),
                        padding=ft.Padding(10, 10, 10, 10),
                        border_radius=4,
                        ink=True,
                    )
                )



        # Grouper en rangées de 3 avec padding uniforme
        apps_list.controls.clear()
        for i in range(0, len(items), 3):
            row_items = items[i:i + 3]
            # Compléter la dernière rangée si incomplète
            while len(row_items) < 3:
                row_items.append(ft.Container(expand=True))
            apps_list.controls.append(
                ft.Row(controls=row_items, expand=True, spacing=8)
            )

        page.update()



    def _build_quick_tools():
        """Construit la colonne d'icônes rondes (outils rapides)."""
        two_in_one_path = os.path.join(app_directory, "Data", "2 en 1.py")
        side_panel_path = os.path.join(app_directory, "Data", "SidePanel.pyw")

        def _round_button(icon, color, tooltip, on_click):
            return ft.Container(
                content=ft.Icon(icon, color=color, size=22),
                bgcolor=GREY,
                border=ft.Border.all(1, color),
                border_radius=50,
                width=44,
                height=44,
                alignment=ft.Alignment(0, 0),
                tooltip=tooltip,
                on_click=on_click,
                ink=True,
            )

        noir_et_blanc_path        = os.path.join(app_directory, "Data", "N&B.py")
        ameliorer_nettete_path    = os.path.join(app_directory, "Data", "Améliorer netteté.py")
        nettoyer_metadonnees_path = os.path.join(app_directory, "Data", "Nettoyer metadonnées.py")
        copyright_path            = os.path.join(app_directory, "Data", "Copyright.py")
        images_en_pdf_path        = os.path.join(app_directory, "Data", "Images en PDF.py")
        remerciements_path        = os.path.join(app_directory, "Data", "Remerciements.py")
        copier_nefs_path          = os.path.join(app_directory, "Data", "Copier NEFs sélection.py")
        copier_selection_path       = os.path.join(app_directory, "Data", "Copier sélection.py")
        separer_raw_jpg_path      = os.path.join(app_directory, "Data", "Séparer RAW et JPG.py")
        recadrage_automatique_path = os.path.join(app_directory, "Data", "Recadrage automatique.py")

        quick_tools_col.controls = [
            _round_button(
                ft.Icons.MONOCHROME_PHOTOS,
                WHITE,
                "N&B",
                lambda e: launch_app("N&B.py", noir_et_blanc_path, False),
            ),
            _round_button(
                ft.Icons.AUTO_GRAPH,
                WHITE,
                "Améliorer netteté",
                lambda e: launch_app("Améliorer netteté.py", ameliorer_nettete_path, False),
            ),
            _round_button(
                ft.Icons.CLEANING_SERVICES,
                RED,
                "Nettoyer métadonnées",
                lambda e: launch_app("Nettoyer metadonnées.py", nettoyer_metadonnees_path, False),
            ),
            _round_button(
                ft.Icons.PICTURE_AS_PDF,
                BLUE,
                "Images en PDF",
                lambda e: launch_app("Images en PDF.py", images_en_pdf_path, False),
            ),
            _round_button(
                ft.CupertinoIcons.BIN_XMARK_FILL,
                VIOLET,
                "Remerciements",
                lambda e: launch_app("Remerciements.py", remerciements_path, False),
            ),
            _round_button(
                ft.Icons.FOLDER_ZIP,
                ORANGE,
                "Zipper la sélection",
                _prompt_and_zip_selection,
            ),
            _round_button(
                ft.Icons.HIDE_IMAGE,
                YELLOW,
                "Séparer RAW et JPG",
                lambda e: launch_app("Séparer RAW et JPG.py", separer_raw_jpg_path, False),
            ),
            _round_button(
                ft.Icons.FOLDER_COPY,
                YELLOW,
                "Copier sélection → SELECTION",
                lambda e: launch_app("Copier sélection.py", copier_selection_path, False),
            ),
            _round_button(
                ft.Icons.IMAGE_SEARCH_OUTLINED,
                YELLOW,
                "Copier NEFs → SELECTION",
                lambda e: launch_app("Copier NEFs sélection.py", copier_nefs_path, False),
            ),
            _round_button(
                ft.Icons.NOTE_ADD,
                BLUE,
                "Créer INFO.txt dans le dossier courant",
                _create_and_open_info_txt,
            ),
        ]



    # ================================================================ #
    #                       ACTIONS FENÊTRE                            #
    # ================================================================ #
    async def pick_folder(e):
        """
        Ouvre le sélecteur de dossier natif et navigue vers le dossier choisi.

        Appelle ``ft.FilePicker.get_directory_path`` de façon asynchrone,
        normalise le chemin résultant et déclenche un rafraîchissement de la preview.
        """
        folder = await ft.FilePicker().get_directory_path(dialog_title="Sélectionner un dossier contenant des images")
        if folder:
            selected_folder["path"] = os.path.normpath(folder)
            current_browse_folder["path"] = selected_folder["path"]
            folder_path.value = _short_path(selected_folder["path"])
            folder_path.update()
            selected_files.clear()
            search_query["value"] = ""
            search_field.value = ""
            preview_page["value"] = 0
            _add_to_recent(selected_folder["path"])
            _rebuild_recent_folders_menu()
            refresh_preview()



    async def close_window(e):
        """Ferme la fenêtre principale de l'application de façon asynchrone."""
        await page.window.close()



    def update_app(e):
        """Sauvegarde les fichiers utilisateur, git pull --rebase, vérifie les dépendances si requirements a changé, relance."""
        log_to_terminal("Mise à jour en cours…", YELLOW)
        def _run_update():



            def run_git_command(*args):
                return subprocess.run(
                    ["git", *args],
                    cwd=app_directory,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                )



            # ── Sauvegarde mémoire des fichiers utilisateur avant git ─
            user_data_filenames = [".recent_folders.json", ".favorites.json", ".pip_cache.json"]
            user_data_backups = {}
            for file_name in user_data_filenames:
                file_path = os.path.join(app_directory, file_name)
                if os.path.isfile(file_path):
                    try:
                        with open(file_path, "r", encoding="utf-8") as file_handle:
                            user_data_backups[file_name] = file_handle.read()
                    except Exception:
                        pass



            def _restore_user_data_files():
                for file_name, content in user_data_backups.items():
                    file_path = os.path.join(app_directory, file_name)
                    try:
                        with open(file_path, "w", encoding="utf-8") as file_handle:
                            file_handle.write(content)
                    except Exception:
                        pass

            try:
                # Stash les changements locaux s'il y en a
                stash_result = run_git_command("stash")
                had_local_changes = "No local changes" not in stash_result.stdout

                # Pull --rebase
                git_pull_result = run_git_command("pull", "--rebase", "origin")
                git_command_output = (git_pull_result.stdout + git_pull_result.stderr).strip()

                if git_pull_result.returncode != 0:
                    if had_local_changes:
                        run_git_command("rebase", "--abort")
                        run_git_command("stash", "pop")
                    _restore_user_data_files()
                    log_to_terminal(f"[ERREUR] Erreur lors de la mise à jour.\n{git_command_output}", RED)
                    return



                # Supprimer le stash (changements de code locaux non désirés)
                if had_local_changes:
                    run_git_command("stash", "drop")



                # Restaurer systématiquement les fichiers utilisateur
                _restore_user_data_files()

                if "Already up to date" in git_command_output or "Déjà à jour" in git_command_output or git_command_output == "":
                    log_to_terminal("[OK] Déjà à jour.", GREEN)
                else:
                    log_to_terminal(f"[OK] Code mis à jour.\n{git_command_output}", GREEN)



                # ── Dépendances : pip uniquement si requirements.txt a changé ──
                requirements_file_path = os.path.join(app_directory, "requirements.txt")
                pip_cache_file_path = os.path.join(app_directory, ".pip_cache.json")
                if not os.path.isfile(requirements_file_path):
                    log_to_terminal("⚠ requirements.txt introuvable, installation ignorée.", YELLOW)
                else:
                    with open(requirements_file_path, "rb") as f:
                        requirements_checksum = hashlib.sha256(f.read()).hexdigest()

                    cached_checksum = None
                    try:
                        with open(pip_cache_file_path, "r", encoding="utf-8") as f:
                            cached_checksum = json.load(f).get("req_hash")
                    except Exception:
                        pass

                    # ── flet + flet-desktop : toujours synchronisés ──────
                    log_to_terminal("🔌 Mise à jour de flet et flet-desktop…", YELLOW)
                    flet_upgrade_proc = subprocess.Popen(
                        [sys.executable, "-m", "pip", "install", "flet", "flet-desktop", "--upgrade"],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.STDOUT,
                        text=True,
                        encoding="utf-8",
                        errors="replace",
                        cwd=app_directory,
                    )
                    for line in flet_upgrade_proc.stdout:
                        line = line.rstrip()
                        if line:
                            log_to_terminal(line, LIGHT_GREY)
                    flet_upgrade_proc.wait()
                    if flet_upgrade_proc.returncode == 0:
                        log_to_terminal("[OK] flet et flet-desktop mis à jour.", GREEN)
                    else:
                        log_to_terminal(f"⚠ flet-desktop : pip a terminé avec le code {flet_upgrade_proc.returncode}.", YELLOW)

                    if cached_checksum == requirements_checksum:
                        log_to_terminal("[OK] Dépendances inchangées, installation ignorée.", GREEN)
                    else:
                        log_to_terminal("📦 Nouvelles dépendances détectées, installation en cours…", YELLOW)
                        pip_install_process = subprocess.Popen(
                            [sys.executable, "-m", "pip", "install", "-r", requirements_file_path, "--upgrade"],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.STDOUT,
                            text=True,
                            encoding="utf-8",
                            errors="replace",
                            cwd=app_directory,
                        )
                        for line in pip_install_process.stdout:
                            line = line.rstrip()
                            if line:
                                log_to_terminal(line, LIGHT_GREY)
                        pip_install_process.wait()
                        if pip_install_process.returncode == 0:
                            log_to_terminal("[OK] Dépendances installées.", GREEN)
                            try:
                                with open(pip_cache_file_path, "w", encoding="utf-8") as f:
                                    json.dump(
                                        {"req_hash": requirements_checksum,
                                         "updated_at": time.strftime("%Y-%m-%d %H:%M")},
                                        f, ensure_ascii=False, indent=2,
                                    )
                            except Exception:
                                pass
                        else:
                            log_to_terminal(f"pip a terminé avec le code {pip_install_process.returncode}.", YELLOW)

                # ── Mise à jour d'Ollama ──────────────────────────────────
                log_to_terminal("🤖 Mise à jour d'Ollama...", YELLOW)
                try:
                    ollama_check = subprocess.run(
                        ["ollama", "--version"],
                        capture_output=True, text=True, timeout=5,
                    )
                    if ollama_check.returncode == 0:
                        log_to_terminal(f"[OK] Ollama détecté : {ollama_check.stdout.strip()}", GREEN)
                        tags_resp = subprocess.run(
                            ["ollama", "list"],
                            capture_output=True, text=True, timeout=10,
                        )
                        if tags_resp.returncode == 0:
                            for model_line in tags_resp.stdout.splitlines()[1:]:
                                model_tag = model_line.split()[0] if model_line.split() else ""
                                if model_tag:
                                    log_to_terminal(f"  ⬇️ ollama pull {model_tag}", LIGHT_GREY)
                                    pull_proc = subprocess.Popen(
                                        ["ollama", "pull", model_tag],
                                        stdout=subprocess.PIPE,
                                        stderr=subprocess.STDOUT,
                                        text=True, encoding="utf-8", errors="replace",
                                    )
                                    for pull_line in pull_proc.stdout:
                                        pull_line = pull_line.rstrip()
                                        if pull_line:
                                            log_to_terminal(f"    {pull_line}", LIGHT_GREY)
                                    pull_proc.wait()
                            log_to_terminal("[OK] Modèles Ollama à jour.", GREEN)
                    else:
                        log_to_terminal("[AVERTISSEMENT] Ollama non détecté, ignoré.", YELLOW)
                except FileNotFoundError:
                    log_to_terminal("Ollama non trouvé, tentative d'installation...", YELLOW)
                    try:
                        if sys.platform == "win32":
                            install_script = (
                                "Invoke-WebRequest -Uri 'https://ollama.com/download/OllamaSetup.exe' "
                                "-OutFile '$env:TEMP\\OllamaSetup.exe'; "
                                "Start-Process '$env:TEMP\\OllamaSetup.exe' -Wait"
                            )
                            subprocess.run(
                                ["powershell", "-Command", install_script],
                                timeout=300,
                            )
                        else:
                            subprocess.run(
                                ["sh", "-c", "curl -fsSL https://ollama.com/install.sh | sh"],
                                timeout=300,
                            )
                        subprocess.run(["ollama", "pull", CONSTANTS.AI_MODEL_TEXT], timeout=3600)
                        log_to_terminal("[OK] Ollama installé.", GREEN)
                    except Exception as install_exc:
                        log_to_terminal(f"[AVERTISSEMENT] Impossible d'installer Ollama : {install_exc}", YELLOW)
                        log_to_terminal("[INFO] Installez-le manuellement : https://ollama.com/download", LIGHT_GREY)
                except Exception as ollama_exc:
                    log_to_terminal(f"[AVERTISSEMENT] Ollama : {ollama_exc}", YELLOW)



                # ── Redémarrage automatique ───────────────────────────────
                log_to_terminal("🔄 Redémarrage du Dashboard…", BLUE)
                dashboard_path = os.path.abspath(__file__)
                async def _restart_after_update():
                    import time as _time
                    _time.sleep(0.4)
                    subprocess.Popen([sys.executable, dashboard_path])
                    _time.sleep(0.2)
                    try:
                        await page.window.close()
                    except Exception:
                        pass
                    os._exit(0)
                page.run_task(_restart_after_update)

            except Exception as exc:
                _restore_user_data_files()
                log_to_terminal(f"[ERREUR] {exc}", RED)

        threading.Thread(target=_run_update, daemon=True).start()



# ===================== CONNEXIONS UI ===================== #
    # ── Champ dossier ────────────────────────────────────────────────
    folder_path.on_submit = on_folder_path_submit
    folder_path.on_focus = lambda e: _suspend_keyboard_shortcuts()
    folder_path.on_blur = lambda e: (on_folder_path_blur(e), _resume_keyboard_shortcuts())



    # ── Preview ───────────────────────────────────────────────────────
    sort_segment.on_change = on_sort_change
    select_toggle_button.on_click = toggle_select_all
    invert_selection_button.on_click = invert_selection
    select_same_date_button.on_click = select_same_date
    filter_sel_btn.on_click = _toggle_show_only_selection
    prev_page_btn.on_click = lambda e: go_to_page(-1)
    next_page_btn.on_click = lambda e: go_to_page(+1)



    # ── Recherche preview ─────────────────────────────────────────────
    search_field.on_change = _on_search_change
    search_field.on_submit = _on_search_change
    search_field.on_focus = lambda e: _suspend_keyboard_shortcuts()
    search_field.on_blur = lambda e: _resume_keyboard_shortcuts()
    search_close_btn.on_click = _clear_search



    # ── Redimensionnement ─────────────────────────────────────────────
    resize_input.on_change = on_resize_input_change
    resize_watermark_input.on_change = on_resize_watermark_input_change
    resize_input.on_focus = lambda e: _suspend_keyboard_shortcuts()
    resize_input.on_blur = lambda e: _resume_keyboard_shortcuts()
    resize_watermark_input.on_focus = lambda e: _suspend_keyboard_shortcuts()
    resize_watermark_input.on_blur = lambda e: _resume_keyboard_shortcuts()


    # Bouton global utilisé dans la grille pour ouvrir/fermer IA + bloc-notes.
    open_panels_button = ft.IconButton(
        icon=ft.Icons.SMART_TOY,
        tooltip="Ouvrir IA & Bloc-notes",
        icon_color=BLUE,
        icon_size=16,
        on_click=lambda e: toggle_panels_open(),
    )
    open_panels_card_state = {"control": None, "title": None, "caption": None}



    # ── Initialisation ────────────────────────────────────────────────
    _rebuild_recent_folders_menu()
    refresh_apps()
    _build_quick_tools()
    _rebuild_favorites_panel()
    _ai_load_history()
    _initial_drives = _get_removable_drives()
    _rebuild_drives_panel(_initial_drives)
    threading.Thread(target=_poll_removable_drives, daemon=True).start()
    overlay_container       = None
    ai_panel_container      = None
    notepad_panel_container = None
    bottom_panel_container  = None
    ai_both_fullscreen_btn  = None
    if _initial_drives is not None:
        def update_overlay_visibility():
            """Affiche ou masque l'overlay (IA à gauche + Notes à droite)."""
            panels_are_open = ai_mode["value"] or note_mode["value"]
            # Nettoyage du mode solo si les panneaux se ferment
            if not panels_are_open and overlay_fullscreen["mode"] in ("ai", "notepad"):
                # Restaurer bottom_panel_container au mode normal
                if bottom_panel_container is not None:
                    bottom_panel_container.top    = None
                    bottom_panel_container.right  = 0
                    bottom_panel_container.width  = None
                    bottom_panel_container.height = CONSTANTS.TERMINAL_HEIGHT
                if ai_panel_container is not None:
                    ai_panel_container.visible = True
                if notepad_panel_container is not None:
                    notepad_panel_container.visible = True
                _terminal_spacer.height = CONSTANTS.TERMINAL_HEIGHT
            if overlay_container is not None:
                overlay_container.visible = panels_are_open
            if not panels_are_open:
                overlay_fullscreen["mode"] = None
                if ai_panel_container is not None:
                    ai_panel_container.visible = True
                if notepad_panel_container is not None:
                    notepad_panel_container.visible = True
            if open_panels_button is not None:
                open_panels_button.icon       = ft.Icons.SMART_TOY
                open_panels_button.icon_color = RED if panels_are_open else BLUE
                open_panels_button.tooltip    = "Fermer IA & Notes" if panels_are_open else "Ouvrir IA & Bloc-notes"
            open_panels_card = open_panels_card_state["control"]
            if open_panels_card is not None:
                open_panels_card.bgcolor = GREY
                open_panels_card.border = ft.Border.all(1, RED if panels_are_open else BLUE)
                title_control = open_panels_card_state["title"]
                caption_control = open_panels_card_state["caption"]
                if title_control is not None:
                    title_control.color = RED if panels_are_open else BLUE
                if caption_control is not None:
                    caption_control.color = RED if panels_are_open else BLUE
                try:
                    if open_panels_button is not None:
                        open_panels_button.update()
                    if title_control is not None:
                        title_control.update()
                    if caption_control is not None:
                        caption_control.update()
                    open_panels_card.update()
                except Exception:
                    pass



        def toggle_panels_open():
            if ai_mode["value"] or note_mode["value"]:
                switch_to_terminal_mode()
            else:
                note_target_file["path"] = notes_file_path
                load_notes()
                note_mode["value"] = True
                ai_mode["value"]   = True
                terminal_output.visible  = False
                terminal_cmd_row.visible = False
                update_overlay_visibility()
                terminal_output.update()
                terminal_cmd_row.update()
                try:
                    page.update()
                except Exception:
                    pass


        bottom_panel_container = ft.Container(
            content=ft.Stack([
                ft.Row([
                    ft.Container(
                        content=ft.Row([
                            ft.Column([
                                terminal_output,
                                app_progress_bar,
                                terminal_cmd_row,
                            ], spacing=4, expand=True),
                            ft.Column([
                                expand_button_terminal,
                                ft.IconButton(
                                    icon=ft.Icons.COPY_ALL,
                                    tooltip="Copier le terminal",
                                    on_click=lambda e: copy_terminal_to_clipboard(),
                                    icon_color=BLUE,
                                    icon_size=16,
                                ),
                                ft.IconButton(
                                    icon=ft.Icons.CLEAR_ALL,
                                    tooltip="Effacer le terminal",
                                    on_click=clear_terminal,
                                    icon_color=RED,
                                    icon_size=16,
                                ),
                                ft.IconButton(
                                    icon=ft.Icons.SEND,
                                    icon_color=GREEN,
                                    icon_size=16,
                                    tooltip="Envoyer la commande",
                                    on_click=on_terminal_command_submit,
                                ),
                            ], alignment=ft.MainAxisAlignment.END, spacing=0),
                        ], spacing=4, expand=True, vertical_alignment=ft.CrossAxisAlignment.STRETCH),
                        expand=True,
                        border=ft.Border.all(1, GREEN),
                        border_radius=8,
                        bgcolor=DARK,
                        padding=5,
                    ),
                    ft.Row([
                        favorites_panel,
                        drives_panel,
                    ], expand=True, spacing=8, vertical_alignment=ft.CrossAxisAlignment.STRETCH),
                ], spacing=8, expand=True),
                overlay_container if overlay_container is not None else ft.Container(),
            ]),
            height=CONSTANTS.TERMINAL_HEIGHT,
            bgcolor=BACKGROUND,
            bottom=0,
            left=0,
            right=0,
        )
        ai_clear_button = ft.IconButton(
            icon=ft.Icons.DELETE_SWEEP,
            icon_color=LIGHT_GREY,
            icon_size=16,
            tooltip="Effacer la conversation IA",
            on_click=lambda e: _clear_ai_conversation(),
        )
        ai_fullscreen_btn = ft.IconButton(
            icon=ft.Icons.FULLSCREEN,
            icon_color=GREEN,
            icon_size=16,
            tooltip="IA en plein écran",
            on_click=lambda e: toggle_ai_true_fullscreen(),
        )
        ai_copy_button = ft.IconButton(
            icon=ft.Icons.COPY_ALL,
            icon_color=BLUE,
            icon_size=16,
            tooltip="Copier la conversation IA",
            on_click=lambda e: _export_ai_conversation(to_notepad=False),
        )
        ai_to_notepad_button = ft.IconButton(
            icon=ft.Icons.SEND_TO_MOBILE,
            icon_color=VIOLET,
            icon_size=16,
            tooltip="Transférer la conversation vers le bloc-notes",
            on_click=lambda e: _export_ai_conversation(to_notepad=True),
        )
        ai_both_fullscreen_btn = ft.IconButton(
            icon=ft.Icons.VERTICAL_SPLIT,
            icon_color=YELLOW,
            icon_size=16,
            tooltip="IA + Bloc-notes côte à côte (plein écran)",
            on_click=lambda e: toggle_both_fullscreen(),
        )
        ai_panel_header = ft.Row([
            ft.Icon(ft.Icons.SMART_TOY, color=BLUE, size=14),
            ft.Text("IA", color=BLUE, size=11, weight=ft.FontWeight.BOLD),
            ft.Container(width=4),
            ai_model_dropdown,
            ft.Container(width=4),
            ft.Container(
                content=ft.Row([ai_stop_button], spacing=0),
                border=ft.Border.all(1, GREY),
                border_radius=6,
                padding=ft.Padding(0, 0, 0, 0),
            ),
            ai_image_size_button,
            ai_image_mode_label,
            ai_clear_button,
            ai_speaker_button,
            ai_copy_button,
            ai_to_notepad_button,
            ft.Container(expand=True),
            ai_both_fullscreen_btn,
            expand_button_overlay,
            ai_fullscreen_btn,
        ], spacing=2, vertical_alignment=ft.CrossAxisAlignment.CENTER)
        ai_panel_container = ft.Container(
            content=ft.Column([ai_panel_header, ai_container], spacing=4, expand=True),
            expand=True,
            bgcolor=DARK,
            border=ft.Border.all(1, BLUE),
            border_radius=8,
            padding=5,
        )

        notepad_clear_button = ft.IconButton(
            icon=ft.Icons.DELETE_SWEEP,
            icon_color=ORANGE,
            icon_size=16,
            tooltip="Effacer tout le bloc-notes",
            on_click=lambda e: _notepad_clear(),
        )
        notepad_fullscreen_btn = ft.IconButton(
            icon=ft.Icons.FULLSCREEN,
            icon_color=GREEN,
            icon_size=16,
            tooltip="Bloc-notes en plein écran",
            on_click=lambda e: toggle_notepad_true_fullscreen(),
        )
        notepad_home_button = ft.IconButton(
            icon=ft.Icons.HOME,
            icon_color=VIOLET,
            icon_size=16,
            tooltip="Charger la note par défaut (.notes.md)",
            on_click=lambda e: switch_to_note(),
        )
        notepad_preview_button = ft.IconButton(
            icon=ft.Icons.VISIBILITY,
            icon_color=LIGHT_GREY,
            icon_size=16,
            tooltip="Prévisualiser en Markdown",
            on_click=lambda e: _notepad_toggle_preview(),
        )
        notepad_save_as_button = ft.IconButton(
            icon=ft.Icons.SAVE_AS,
            icon_color=BLUE,
            icon_size=16,
            tooltip="Sauvegarder les notes sous…",
            on_click=lambda e: page.run_task(_notepad_save_as),
        )
        notepad_panel_header = ft.Row([
            notepad_header_icon,
            notepad_header_title,
            notepad_clear_button,
            notepad_home_button,
            notepad_preview_button,
            notepad_save_as_button,
            ft.Container(expand=True),
            expand_button_notepad,
            notepad_fullscreen_btn,
        ], spacing=4, vertical_alignment=ft.CrossAxisAlignment.CENTER)
        notepad_panel_container = ft.Container(
            content=ft.Column([notepad_panel_header, notepad_container], spacing=4, expand=True),
            expand=True,
            bgcolor=DARK,
            border=ft.Border.all(1, VIOLET),
            border_radius=8,
            padding=5,
        )

    overlay_container = ft.Container(
        content=ft.Row([
            ai_panel_container,
            notepad_panel_container,
        ], expand=True, spacing=8),
        visible=False,
        bgcolor=BACKGROUND,
        left=0, right=0, top=0, bottom=0,
    )

    if bottom_panel_container is not None:
        try:
            bottom_panel_container.content.controls[1] = overlay_container
        except Exception:
            pass

    # ── Spacer et conteneur solo (mode pleine hauteur gauche) ─────────────────
    # _terminal_spacer : réserve la hauteur du terminal dans la colonne principale.
    # En mode solo, sa hauteur passe à 0 pour que la colonne gauche prenne toute la hauteur.
    _terminal_spacer = ft.Container(height=CONSTANTS.TERMINAL_HEIGHT)

    # _solo_left_container : overlay Stack positionné à gauche, visible uniquement en mode solo.
    # Sa largeur correspond à la colonne apps (expand=6 sur 15) et il couvre la hauteur entière.
    _solo_left_inner     = ft.Column([], expand=True)
    _solo_left_container = ft.Container(
        content=_solo_left_inner,
        visible=False,
        bgcolor=BACKGROUND,
        left=0, top=0, bottom=0,
    )
    _solo_left_state["container"] = _solo_left_container



    def _expanded_terminal_height():
        """Calcule une hauteur terminal étendue robuste même pendant un resize/maximize."""
        win_h = page.window.height or CONSTANTS.WINDOW_HEIGHT
        return max(CONSTANTS.TERMINAL_HEIGHT, int(win_h - CONSTANTS.WDA_HEIGHT))



    def _enter_solo_mode(panel_container, mode_name, do_update=True):
        """Bascule un panneau en mode solo pleine hauteur à gauche."""
        overlay_fullscreen["mode"] = mode_name
        ai_is_solo = mode_name in ("ai", "ai_full")
        notepad_is_solo = mode_name in ("notepad", "notepad_full")
        expand_button_overlay.icon = ft.Icons.CLOSE_FULLSCREEN if ai_is_solo else ft.Icons.OPEN_IN_FULL
        expand_button_notepad.icon = ft.Icons.CLOSE_FULLSCREEN if notepad_is_solo else ft.Icons.OPEN_IN_FULL
        expand_button_overlay.tooltip = "Réduire IA seule (Ctrl/Cmd+←)" if ai_is_solo else "IA seule (Ctrl/Cmd+←)"
        expand_button_notepad.tooltip = "Réduire Bloc-notes seul (Ctrl/Cmd+→)" if notepad_is_solo else "Bloc-notes seul (Ctrl/Cmd+→)"
        # Repositionner le panneau du bas selon le mode.
        if mode_name in ("ai_full", "notepad_full"):
            # Plein écran réel (moins WDA), conserve la preview en dessous.
            bottom_panel_container.width = None
            bottom_panel_container.left = 0
            bottom_panel_container.right = 0
        else:
            # Mode colonne gauche avec preview_list visible à droite.
            win_w = page.window.width or CONSTANTS.WINDOW_WIDTH
            bottom_panel_container.width = int((win_w - 8) * 6 / 15 + 4)
            bottom_panel_container.left = 0
            bottom_panel_container.right = None
        bottom_panel_container.top    = 0
        bottom_panel_container.height = None
        # Masquer le panneau inactif ; l'actif reste dans overlay_container
        if ai_panel_container is not None:
            ai_panel_container.visible = (mode_name in ("ai", "ai_full"))
        if notepad_panel_container is not None:
            notepad_panel_container.visible = (mode_name in ("notepad", "notepad_full"))
        if overlay_container is not None:
            overlay_container.visible = True
        _terminal_spacer.height = 0
        # Mettre à jour les boutons plein écran (header IA / Notes)
        if mode_name == "ai_full":
            ai_fullscreen_btn.icon = ft.Icons.FULLSCREEN_EXIT
            ai_fullscreen_btn.tooltip = "Quitter le plein écran IA"
            notepad_fullscreen_btn.icon = ft.Icons.FULLSCREEN
            notepad_fullscreen_btn.tooltip = "Bloc-notes en plein écran"
        elif mode_name == "notepad_full":
            notepad_fullscreen_btn.icon = ft.Icons.FULLSCREEN_EXIT
            notepad_fullscreen_btn.tooltip = "Quitter le plein écran Bloc-notes"
            ai_fullscreen_btn.icon = ft.Icons.FULLSCREEN
            ai_fullscreen_btn.tooltip = "IA en plein écran"
        else:
            ai_fullscreen_btn.icon = ft.Icons.FULLSCREEN
            ai_fullscreen_btn.tooltip = "IA en plein écran"
            notepad_fullscreen_btn.icon = ft.Icons.FULLSCREEN
            notepad_fullscreen_btn.tooltip = "Bloc-notes en plein écran"
        if do_update:
            page.update()



    def _exit_solo_mode(do_update=True):
        """Restaure le mode deux panneaux depuis le mode solo."""
        overlay_fullscreen["mode"] = None
        expand_button_overlay.icon = ft.Icons.OPEN_IN_FULL
        expand_button_notepad.icon = ft.Icons.OPEN_IN_FULL
        expand_button_overlay.tooltip = "IA seule (Ctrl/Cmd+←)"
        expand_button_notepad.tooltip = "Bloc-notes seul (Ctrl/Cmd+→)"
        # Restaurer bottom_panel_container en barre de fond
        bottom_panel_container.top    = None
        bottom_panel_container.right  = 0
        bottom_panel_container.left   = 0
        bottom_panel_container.width  = None
        bottom_panel_container.height = (
            _expanded_terminal_height()
            if terminal_is_expanded["value"]
            else CONSTANTS.TERMINAL_HEIGHT
        )
        if ai_panel_container is not None:
            ai_panel_container.visible = True
        if notepad_panel_container is not None:
            notepad_panel_container.visible = True
        _terminal_spacer.height = CONSTANTS.TERMINAL_HEIGHT
        # Restaurer les deux boutons à leur icône d'origine
        ai_fullscreen_btn.icon         = ft.Icons.FULLSCREEN
        ai_fullscreen_btn.tooltip      = "IA en plein écran"
        notepad_fullscreen_btn.icon    = ft.Icons.FULLSCREEN
        notepad_fullscreen_btn.tooltip = "Bloc-notes en plein écran"
        if ai_both_fullscreen_btn is not None:
            ai_both_fullscreen_btn.icon    = ft.Icons.VERTICAL_SPLIT
            ai_both_fullscreen_btn.tooltip = "IA + Bloc-notes côte à côte (plein écran)"
        if do_update:
            page.update()



    def toggle_ai_fullscreen():
        """Mode solo IA : panel IA pleine hauteur à gauche, preview_list visible à droite."""
        if overlay_fullscreen["mode"] == "ai":
            _exit_solo_mode()
        elif overlay_fullscreen["mode"] in ("notepad", "both_full"):
            _exit_solo_mode(do_update=False)
            _enter_solo_mode(ai_panel_container, "ai")
        else:
            _enter_solo_mode(ai_panel_container, "ai")



    def toggle_notepad_fullscreen():
        """Mode solo Bloc-notes : panel notes pleine hauteur à gauche, preview_list visible à droite."""
        if overlay_fullscreen["mode"] == "notepad":
            _exit_solo_mode()
        elif overlay_fullscreen["mode"] == "ai":
            _exit_solo_mode(do_update=False)
            _enter_solo_mode(notepad_panel_container, "notepad")
        else:
            _enter_solo_mode(notepad_panel_container, "notepad")



    def toggle_ai_true_fullscreen():
        """Plein écran réel IA (moins WDA)."""
        if overlay_fullscreen["mode"] == "ai_full":
            _exit_solo_mode()
        elif overlay_fullscreen["mode"] in ("ai", "notepad", "notepad_full", "both_full"):
            _exit_solo_mode(do_update=False)
            _enter_solo_mode(ai_panel_container, "ai_full")
        else:
            _enter_solo_mode(ai_panel_container, "ai_full")



    def toggle_notepad_true_fullscreen():
        """Plein écran réel Bloc-notes (moins WDA)."""
        if overlay_fullscreen["mode"] == "notepad_full":
            _exit_solo_mode()
        elif overlay_fullscreen["mode"] in ("notepad", "ai", "ai_full", "both_full"):
            _exit_solo_mode(do_update=False)
            _enter_solo_mode(notepad_panel_container, "notepad_full")
        else:
            _enter_solo_mode(notepad_panel_container, "notepad_full")



    def toggle_both_fullscreen():
        """IA + Bloc-notes côte à côte en plein écran réel."""
        if overlay_fullscreen["mode"] == "both_full":
            _exit_solo_mode()
            return
        if overlay_fullscreen["mode"] is not None:
            _exit_solo_mode(do_update=False)
        overlay_fullscreen["mode"] = "both_full"
        ai_panel_container.visible      = True
        notepad_panel_container.visible = True
        overlay_container.visible       = True
        bottom_panel_container.top    = 0
        bottom_panel_container.height = None
        bottom_panel_container.left   = 0
        bottom_panel_container.right  = 0
        bottom_panel_container.width  = None
        _terminal_spacer.height = 0
        expand_button_overlay.icon     = ft.Icons.OPEN_IN_FULL
        expand_button_notepad.icon     = ft.Icons.OPEN_IN_FULL
        expand_button_overlay.tooltip  = "IA seule (Ctrl/Cmd+←)"
        expand_button_notepad.tooltip  = "Bloc-notes seul (Ctrl/Cmd+→)"
        ai_both_fullscreen_btn.icon    = ft.Icons.CLOSE_FULLSCREEN
        ai_both_fullscreen_btn.tooltip = "Fermer le mode côte à côte"
        ai_fullscreen_btn.icon         = ft.Icons.FULLSCREEN
        notepad_fullscreen_btn.icon    = ft.Icons.FULLSCREEN
        page.update()



    def update_overlay_visibility():
        """Affiche ou masque l'overlay (IA à gauche + Notes à droite)."""
        panels_are_open = ai_mode["value"] or note_mode["value"]
        # Nettoyage du mode solo si les panneaux se ferment
        if not panels_are_open and overlay_fullscreen["mode"] in ("ai", "notepad", "ai_full", "notepad_full"):
            # Restaurer bottom_panel_container au mode normal
            bottom_panel_container.top    = None
            bottom_panel_container.right  = 0
            bottom_panel_container.left   = 0
            bottom_panel_container.width  = None
            bottom_panel_container.height = CONSTANTS.TERMINAL_HEIGHT
            ai_panel_container.visible      = True
            notepad_panel_container.visible = True
            _terminal_spacer.height = CONSTANTS.TERMINAL_HEIGHT
        overlay_container.visible = panels_are_open
        if not panels_are_open:
            overlay_fullscreen["mode"] = None
            ai_panel_container.visible = True
            notepad_panel_container.visible = True
            expand_button_overlay.icon = ft.Icons.OPEN_IN_FULL
            expand_button_notepad.icon = ft.Icons.OPEN_IN_FULL
            expand_button_overlay.tooltip = "IA seule (Ctrl/Cmd+←)"
            expand_button_notepad.tooltip = "Bloc-notes seul (Ctrl/Cmd+→)"
        if open_panels_button is not None:
            open_panels_button.icon       = ft.Icons.SMART_TOY
            open_panels_button.icon_color = RED if panels_are_open else BLUE
            open_panels_button.tooltip    = "Fermer IA & Notes" if panels_are_open else "Ouvrir IA & Bloc-notes"
        open_panels_card = open_panels_card_state["control"]
        if open_panels_card is not None:
            open_panels_card.bgcolor = GREY
            open_panels_card.border = ft.Border.all(1, RED if panels_are_open else BLUE)
            title_control = open_panels_card_state["title"]
            caption_control = open_panels_card_state["caption"]
            if title_control is not None:
                title_control.color = RED if panels_are_open else BLUE
            if caption_control is not None:
                caption_control.color = RED if panels_are_open else BLUE
            try:
                if open_panels_button is not None:
                    open_panels_button.update()
                if title_control is not None:
                    title_control.update()
                if caption_control is not None:
                    caption_control.update()
                open_panels_card.update()
            except Exception:
                pass



    def toggle_terminal_overlay():
        if overlay_fullscreen["mode"] in ("ai", "notepad", "ai_full", "notepad_full"):
            return
        terminal_is_expanded["value"] = not terminal_is_expanded["value"]
        is_expanded = terminal_is_expanded["value"]
        bottom_panel_container.height = _expanded_terminal_height() if is_expanded else CONSTANTS.TERMINAL_HEIGHT
        new_icon    = ft.Icons.CLOSE_FULLSCREEN if is_expanded else ft.Icons.VERTICAL_SPLIT
        new_tooltip = "Réduire  (Ctrl+↑)" if is_expanded else "Agrandir  (Ctrl+↑)"
        for expand_button in (expand_button_terminal,):
            expand_button.icon    = new_icon
            expand_button.tooltip = new_tooltip
        page.update()
        # Réaffirmer la visibilité de l'overlay après le page.update() pour éviter
        # que Flet ne la réinitialise à sa valeur initiale (False) lors du re-render.
        if ai_mode["value"] or note_mode["value"]:
            if overlay_container is not None:
                overlay_container.visible = True
                overlay_container.update()



    expand_button_terminal.on_click = lambda e: toggle_terminal_overlay()
    expand_button_overlay.on_click  = lambda e: toggle_ai_fullscreen()
    expand_button_notepad.on_click  = lambda e: toggle_notepad_fullscreen()



    def _open_bluetooth():
        if not _strip_state["active"]:
            _toggle_strip()
        if platform.system() == "Windows":
            subprocess.Popen(["fsquirt.exe", "/Receive"])
        else:
            subprocess.Popen(["open", "-a", "Bluetooth File Exchange"])



# ── Strip mode (réduction en bandeau pour écrans tactiles) ────────────────────
    _strip_state = {"active": False, "saved_height": CONSTANTS.WINDOW_HEIGHT, "was_maximized": False}
    _main_stack_ref = ft.Ref[ft.Stack]()
    strip_btn = ft.IconButton(
        icon=ft.Icons.UNFOLD_LESS,
        tooltip="Réduire en bandeau (écran tactile)",
        icon_color=LIGHT_GREY,
        bgcolor=DARK,
        icon_size=18,
    )



    def _toggle_strip(e=None):
        stack = _main_stack_ref.current
        is_mac = platform.system() == "Darwin"
        if not _strip_state["active"]:
            _strip_state["was_maximized"] = bool(page.window.maximized)
            _strip_state["saved_height"] = page.window.height or CONSTANTS.WINDOW_HEIGHT
            _strip_state["active"] = True
            # Sur macOS, changer window.height n'a aucun effet si la fenêtre est
            # maximisée : il faut d'abord la dé-maximiser explicitement.
            if is_mac and _strip_state["was_maximized"]:
                page.window.maximized = False
            stack.visible = False
            page.window.height = CONSTANTS.WDA_HEIGHT
            strip_btn.icon = ft.Icons.UNFOLD_MORE
            strip_btn.tooltip = "Restaurer la fenêtre"
            strip_btn.icon_color = BLUE
        else:
            _strip_state["active"] = False
            stack.visible = True
            if is_mac and _strip_state["was_maximized"]:
                page.window.maximized = True
            else:
                page.window.height = _strip_state["saved_height"]
            strip_btn.icon = ft.Icons.UNFOLD_LESS
            strip_btn.tooltip = "Réduire en bandeau (écran tactile)"
            strip_btn.icon_color = LIGHT_GREY
        page.update()



    strip_btn.on_click = _toggle_strip


# ===================== INTERFACE FLET ===================== #
    page.add(

        # ── Dessus ────────────────────────────
        ft.WindowDragArea(
            ft.Row([
                ft.Container(
                    ft.Text(f"DASHBOARD {__version__}", size=24, color=WHITE),
                    bgcolor=BACKGROUND,
                    padding=10,
                ),
                ft.IconButton(
                    icon=ft.Icons.SYSTEM_UPDATE_ALT,
                    tooltip="Mettre à jour (git pull --rebase)",
                    on_click=update_app,
                    icon_color=LIGHT_GREY,
                    bgcolor=DARK,
                    icon_size=18,
                ),
                ft.Container(
                    content=ft.Row(
                        [
                            ft.IconButton(
                                icon=ft.Icons.PHOTO_LIBRARY,
                                tooltip="Ouvrir le Kiosk",
                                on_click=lambda e: _launch_kiosk_flet(),
                                icon_color=VIOLET,
                                bgcolor=DARK,
                                icon_size=18,
                            ),
                            kiosk_tariff_btn,
                        ],
                        spacing=4,
                        tight=True,
                        vertical_alignment=ft.CrossAxisAlignment.CENTER,
                    ),
                    border=ft.Border.all(1, VIOLET),
                    border_radius=10,
                    padding=ft.Padding(4, 2, 8, 2),
                    margin=ft.Margin(6, 0, 6, 0),
                ),
                ft.IconButton(
                    icon=ft.Icons.VIEW_SIDEBAR,
                    tooltip="Ouvrir le Side Panel",
                    on_click=lambda e: _launch_side_panel(),
                    icon_color=BLUE,
                    bgcolor=DARK,
                    icon_size=18,
                ),
                ft.Container(expand=True),
                folder_path,
                recent_folders_btn,
                ft.IconButton(
                    icon=ft.Icons.FOLDER_OPEN,
                    icon_color=RED,
                    bgcolor=GREY,
                    tooltip="Parcourir…",
                    on_click=pick_folder,
                ),
                ft.IconButton(
                    icon=ft.Icons.REFRESH,
                    icon_color=BLUE,
                    bgcolor=GREY,
                    tooltip="Rafraîchir (Ctrl+R)",
                    on_click=lambda e: run_refresh_preview_command(),
                ),
                ft.IconButton(
                    icon=ft.Icons.OPEN_IN_NEW,
                    icon_color=GREEN,
                    bgcolor=GREY,
                    tooltip="Ouvrir l'explorateur",
                    on_click=lambda e: (open_in_file_explorer(current_browse_folder["path"] or selected_folder["path"]), _toggle_strip() if not _strip_state["active"] else None),
                ),
                ft.IconButton(
                    icon=ft.Icons.BLUETOOTH,
                    icon_color=ft.Colors.LIGHT_BLUE_300,
                    bgcolor=GREY,
                    tooltip="Recevoir un fichier via Bluetooth",
                    on_click=lambda e: _open_bluetooth(),
                ),
                ft.IconButton(
                    icon=ft.Icons.PRINT,
                    icon_color=ORANGE,
                    bgcolor=GREY,
                    tooltip="Imprimer les images sélectionnées",
                    on_click=lambda e: _print_files_with_default_app(selected_files),
                ),


                ft.Container(expand=True),
                strip_btn,
                ft.IconButton(
                    icon=ft.Icons.MINIMIZE, on_click=lambda e: setattr(page.window, 'minimized', True),),
                ft.IconButton(
                    icon=ft.Icons.FULLSCREEN,
                    on_click=lambda e: (setattr(page.window, 'maximized', not page.window.maximized), page.update()),
                    tooltip="Maximiser / Restaurer",
                ),
                ft.IconButton(ft.Icons.CLOSE, on_click=close_window),
            ])
        ),
        ft.Stack([
            ft.Column([
            ft.Divider(),
            ft.Row([
                ft.Column([

                    # ── Zone gauche ────────────────────────────
                    ft.Row([
                        ft.Container(
                            content=ft.Text("Applications disponibles", weight=ft.FontWeight.BOLD, size=14, color=WHITE),
                            margin=ft.Margin.only(top=10, bottom=10, left=10),
                        ),
                        ft.Container(width=32),  # Espacement entre le titre et les boutons
                        ft.IconButton(
                            icon=ft.Icons.KEYBOARD_DOUBLE_ARROW_LEFT_SHARP,
                            tooltip="Kiosk gauche",
                            on_click=lambda e: launch_app("Kiosk gauche.py", os.path.join(app_directory, "Data", "Kiosk gauche.py"), True),
                            icon_color=VIOLET,
                            bgcolor=DARK,
                            icon_size=18,
                        ),
                        ft.IconButton(
                            icon=ft.Icons.KEYBOARD_DOUBLE_ARROW_RIGHT_SHARP,
                            tooltip="Kiosk droite",
                            on_click=lambda e: launch_app("Kiosk droite.py", os.path.join(app_directory, "Data", "Kiosk droite.py"), True),
                            icon_color=VIOLET,
                            bgcolor=DARK,
                            icon_size=18,
                        ),
                        ft.IconButton(
                            icon=ft.Icons.AUTO_DELETE,
                            tooltip="Nettoyer anciens fichiers (> 60 jours)",
                            on_click=lambda e: launch_app("Nettoyer anciens fichiers.py", os.path.join(app_directory, "Data", "Nettoyer anciens fichiers.py"), True),
                            icon_color=ORANGE,
                            bgcolor=GREY,
                            icon_size=18,
                        ),
                    ]),

                    # ── Apps ────────────────────────────
                    ft.Row([
                        ft.Container(
                            content=apps_list,
                            expand=True,
                            border=ft.Border.all(1, GREY),
                            border_radius=8,
                            bgcolor=DARK,
                            padding=ft.Padding(8, 8, 8, 8),
                        ),
                        ft.Container(
                            content=ft.ListView(
                                controls=[quick_tools_col],
                                expand=True,
                                auto_scroll=False,
                            ),
                            bgcolor=DARK,
                            border=ft.Border.all(1, GREY),
                            border_radius=8,
                            padding=ft.Padding.symmetric(vertical=8, horizontal=4),
                            width=56,
                        ),
                    ], expand=True, spacing=8),
                ], expand=6),
                ft.Column([

                    # ── Zone droite ────────────────────────────
                    ft.Row([
                        ft.Text("Contenu du dossier", weight=ft.FontWeight.BOLD, size=14, color=WHITE, margin=ft.Margin.only(left=10)),
                        ft.IconButton(
                            icon=ft.Icons.ARROW_UPWARD,
                            tooltip="Dossier parent",
                            on_click=go_to_parent_folder,
                            icon_color=BLUE,
                            icon_size=20,
                        ),
                        filter_sel_btn,
                        select_toggle_button,
                        invert_selection_button,
                        select_same_date_button,
                        ft.IconButton(
                            icon=ft.Icons.DELETE_SWEEP,
                            tooltip="Supprimer les fichiers sélectionnés",
                            on_click=delete_selected_files,
                            icon_color=RED,
                            icon_size=20,
                        ),
                        ft.Container(expand=True),
                        file_count_text,
                        ft.Container(width=4),
                        prev_page_btn,
                        page_indicator_text,
                        next_page_btn,
                    ], spacing=5, vertical_alignment=ft.CrossAxisAlignment.CENTER, height=36),
                    ft.Container(
                        content=ft.Row([
                        search_active_row,
                        selection_count_text,
                        ft.Container(expand=True),
                        ft.IconButton(
                            icon=ft.Icons.CONTENT_COPY,
                            tooltip="Copier les fichiers sélectionnés (Ctrl+C)",
                            on_click=copy_selected_files,
                            icon_color=BLUE,
                            icon_size=18,
                        ),
                        ft.IconButton(
                            icon=ft.Icons.CONTENT_CUT,
                            tooltip="Couper les fichiers sélectionnés (Ctrl+X)",
                            on_click=cut_selected_files,
                            icon_color=ORANGE,
                            icon_size=18,
                        ),
                        ft.IconButton(
                            icon=ft.Icons.CONTENT_PASTE,
                            tooltip="Coller les fichiers (Ctrl+V)",
                            on_click=paste_files,
                            icon_color=YELLOW,
                            icon_size=18,
                        ),
                        ft.IconButton(
                            icon=ft.Icons.CREATE_NEW_FOLDER,
                            tooltip="Créer un nouveau dossier (Ctrl+N)",
                            on_click=create_new_folder,
                            icon_color=GREEN,
                            icon_size=18,
                        ),
                        sort_segment
                    ], spacing=5, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                        padding=ft.Padding(0, 5, 0, 0), # Ajustement pour aligner le dessus avec les zones à gauche 
                    ),

                    # ── Preview list (droite) ────────────────────────────
                    ft.Container(
                        content=ft.Stack([
                            preview_list,
                            preview_loading,
                        ]),
                        expand=True,
                        border=ft.Border.all(1, GREY),
                        border_radius=8,
                        bgcolor=DARK,
                    )
                ], expand=9)
            ], expand=True, spacing=8),
            _terminal_spacer,
        ], expand=True, spacing=8),
        bottom_panel_container,
        _solo_left_container,
        ], expand=True, ref=_main_stack_ref),
    )

    if CONSTANTS.MAXIMIZED:
        async def _delayed_maximize():
            await asyncio.sleep(0.15)
            if platform.system() == "Darwin":
                page.window.maximized = False
                page.update()
                await asyncio.sleep(0.05)
            page.window.maximized = True
            page.update()
        page.run_task(_delayed_maximize)



#############################################################
#                         DÉMARRAGE                         #
#############################################################
# Neutralise l'erreur asyncio Windows "ConnectionResetError: [WinError 10054]"
# qui apparaît lors de la fermeture des pipes des sous-processus.
# C'est un bug connu de la boucle ProactorEventLoop sous Windows — sans impact fonctionnel.
if sys.platform == "win32":
    _original_exception_handler = None

    def _silence_proactor_pipe_errors(loop, context):
        exc = context.get("exception")
        if isinstance(exc, (ConnectionResetError, BrokenPipeError)):
            return
        if _original_exception_handler:
            _original_exception_handler(loop, context)
        else:
            loop.default_exception_handler(context)

    _loop = asyncio.new_event_loop()
    _original_exception_handler = _loop.get_exception_handler()
    _loop.set_exception_handler(_silence_proactor_pipe_errors)
    asyncio.set_event_loop(_loop)



ft.run(main, assets_dir="assets")
